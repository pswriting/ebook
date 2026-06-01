# -*- coding: utf-8 -*-
import os
import logging

# 환경 변수
os.environ['PYTHONIOENCODING'] = 'utf-8'

# 로깅 비활성화
logging.getLogger('anthropic').setLevel(logging.ERROR)
logging.getLogger('httpx').setLevel(logging.ERROR)

import streamlit as st
import streamlit.components.v1 as components
import re
import json
import html
import base64
import urllib.parse
import uuid
import platform
import hashlib
import requests
from datetime import datetime, timedelta
from pathlib import Path

# Claude API
try:
    import anthropic
    CLAUDE_AVAILABLE = True
except ImportError:
    CLAUDE_AVAILABLE = False

# Gemini (이미지 생성용으로만 사용)
try:
    import google.generativeai as genai
    from google import genai as google_genai
    from google.genai import types as genai_types
    GEMINI_AVAILABLE = True
    IMAGEN_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False
    IMAGEN_AVAILABLE = False

# YouTube 자막 추출용
try:
    from youtube_transcript_api import YouTubeTranscriptApi
    YOUTUBE_TRANSCRIPT_AVAILABLE = True
except ImportError:
    YOUTUBE_TRANSCRIPT_AVAILABLE = False

# 브라우저 ID용 (클라우드 배포 시 필요)
try:
    from streamlit_javascript import st_javascript
    BROWSER_ID_AVAILABLE = True
except ImportError:
    BROWSER_ID_AVAILABLE = False

# 쿠키 매니저 (데이터 저장용)
try:
    import extra_streamlit_components as stx
    COOKIE_AVAILABLE = True
except ImportError:
    COOKIE_AVAILABLE = False

# Word 문서 생성용
try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_SECTION
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import io
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ==========================================
# 설정
# ==========================================
def get_config_path():
    return Path.home() / ".ebook_app_config.json"

def load_config():
    try:
        if get_config_path().exists():
            with open(get_config_path(), 'r') as f:
                return json.load(f)
    except:
        pass
    return {}

def save_config(data):
    try:
        config = load_config()
        config.update(data)
        with open(get_config_path(), 'w') as f:
            json.dump(config, f)
    except:
        pass

def load_saved_api_key():
    return load_config().get('api_key', '')

def save_api_key(api_key):
    save_config({'api_key': api_key})

def is_authenticated():
    return load_config().get('authenticated', False)

def save_authenticated():
    save_config({'authenticated': True})

# ==========================================
# 간단 비밀번호 인증
# ==========================================
CORRECT_PASSWORD = "cashmaker2024"  # ← 비밀번호 변경하려면 여기만 수정


# 비디오 배경용 base64 인코딩
@st.cache_data
def get_video_base64(video_path):
    try:
        with open(video_path, "rb") as f:
            return base64.b64encode(f.read()).decode()
    except:
        return None

st.set_page_config(page_title="Writey", layout="wide", page_icon="✍")

# 쿠키 매니저 초기화 및 데이터 불러오기/저장
# 승인 유효기간: 365일 + 롤링 만료 (방문할 때마다 자동 갱신)
COOKIE_LIFETIME_DAYS = 365  # 한 번 로그인 후 365일 유지. 방문할 때마다 자동 연장됨

if COOKIE_AVAILABLE:
    import time
    cookie_manager = stx.CookieManager(key="writey_cookies")
    cookies = cookie_manager.get_all()

    # [중요] CookieManager는 새 세션의 '첫 실행'에서는 브라우저와의 통신이 끝나기 전이라
    # 쿠키가 실제로 있어도 빈 dict({})를 반환한다. 그 상태로 자동 로그인을 판정하면
    # 항상 로그인 화면이 떠서 "로그인 정보가 기억되지 않는" 증상이 발생한다.
    # → 첫 실행에서 쿠키가 비어 있으면 딱 한 번만 재실행해서 쿠키를 읽을 기회를 준다.
    if not cookies and not st.session_state.get('_cookie_loaded'):
        st.session_state['_cookie_loaded'] = True
        time.sleep(0.3)
        st.rerun()
    st.session_state['_cookie_loaded'] = True

    # 쿠키에서 비밀번호/API키 복원
    if cookies:
        if 'writey_password' in cookies and cookies['writey_password']:
            if 'saved_password' not in st.session_state:
                st.session_state['saved_password'] = cookies['writey_password']
        if 'writey_api_key' in cookies and cookies['writey_api_key']:
            if 'saved_api_key' not in st.session_state:
                st.session_state['saved_api_key'] = cookies['writey_api_key']

    # 롤링 만료: 매 세션마다 만료일 연장
    if 'cookie_rolling_refreshed' not in st.session_state:
        if st.session_state.get('saved_password'):
            st.session_state['pending_save_password'] = st.session_state['saved_password']
        if st.session_state.get('saved_api_key'):
            st.session_state['pending_save_api'] = st.session_state['saved_api_key']
        st.session_state['cookie_rolling_refreshed'] = True

    # pending 값을 쿠키에 저장
    if 'pending_save_password' in st.session_state:
        cookie_manager.set('writey_password', st.session_state['pending_save_password'], expires_at=datetime.now() + timedelta(days=COOKIE_LIFETIME_DAYS))
        del st.session_state['pending_save_password']
    if 'pending_save_api' in st.session_state:
        cookie_manager.set('writey_api_key', st.session_state['pending_save_api'], expires_at=datetime.now() + timedelta(days=COOKIE_LIFETIME_DAYS))
        del st.session_state['pending_save_api']
else:
    cookie_manager = None

def save_password_to_browser(password):
    """비밀번호를 쿠키에 저장"""
    st.session_state['saved_password'] = password
    st.session_state['pending_save_password'] = password

def save_api_key_to_browser(api_key):
    """API 키를 쿠키에 저장"""
    st.session_state['saved_api_key'] = api_key
    st.session_state['pending_save_api'] = api_key

def get_saved_api_key():
    """저장된 API 키 반환"""
    return st.session_state.get('saved_api_key', None)

# ==========================================
# APPLE STYLE CSS
# ==========================================
st.markdown("""
<style>
@import url('https://cdn.jsdelivr.net/gh/orioncactus/pretendard/dist/web/static/pretendard.css');
@import url('https://fonts.googleapis.com/css2?family=Noto+Serif+KR:wght@400;500;600;700;900&display=swap');
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@400;500;600;700;800;900&display=swap');
@import url('https://fonts.googleapis.com/css2?family=Cormorant+Garamond:wght@400;500;600;700&display=swap');
@import url('https://fonts.googleapis.com/css2?family=Cinzel:wght@400;500;600;700&display=swap');

/* S-Core Dream 폰트 */
@font-face {
    font-family: 'S-CoreDream';
    src: url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_six@1.2/S-CoreDream-1Thin.woff') format('woff');
    font-weight: 100;
}
@font-face {
    font-family: 'S-CoreDream';
    src: url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_six@1.2/S-CoreDream-2ExtraLight.woff') format('woff');
    font-weight: 200;
}
@font-face {
    font-family: 'S-CoreDream';
    src: url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_six@1.2/S-CoreDream-3Light.woff') format('woff');
    font-weight: 300;
}
@font-face {
    font-family: 'S-CoreDream';
    src: url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_six@1.2/S-CoreDream-4Regular.woff') format('woff');
    font-weight: 400;
}
@font-face {
    font-family: 'S-CoreDream';
    src: url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_six@1.2/S-CoreDream-5Medium.woff') format('woff');
    font-weight: 500;
}
@font-face {
    font-family: 'S-CoreDream';
    src: url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_six@1.2/S-CoreDream-6Bold.woff') format('woff');
    font-weight: 600;
}
@font-face {
    font-family: 'S-CoreDream';
    src: url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_six@1.2/S-CoreDream-7ExtraBold.woff') format('woff');
    font-weight: 700;
}
@font-face {
    font-family: 'S-CoreDream';
    src: url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_six@1.2/S-CoreDream-8Heavy.woff') format('woff');
    font-weight: 800;
}
@font-face {
    font-family: 'S-CoreDream';
    src: url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_six@1.2/S-CoreDream-9Black.woff') format('woff');
    font-weight: 900;
}

:root {
    --gold: #C9A24B;
    --gold-light: #E0C074;
    --gold-dark: #A8852F;
    --rose-gold: #b76e79;
    --cream: #FAF8F4;
    --charcoal: #141416;
    --dark: #0B0B0D;
    --card: rgba(255,255,255,0.025);
    --card2: rgba(255,255,255,0.05);
    --text: #F5F3EF;
    --text2: #8A8780;
    --text3: #7A776F;
    --line: rgba(201,162,75,0.18);
    --line2: rgba(255,255,255,0.06);
    --glow: rgba(201,162,75,0.32);
    --success: #6FA86F;
    --warning: #E0C074;
    --danger: #C97A6F;
}

/* 애니메이션 정의 */
@keyframes fadeInUp {
    from { opacity: 0; transform: translateY(30px); }
    to { opacity: 1; transform: translateY(0); }
}
@keyframes shimmer {
    0% { background-position: -200% 0; }
    100% { background-position: 200% 0; }
}
@keyframes pulse {
    0%, 100% { opacity: 1; transform: scale(1); }
    50% { opacity: 0.8; transform: scale(1.02); }
}
@keyframes borderGlow {
    0%, 100% { box-shadow: 0 0 5px var(--glow), inset 0 0 5px rgba(201,162,75,0.1); }
    50% { box-shadow: 0 0 20px var(--glow), inset 0 0 10px rgba(201,162,75,0.2); }
}
@keyframes float {
    0%, 100% { transform: translateY(0); }
    50% { transform: translateY(-5px); }
}
@keyframes goldShine {
    0% { background-position: -100% 0; }
    100% { background-position: 200% 0; }
}

*:not([data-testid*="Icon"]):not(.material-icons):not([class*="icon"]):not(span[aria-hidden="true"]) {
    font-family: 'S-CoreDream', 'Pretendard', -apple-system, sans-serif !important;
}
/* 아이콘 폰트 복원 */
[data-testid*="Icon"], .material-icons, span[aria-hidden="true"], button[kind="header"] span {
    font-family: 'Material Symbols Rounded', 'Material Icons', sans-serif !important;
}
.stDeployButton, footer, #MainMenu { display: none !important; }
/* 헤더 투명하게 (사이드바 버튼은 보임) */
header[data-testid="stHeader"] {
    background: transparent !important;
}

/* 럭셔리 배경 - 미세한 그라데이션 */
.stApp {
    background:
        radial-gradient(ellipse at 20% 0%, rgba(201,162,75,0.04) 0%, transparent 55%),
        radial-gradient(ellipse at 80% 100%, rgba(201,162,75,0.025) 0%, transparent 55%),
        linear-gradient(180deg, #0B0B0D 0%, #08080A 50%, #0B0B0D 100%) !important;
    background-attachment: fixed;
}

.main .block-container { max-width: 1000px; padding: 3rem 2rem; }

/* 사이드바 - 미니멀 */
[data-testid="stSidebar"] {
    background: var(--charcoal) !important;
    border-right: 1px solid var(--line);
}
[data-testid="stSidebar"] * { color: var(--text2) !important; }

/* 타이포그래피 - 가독성 향상 */
h1, h2, h3 { color: var(--text) !important; font-weight: 300 !important; letter-spacing: 0.5px; }
h1 { font-size: 34px !important; color: var(--cream) !important; font-weight: 300 !important; }
h2 { font-size: 26px !important; margin-bottom: 20px !important; font-weight: 300 !important; }
h3 { font-size: 21px !important; color: var(--gold) !important; font-weight: 400 !important; }
p, span, label, div { color: var(--text) !important; font-size: 16px !important; line-height: 1.7 !important; }
li { font-size: 16px !important; line-height: 1.8 !important; }

/* 버튼 - 채워진 골드 그라데이션 (첨부 디자인) */
.stButton > button {
    background: linear-gradient(135deg, #E0C074 0%, #C9A24B 100%) !important;
    color: #0B0B0D !important;
    -webkit-text-fill-color: #0B0B0D !important;
    border: none !important;
    border-radius: 12px;
    font-weight: 600;
    font-size: 15px !important;
    padding: 15px 36px;
    letter-spacing: 0.4px;
    text-transform: none;
    transition: all 0.35s cubic-bezier(0.4, 0, 0.2, 1);
    position: relative;
    overflow: hidden;
    box-shadow: 0 6px 20px rgba(201,162,75,0.22);
}
.stButton > button * {
    color: #0B0B0D !important;
    -webkit-text-fill-color: #0B0B0D !important;
}
.stButton > button::before {
    content: '';
    position: absolute;
    top: 0;
    left: -100%;
    width: 100%;
    height: 100%;
    background: linear-gradient(90deg, transparent, rgba(255,255,255,0.45), transparent);
    transition: left 0.6s ease;
}
.stButton > button:hover::before {
    left: 100%;
}
.stButton > button:hover {
    background: linear-gradient(135deg, #EBCE86 0%, #D4AC56 100%) !important;
    box-shadow: 0 10px 32px rgba(201,162,75,0.4);
    transform: translateY(-2px);
}
.stButton > button:active {
    transform: translateY(0);
    box-shadow: 0 4px 15px rgba(201,162,75,0.3);
}

/* 입력 필드 - 밝은 배경 + 검은 글씨 */
.stTextInput input, .stTextArea textarea, .stNumberInput input {
    background: #ffffff !important;
    background-color: #ffffff !important;
    border: 0.5px solid var(--line) !important;
    border-radius: 10px !important;
    color: #000000 !important;
    -webkit-text-fill-color: #000000 !important;
    caret-color: #000000 !important;
    padding: 18px !important;
    font-size: 17px !important;
}
.stTextInput input:focus, .stTextArea textarea:focus, .stNumberInput input:focus {
    border-color: var(--gold) !important;
    box-shadow: 0 0 0 2px rgba(201,162,75,0.2) !important;
}

/* 셀렉트박스 컨테이너 */
.stSelectbox > div > div {
    background: var(--card) !important;
    border: 0.5px solid var(--line) !important;
    border-radius: 10px;
}
/* 셀렉트박스 선택된 값 - 흰색 */
.stSelectbox [data-baseweb="select"] > div {
    color: #ffffff !important;
    -webkit-text-fill-color: #ffffff !important;
}

/* 스코어 카드 - 럭셔리 */
.score-card {
    background: linear-gradient(145deg, var(--card) 0%, rgba(30,30,30,0.95) 100%) !important;
    border: 0.5px solid var(--gold);
    border-radius: 20px;
    padding: 50px 40px;
    text-align: center;
    animation: fadeInUp 0.6s ease-out;
    transition: all 0.5s cubic-bezier(0.4, 0, 0.2, 1);
    position: relative;
    overflow: hidden;
    box-shadow: 0 10px 40px rgba(201,162,75,0.15);
}
.score-card::before {
    content: '';
    position: absolute;
    top: 0;
    left: 0;
    right: 0;
    height: 3px;
    background: linear-gradient(90deg, transparent, var(--gold), transparent);
    opacity: 1;
}
.score-card:hover {
    border-color: var(--gold);
    box-shadow: 0 20px 60px rgba(201,162,75,0.3), inset 0 1px 0 rgba(201,162,75,0.1);
    transform: translateY(-5px);
}
.score-card:hover::before {
    opacity: 1;
}
.score-number {
    font-size: 140px;
    font-weight: 300;
    background: linear-gradient(135deg, var(--gold-light) 0%, var(--gold) 50%, var(--gold-dark) 100%);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
    line-height: 1;
    letter-spacing: -4px;
    animation: fadeInUp 0.8s ease-out;
    filter: drop-shadow(0 2px 4px rgba(201,162,75,0.3));
}

/* 정보 카드 + 애니메이션 */
.info-card {
    background: transparent !important;
    border: none;
    border-left: 2px solid var(--gold);
    padding: 20px 24px;
    margin: 20px 0;
    animation: fadeInUp 0.5s ease-out;
    transition: all 0.3s ease;
}
.info-card:hover {
    background: rgba(201,169,98,0.05) !important;
    border-left-width: 4px;
    padding-left: 22px;
}

/* 스탯 박스 + 애니메이션 */
.stat-box {
    background: var(--card) !important;
    border: 0.5px solid var(--line2);
    border-radius: 14px;
    padding: 32px;
    text-align: center;
    animation: fadeInUp 0.5s ease-out;
    transition: all 0.4s ease;
}
.stat-box:hover {
    transform: translateY(-4px);
    box-shadow: 0 8px 25px rgba(0,0,0,0.3);
    border-color: var(--gold);
}
.stat-value {
    font-size: 42px;
    font-weight: 200;
    color: var(--gold) !important;
    letter-spacing: -2px;
    transition: transform 0.3s ease;
}
.stat-box:hover .stat-value {
    transform: scale(1.05);
}
.stat-label {
    font-size: 11px;
    color: var(--text2) !important;
    margin-top: 12px;
    text-transform: uppercase;
    letter-spacing: 3px;
}

/* 데이터 카드 + 애니메이션 */
.data-card {
    background: var(--card) !important;
    border-left: 2px solid var(--gold);
    padding: 20px 24px;
    margin: 16px 0;
    animation: fadeInUp 0.4s ease-out;
    transition: all 0.3s ease;
}
.data-card:hover {
    border-left-width: 4px;
    background: var(--card2) !important;
}

/* 서머리 허브 + 애니메이션 */
.summary-hub {
    background: var(--card) !important;
    border: 0.5px solid var(--line2);
    border-radius: 14px;
    padding: 40px;
    animation: fadeInUp 0.5s ease-out;
    transition: all 0.4s ease;
}
.summary-hub:hover {
    border-color: var(--gold);
}

/* 배지 - 미니멀 + 펄스 */
.verdict-go {
    background: transparent !important;
    color: var(--success) !important;
    border: 1px solid var(--success);
    padding: 12px 32px;
    border-radius: 20px;
    font-weight: 400;
    font-size: 12px;
    letter-spacing: 3px;
    text-transform: uppercase;
    animation: fadeInUp 0.6s ease-out;
}
.verdict-wait {
    background: transparent !important;
    color: var(--warning) !important;
    border: 1px solid var(--warning);
    padding: 12px 32px;
    border-radius: 20px;
    font-weight: 400;
    font-size: 12px;
    letter-spacing: 3px;
    animation: fadeInUp 0.6s ease-out;
}
.verdict-no {
    background: transparent !important;
    color: var(--danger) !important;
    border: 1px solid var(--danger);
    padding: 12px 32px;
    border-radius: 20px;
    font-weight: 400;
    font-size: 12px;
    letter-spacing: 3px;
    animation: fadeInUp 0.6s ease-out;
}

/* 네비게이션 */
.premium-nav-container {
    background: transparent;
    border-top: 1px solid var(--line);
    border-bottom: 1px solid var(--line);
    padding: 0;
    margin-bottom: 48px;
}
.nav-item {
    padding: 18px 12px;
    text-align: center;
    font-size: 14px;
    color: var(--text2);
    letter-spacing: 1px;
    transition: all 0.3s ease;
}
.nav-item.active {
    background: linear-gradient(135deg, rgba(201,162,75,0.2) 0%, rgba(201,162,75,0.1) 100%);
    color: var(--gold) !important;
    font-weight: 600;
    border-bottom: 3px solid var(--gold);
    box-shadow: 0 4px 15px rgba(201,162,75,0.2);
}

/* 섹션 타이틀 - 미니멀 (첨부 디자인) */
.section-title-box {
    background: rgba(255,255,255,0.025);
    border: 0.5px solid var(--line2);
    border-radius: 14px;
    padding: 30px 36px;
    margin-bottom: 35px;
    text-align: left;
    position: relative;
    overflow: hidden;
    animation: fadeInUp 0.5s ease-out;
}
.section-title-box h2 {
    font-size: 26px !important;
    color: #FAF8F4 !important;
    margin: 0 0 8px 0 !important;
    font-weight: 300 !important;
    letter-spacing: 0.5px;
}
.section-title-box p {
    color: var(--text2) !important;
    font-size: 14px !important;
    margin: 0 !important;
}
.section-step {
    display: inline-block;
    background: transparent;
    color: var(--gold) !important;
    font-size: 11px;
    font-weight: 500;
    padding: 0;
    margin-bottom: 12px;
    letter-spacing: 0.22em;
}

/* 제목 카드 + 애니메이션 */
.title-card {
    background: var(--card);
    border: 0.5px solid var(--line2);
    border-radius: 12px;
    padding: 28px;
    margin: 16px 0;
    transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1);
    animation: fadeInUp 0.5s ease-out;
}
.title-card:hover {
    border-color: var(--gold);
    background: rgba(201,169,98,0.05);
    transform: translateX(8px);
    box-shadow: -4px 0 20px rgba(201,169,98,0.15);
}
.title-main {
    font-size: 18px;
    font-weight: 400;
    color: var(--text) !important;
    letter-spacing: 1px;
    transition: color 0.3s ease;
}
.title-card:hover .title-main {
    color: var(--gold) !important;
}
.title-sub {
    font-size: 13px;
    color: var(--text2) !important;
    margin-top: 8px;
}

/* 로그인 - 럭셔리 */
.login-card {
    max-width: 420px;
    margin: 100px auto;
    padding: 70px 50px;
    background: linear-gradient(145deg, rgba(26,26,31,0.98) 0%, rgba(11,11,13,0.98) 100%);
    border: 0.5px solid var(--line);
    border-radius: 16px;
    text-align: center;
    animation: fadeInUp 0.8s ease-out;
    position: relative;
    box-shadow: 0 25px 80px rgba(0,0,0,0.5), 0 0 40px rgba(201,162,75,0.05);
}
.login-card::before {
    content: '';
    position: absolute;
    top: -1px;
    left: 20%;
    right: 20%;
    height: 2px;
    background: linear-gradient(90deg, transparent, var(--gold), transparent);
}
.login-card::after {
    content: '';
    position: absolute;
    bottom: -1px;
    left: 20%;
    right: 20%;
    height: 1px;
    background: linear-gradient(90deg, transparent, var(--gold-dark), transparent);
}
.login-title {
    font-size: 32px;
    font-weight: 300;
    background: linear-gradient(135deg, var(--cream) 0%, var(--gold-light) 50%, var(--gold) 100%);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
    letter-spacing: 7px;
    animation: fadeInUp 1s ease-out;
}
.login-subtitle {
    font-size: 11px;
    color: var(--gold) !important;
    margin-top: 20px;
    letter-spacing: 4px;
    text-transform: uppercase;
    animation: fadeInUp 1.2s ease-out;
    opacity: 0.8;
}

/* 헤더 - 럭셔리 */
.main-header {
    text-align: center;
    padding: 80px 20px 60px;
    margin-bottom: 50px;
    border-bottom: 1px solid var(--line);
    animation: fadeInUp 0.6s ease-out;
    position: relative;
    background: linear-gradient(180deg, rgba(201,162,75,0.02) 0%, transparent 100%);
}
.main-header::after {
    content: '';
    position: absolute;
    bottom: -1px;
    left: 10%;
    right: 10%;
    height: 1px;
    background: linear-gradient(90deg, transparent, var(--gold), transparent);
}
.main-header-brand {
    font-size: 11px;
    color: var(--gold) !important;
    letter-spacing: 10px;
    text-transform: uppercase;
    animation: fadeInUp 0.8s ease-out;
    text-shadow: 0 0 20px rgba(201,162,75,0.3);
}
.main-header-title {
    font-size: 42px;
    font-weight: 200;
    background: linear-gradient(135deg, var(--cream) 0%, var(--gold-light) 100%);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
    letter-spacing: 8px;
    margin-top: 24px;
    animation: fadeInUp 1s ease-out;
}
.header-tagline {
    font-size: 13px;
    color: var(--text2) !important;
    margin-top: 24px;
    letter-spacing: 3px;
    animation: fadeInUp 1.2s ease-out;
}

/* Expander + 애니메이션 */
.stExpander {
    background: var(--card) !important;
    border: 0.5px solid var(--line2) !important;
    border-radius: 12px !important;
    animation: fadeInUp 0.4s ease-out;
    transition: border-color 0.3s ease;
}
.stExpander:hover {
    border-color: var(--gold) !important;
}
.stProgress > div > div > div {
    background: linear-gradient(90deg, var(--gold-dark), var(--gold), var(--gold-light), var(--gold), var(--gold-dark)) !important;
    background-size: 300% 100%;
    animation: goldShine 3s ease infinite;
    border-radius: 4px;
    box-shadow: 0 0 15px rgba(201,162,75,0.4);
}
.stProgress > div > div {
    background: rgba(20,20,20,0.8);
    border-radius: 4px;
    border: 1px solid var(--line);
}

/* 라디오 & 탭 */
.stRadio > div { background: transparent; border: 1px solid var(--line); padding: 16px; }
.stTabs [data-baseweb="tab-list"] { background: transparent; border-bottom: 1px solid var(--line); }
.stTabs [aria-selected="true"] {
    background: transparent !important;
    color: var(--gold) !important;
    border-bottom: 2px solid var(--gold) !important;
}

/* 알림 */
.stSuccess > div { background: rgba(111,168,111,0.1) !important; border: 0.5px solid rgba(111,168,111,0.3) !important; border-radius: 10px; }
.stWarning > div { background: rgba(224,192,116,0.1) !important; border: 0.5px solid rgba(224,192,116,0.3) !important; border-radius: 10px; }
.stError > div { background: rgba(201,122,111,0.1) !important; border: 0.5px solid rgba(201,122,111,0.3) !important; border-radius: 10px; }
.stInfo > div { background: rgba(201,162,75,0.08) !important; border: 0.5px solid var(--line) !important; border-radius: 10px; }

/* 스크롤바 */
::-webkit-scrollbar { width: 6px; }
::-webkit-scrollbar-track { background: var(--dark); }
::-webkit-scrollbar-thumb { background: var(--gold-dark); }

/* 다운로드 버튼 - 럭셔리 골드 */
.stDownloadButton button {
    background: linear-gradient(135deg, var(--gold) 0%, var(--gold-dark) 50%, var(--gold) 100%) !important;
    background-size: 200% 100%;
    color: var(--dark) !important;
    border: none !important;
    border-radius: 12px;
    font-weight: 600;
    letter-spacing: 1.5px;
    box-shadow: 0 4px 20px rgba(201,162,75,0.3);
    transition: all 0.4s ease;
    text-shadow: 0 1px 1px rgba(255,255,255,0.2);
}
.stDownloadButton button:hover {
    background-position: 100% 0 !important;
    box-shadow: 0 8px 35px rgba(201,162,75,0.5);
    transform: translateY(-2px);
}

/* 구분선 */
hr { border: none; height: 1px; background: var(--line); margin: 40px 0; }

/* 표지 미리보기 - 실제 책처럼 */
.book-wrapper {
    perspective: 1000px;
    display: flex;
    justify-content: center;
    padding: 30px;
    background: linear-gradient(135deg, #1a1a1a 0%, #2d2d2d 100%);
    border-radius: 8px;
}
.ebook-cover {
    font-family: 'Pretendard', sans-serif !important;
    box-shadow:
        0 0 5px rgba(0,0,0,0.3),
        5px 5px 15px rgba(0,0,0,0.4),
        10px 10px 30px rgba(0,0,0,0.3),
        15px 15px 50px rgba(0,0,0,0.2),
        inset -3px 0 10px rgba(0,0,0,0.2);
    transform: rotateY(-3deg);
    border-radius: 0 3px 3px 0;
    position: relative;
}
.ebook-cover::before {
    content: '';
    position: absolute;
    left: 0;
    top: 0;
    bottom: 0;
    width: 25px;
    background: linear-gradient(90deg,
        rgba(0,0,0,0.4) 0%,
        rgba(0,0,0,0.1) 30%,
        rgba(255,255,255,0.05) 50%,
        rgba(0,0,0,0.1) 70%,
        rgba(0,0,0,0.3) 100%);
    border-radius: 3px 0 0 3px;
}
.ebook-cover::after {
    content: '';
    position: absolute;
    right: 0;
    top: 2px;
    bottom: 2px;
    width: 8px;
    background: linear-gradient(90deg,
        rgba(255,255,255,0.03) 0%,
        rgba(255,255,255,0.08) 50%,
        rgba(0,0,0,0.1) 100%);
}
.ebook-cover * {
    color: inherit !important;
    -webkit-text-fill-color: inherit !important;
}

/* ============================================
   입력 필드 텍스트 색상 - 최우선 적용
   ============================================ */

/* 모든 입력 필드 - 흰 배경 + 검은 글씨 */
.stTextInput input,
.stTextArea textarea,
.stNumberInput input,
[data-testid="stTextInput"] input,
[data-testid="stTextArea"] textarea,
[data-testid="stNumberInput"] input {
    background: #ffffff !important;
    background-color: #ffffff !important;
    color: #000000 !important;
    -webkit-text-fill-color: #000000 !important;
    caret-color: #000000 !important;
}

/* Placeholder 색상 */
input::placeholder,
textarea::placeholder {
    color: #888888 !important;
    -webkit-text-fill-color: #888888 !important;
}

/* 셀렉트박스 - 선택된 값 (어두운 배경에 흰 글씨) */
.stSelectbox [data-baseweb="select"] > div,
.stSelectbox [data-baseweb="select"] span,
.stSelectbox > div > div > div {
    color: #ffffff !important;
    -webkit-text-fill-color: #ffffff !important;
}

/* ============================================
   드롭다운/팝오버 - 검은 글씨 (흰 배경)
   ============================================ */
[data-baseweb="popover"],
[data-baseweb="popover"] *,
[data-baseweb="menu"],
[data-baseweb="menu"] *,
[data-baseweb="list"],
[data-baseweb="list"] *,
[role="listbox"],
[role="listbox"] *,
[role="option"],
[role="option"] *,
.stSelectbox ul,
.stSelectbox ul *,
.stSelectbox li,
.stSelectbox li * {
    background: #ffffff !important;
    background-color: #ffffff !important;
    color: #000000 !important;
    -webkit-text-fill-color: #000000 !important;
}

/* 드롭다운 옵션 호버 */
[role="option"]:hover,
[data-baseweb="menu"] li:hover,
.stSelectbox li:hover {
    background: #f0f0f0 !important;
    background-color: #f0f0f0 !important;
}

/* select 요소 */
select,
select option {
    color: #000000 !important;
    background: #ffffff !important;
}

/* Expander 스타일 정리 */
.stExpander details summary {
    background: var(--card) !important;
    overflow: hidden !important;
}
/* 모든 텍스트 숨기기 (keyboard_arrow 등 영어 텍스트 포함) */
.stExpander details summary * {
    font-size: 0 !important;
    color: transparent !important;
    -webkit-text-fill-color: transparent !important;
}
/* 한국어 제목만 보이게 */
.stExpander details summary p {
    font-size: 15px !important;
    color: var(--text) !important;
    -webkit-text-fill-color: var(--text) !important;
}
/* 화살표 아이콘만 보이게 */
.stExpander details summary svg {
    width: 20px !important;
    height: 20px !important;
    color: var(--gold) !important;
    fill: var(--gold) !important;
}

/* 버튼 앞 불필요한 라벨 숨기기 */
.stButton > div:not([data-testid="baseButton-secondary"]):not([data-testid="baseButton-primary"]) > p,
.stButton > div > div > p:first-child:not(:last-child),
.stButton label,
.stExpander .stButton > div:first-child > p {
    display: none !important;
}
/* 링크버튼 라벨 숨기기 */
.stLinkButton > div:first-child > p {
    display: none !important;
}
</style>
""", unsafe_allow_html=True)

st.markdown("""
<style>
/* ============================================
   WRITEY PRO UI OVERRIDE
   더 차분하고 고급스러운 제작 도구형 인터페이스
   ============================================ */

@import url('https://cdn.jsdelivr.net/gh/orioncactus/pretendard/dist/web/static/pretendard.css');
@import url('https://fonts.googleapis.com/css2?family=Noto+Serif+KR:wght@500;600;700;900&display=swap');

:root {
    --bg: #08090c;
    --bg-soft: #101217;
    --surface: rgba(16, 18, 24, 0.78);
    --surface-2: rgba(22, 25, 33, 0.90);
    --surface-3: rgba(246, 242, 232, 0.075);
    --text: #f5f1e8;
    --text-bright: #fffaf0;
    --text-dim: #a6a096;
    --text-muted: #746f67;
    --line: rgba(224, 211, 184, 0.18);
    --line-strong: rgba(224, 211, 184, 0.36);
    --gold: #d7b86a;
    --gold-light: #f0d58b;
    --gold-dark: #a98134;
    --accent: #7bd3c8;
    --accent-2: #ee846f;
    --blue: #8ea7ff;
    --success: #7bc99e;
    --warning: #f0c66b;
    --danger: #ee7f73;
    --card: rgba(255, 255, 255, 0.055);
    --card2: rgba(255, 255, 255, 0.085);
    --dark: #0f1115;
    --radius: 8px;
    --shadow: 0 24px 70px rgba(0, 0, 0, 0.44);
}

html, body, [class*="css"] {
    font-family: 'Pretendard', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif !important;
}

.stApp, .stApp *:not(svg):not(path) {
    font-family: 'Pretendard', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif !important;
}

[data-testid*="Icon"],
.material-icons,
span[aria-hidden="true"],
button[kind="header"] span,
svg,
svg * {
    font-family: initial !important;
}

.stDeployButton, footer, #MainMenu {
    display: none !important;
}

header[data-testid="stHeader"] {
    background: transparent !important;
}

.stApp {
    color: var(--text) !important;
    background:
        radial-gradient(circle at 16% -8%, rgba(215, 184, 106, 0.16), transparent 32%),
        radial-gradient(circle at 86% 8%, rgba(123, 211, 200, 0.12), transparent 30%),
        linear-gradient(rgba(255, 255, 255, 0.018) 1px, transparent 1px),
        linear-gradient(90deg, rgba(255, 255, 255, 0.013) 1px, transparent 1px),
        linear-gradient(180deg, #0c0d11 0%, #07080b 48%, #0d0f14 100%) !important;
    background-size: auto, auto, 64px 64px, 64px 64px, auto;
    background-attachment: fixed;
}

.main .block-container {
    max-width: 1240px !important;
    padding: 1.2rem 2rem 4rem !important;
}

.stApp div,
.stApp p,
.stApp span,
.stApp label,
.stApp li {
    line-height: 1.62 !important;
    letter-spacing: 0 !important;
}

h1, h2, h3, h4 {
    color: var(--text-bright) !important;
    letter-spacing: 0 !important;
}

h1 {
    font-size: clamp(30px, 4vw, 46px) !important;
    font-weight: 760 !important;
}

h2 {
    font-size: clamp(24px, 3vw, 34px) !important;
    font-weight: 740 !important;
}

h3 {
    font-size: 20px !important;
    font-weight: 700 !important;
    color: var(--text-bright) !important;
}

h4 {
    font-size: 17px !important;
    font-weight: 700 !important;
}

p, li, label {
    color: var(--text) !important;
}

small, .stCaption, [data-testid="stCaptionContainer"] {
    color: var(--text-dim) !important;
}

hr {
    border: none !important;
    height: 1px !important;
    background: var(--line) !important;
    margin: 28px 0 !important;
}

/* Sidebar */
[data-testid="stSidebar"] {
    background:
        linear-gradient(180deg, rgba(25, 28, 36, 0.98) 0%, rgba(14, 16, 22, 0.98) 100%) !important;
    border-right: 1px solid var(--line) !important;
}

[data-testid="stSidebar"] > div {
    padding-top: 1.4rem !important;
}

[data-testid="stSidebar"] h3 {
    color: var(--text-bright) !important;
    font-size: 15px !important;
    font-weight: 760 !important;
    margin: 16px 0 10px !important;
}

[data-testid="stSidebar"] p,
[data-testid="stSidebar"] span,
[data-testid="stSidebar"] label,
[data-testid="stSidebar"] div {
    color: var(--text-dim) !important;
}

[data-testid="stSidebar"] hr {
    margin: 22px 0 !important;
}

/* Forms */
.stTextInput input,
.stTextArea textarea,
.stNumberInput input,
[data-testid="stTextInput"] input,
[data-testid="stTextArea"] textarea,
[data-testid="stNumberInput"] input {
    background: #f7f4ec !important;
    border: 1px solid transparent !important;
    border-radius: var(--radius) !important;
    color: #14171d !important;
    -webkit-text-fill-color: #14171d !important;
    caret-color: #14171d !important;
    font-size: 15px !important;
    padding: 13px 14px !important;
    box-shadow: inset 0 0 0 1px rgba(15, 17, 21, 0.08), 0 1px 0 rgba(255,255,255,0.08) !important;
}

.stTextArea textarea {
    min-height: 92px;
}

.stTextInput input:focus,
.stTextArea textarea:focus,
.stNumberInput input:focus {
    border-color: var(--accent) !important;
    box-shadow: inset 0 0 0 1px rgba(15, 17, 21, 0.08), 0 0 0 3px rgba(123, 211, 200, 0.25) !important;
}

input::placeholder,
textarea::placeholder {
    color: #8d877d !important;
    -webkit-text-fill-color: #8d877d !important;
}

.stSelectbox [data-baseweb="select"] > div {
    min-height: 46px !important;
    background: rgba(247, 244, 236, 0.96) !important;
    border: 1px solid transparent !important;
    border-radius: var(--radius) !important;
    box-shadow: inset 0 0 0 1px rgba(15, 17, 21, 0.08) !important;
}

.stSelectbox [data-baseweb="select"] span,
.stSelectbox [data-baseweb="select"] div {
    color: #14171d !important;
    -webkit-text-fill-color: #14171d !important;
}

[data-baseweb="popover"],
[data-baseweb="menu"],
[role="listbox"] {
    border-radius: var(--radius) !important;
    overflow: hidden !important;
    box-shadow: 0 18px 45px rgba(0,0,0,0.28) !important;
}

[role="option"],
[role="option"] * {
    background: #fbfaf6 !important;
    color: #14171d !important;
    -webkit-text-fill-color: #14171d !important;
}

[role="option"]:hover,
[aria-selected="true"] {
    background: #efe8d8 !important;
}

.stCheckbox label,
.stRadio label {
    color: var(--text) !important;
}

/* Buttons */
.stButton > button,
.stFormSubmitButton > button,
.stDownloadButton > button,
.stLinkButton a {
    min-height: 46px !important;
    border-radius: var(--radius) !important;
    border: 1px solid var(--line-strong) !important;
    background: rgba(255, 255, 255, 0.055) !important;
    color: var(--text-bright) !important;
    -webkit-text-fill-color: var(--text-bright) !important;
    font-size: 14px !important;
    font-weight: 720 !important;
    letter-spacing: 0 !important;
    box-shadow: none !important;
    transition: border-color 160ms ease, background 160ms ease, transform 160ms ease, box-shadow 160ms ease !important;
}

.stButton > button *,
.stFormSubmitButton > button *,
.stDownloadButton > button *,
.stLinkButton a * {
    color: inherit !important;
    -webkit-text-fill-color: inherit !important;
}

.stLinkButton > div:first-child > p,
.stButton label {
    display: block !important;
}

.stButton > button::before,
.stFormSubmitButton > button::before,
.stDownloadButton > button::before {
    display: none !important;
}

.stButton > button:hover,
.stFormSubmitButton > button:hover,
.stLinkButton a:hover {
    transform: translateY(-1px) !important;
    background: rgba(255, 255, 255, 0.095) !important;
    border-color: rgba(123, 211, 200, 0.56) !important;
    box-shadow: 0 10px 26px rgba(0, 0, 0, 0.22) !important;
}

.stButton > button[kind="primary"],
.stFormSubmitButton > button[kind="primary"],
[data-testid="baseButton-primary"] {
    background: linear-gradient(135deg, #f0d58b 0%, #d7b86a 54%, #b58a38 100%) !important;
    color: #101216 !important;
    -webkit-text-fill-color: #101216 !important;
    border-color: rgba(240, 213, 139, 0.72) !important;
    box-shadow: 0 14px 34px rgba(215, 184, 106, 0.2) !important;
}

.stDownloadButton > button {
    background: linear-gradient(135deg, #8be0d4 0%, #62bfb3 100%) !important;
    color: #081412 !important;
    -webkit-text-fill-color: #081412 !important;
    border-color: rgba(139, 224, 212, 0.72) !important;
}

.stButton > button:disabled,
.stDownloadButton > button:disabled {
    opacity: 0.48 !important;
    transform: none !important;
    box-shadow: none !important;
}

/* Streamlit status */
.stSuccess > div,
.stWarning > div,
.stError > div,
.stInfo > div {
    border-radius: var(--radius) !important;
    border: 1px solid var(--line) !important;
    box-shadow: none !important;
}

.stSuccess > div {
    background: rgba(123, 201, 158, 0.12) !important;
    border-color: rgba(123, 201, 158, 0.34) !important;
}

.stWarning > div {
    background: rgba(240, 198, 107, 0.12) !important;
    border-color: rgba(240, 198, 107, 0.34) !important;
}

.stError > div {
    background: rgba(238, 127, 115, 0.12) !important;
    border-color: rgba(238, 127, 115, 0.34) !important;
}

.stInfo > div {
    background: rgba(142, 167, 255, 0.12) !important;
    border-color: rgba(142, 167, 255, 0.30) !important;
}

/* Tabs, expanders, progress */
.stTabs [data-baseweb="tab-list"] {
    gap: 6px !important;
    background: rgba(255,255,255,0.045) !important;
    border: 1px solid var(--line) !important;
    border-radius: var(--radius) !important;
    padding: 5px !important;
}

.stTabs [data-baseweb="tab"] {
    border-radius: 6px !important;
    color: var(--text-dim) !important;
    font-weight: 700 !important;
}

.stTabs [aria-selected="true"] {
    background: rgba(247, 244, 236, 0.12) !important;
    color: var(--text-bright) !important;
    border-bottom: none !important;
}

.stExpander {
    background: var(--surface) !important;
    border: 1px solid var(--line) !important;
    border-radius: var(--radius) !important;
    overflow: hidden !important;
}

.stExpander details summary,
.stExpander details summary * {
    background: transparent !important;
    color: var(--text-bright) !important;
    -webkit-text-fill-color: var(--text-bright) !important;
    font-size: 15px !important;
}

.stExpander details summary svg {
    color: var(--accent) !important;
    fill: var(--accent) !important;
}

button[data-testid="stExpandSidebarButton"] [data-testid="stIconMaterial"],
.stExpander details summary [data-testid="stIconMaterial"] {
    color: transparent !important;
    -webkit-text-fill-color: transparent !important;
    font-size: 0 !important;
    line-height: 1 !important;
}

button[data-testid="stExpandSidebarButton"] [data-testid="stIconMaterial"]::before {
    content: "";
    display: block;
    width: 18px;
    height: 12px;
    border-top: 2px solid var(--text-bright);
    border-bottom: 2px solid var(--text-bright);
    background: linear-gradient(var(--text-bright), var(--text-bright)) center / 18px 2px no-repeat;
}

.stExpander details summary [data-testid="stIconMaterial"]::before {
    content: "";
    display: inline-block;
    width: 8px;
    height: 8px;
    border-right: 2px solid var(--accent);
    border-bottom: 2px solid var(--accent);
    transform: rotate(-45deg);
}

.stExpander details[open] summary [data-testid="stIconMaterial"]::before {
    transform: rotate(45deg);
}

.stProgress > div > div {
    background: rgba(255,255,255,0.08) !important;
    border: 1px solid rgba(255,255,255,0.06) !important;
    border-radius: 999px !important;
    height: 9px !important;
}

.stProgress > div > div > div {
    background: linear-gradient(90deg, var(--accent), var(--gold)) !important;
    border-radius: 999px !important;
    box-shadow: none !important;
}

/* App shell */
.workspace-hero {
    border: 1px solid var(--line);
    border-radius: var(--radius);
    padding: 30px;
    margin: 8px 0 18px;
    background:
        linear-gradient(135deg, rgba(247, 244, 236, 0.12), rgba(123, 211, 200, 0.04) 45%, rgba(247, 244, 236, 0.035)),
        linear-gradient(180deg, rgba(13, 15, 20, 0.88), rgba(10, 12, 17, 0.92));
    box-shadow: var(--shadow);
    backdrop-filter: blur(20px) saturate(1.18);
    -webkit-backdrop-filter: blur(20px) saturate(1.18);
    position: relative;
    overflow: hidden;
}

.workspace-hero::before {
    content: "";
    position: absolute;
    inset: 0;
    pointer-events: none;
    border-top: 1px solid rgba(255, 255, 255, 0.16);
    background:
        linear-gradient(120deg, rgba(255,255,255,0.11), transparent 28%),
        radial-gradient(circle at 92% 14%, rgba(215,184,106,0.18), transparent 28%);
}

.workspace-hero-top {
    display: flex;
    align-items: flex-start;
    justify-content: space-between;
    gap: 20px;
    position: relative;
    z-index: 1;
}

.workspace-kicker {
    color: var(--accent) !important;
    font-size: 12px !important;
    font-weight: 800 !important;
    letter-spacing: 0.12em !important;
    text-transform: uppercase;
    margin-bottom: 8px;
}

.workspace-title {
    color: var(--text-bright) !important;
    font-size: clamp(29px, 4vw, 46px) !important;
    font-weight: 820 !important;
    line-height: 1.12 !important;
    margin: 0 !important;
}

.workspace-desc {
    color: var(--text-dim) !important;
    font-size: 15px !important;
    margin: 12px 0 0 !important;
    max-width: 720px;
}

.workspace-badge {
    flex: 0 0 auto;
    border: 1px solid rgba(123, 211, 200, 0.34);
    color: var(--accent) !important;
    background: rgba(123, 211, 200, 0.09);
    border-radius: 999px;
    padding: 8px 12px;
    font-size: 12px !important;
    font-weight: 800;
}

.writey-brandbar {
    display: flex;
    align-items: center;
    justify-content: space-between;
    gap: 18px;
    padding: 13px 16px;
    margin: 0 0 18px;
    border: 1px solid var(--line);
    border-radius: 999px;
    background: rgba(255, 255, 255, 0.045);
    box-shadow: inset 0 1px 0 rgba(255,255,255,0.06);
    backdrop-filter: blur(14px);
    -webkit-backdrop-filter: blur(14px);
}

.writey-brand-left {
    display: flex;
    align-items: baseline;
    gap: 12px;
    flex-wrap: wrap;
}

.writey-wordmark {
    color: var(--text-bright) !important;
    font-size: 18px !important;
    font-weight: 850 !important;
    letter-spacing: 0.04em !important;
}

.writey-cashtag,
.writey-author {
    color: var(--text-dim) !important;
    font-size: 12px !important;
    font-weight: 760 !important;
}

.writey-author {
    border: 1px solid rgba(215,184,106,0.26);
    border-radius: 999px;
    padding: 5px 10px;
    background: rgba(215,184,106,0.08);
    color: var(--gold-light) !important;
}

.interview-hero {
    border: 1px solid var(--line);
    border-radius: var(--radius);
    padding: 50px 28px;
    margin-bottom: 26px;
    text-align: center;
    background:
        linear-gradient(135deg, rgba(240, 213, 139, 0.14), rgba(123, 211, 200, 0.06)),
        rgba(255,255,255,0.055);
    box-shadow: var(--shadow);
    backdrop-filter: blur(18px);
    -webkit-backdrop-filter: blur(18px);
    position: relative;
    overflow: hidden;
}

.interview-hero::before {
    content: "";
    position: absolute;
    inset: 0;
    pointer-events: none;
    background:
        radial-gradient(circle at 50% -20%, rgba(240, 213, 139, 0.18), transparent 34%),
        linear-gradient(90deg, transparent, rgba(255,255,255,0.12), transparent);
}

.interview-hero > * {
    position: relative;
    z-index: 1;
}

.interview-hero .eyebrow {
    color: var(--accent) !important;
    font-size: 12px !important;
    font-weight: 850 !important;
    letter-spacing: 0.12em !important;
    text-transform: uppercase;
    margin-bottom: 12px;
}

.interview-hero .title {
    color: var(--text-bright) !important;
    font-size: clamp(40px, 7vw, 64px) !important;
    font-weight: 880 !important;
    line-height: 1 !important;
    margin: 0 !important;
}

.interview-hero .divider {
    width: 38px;
    height: 2px;
    background: var(--gold);
    margin: 20px auto;
}

.interview-hero .tagline {
    color: var(--text-dim) !important;
    font-size: 15px !important;
    margin: 0 !important;
}

.workspace-metrics {
    display: grid;
    grid-template-columns: repeat(4, minmax(0, 1fr));
    gap: 10px;
    margin-top: 22px;
    position: relative;
    z-index: 1;
}

.workspace-metric {
    border: 1px solid rgba(255,255,255,0.08);
    border-radius: var(--radius);
    background: rgba(255,255,255,0.06);
    padding: 14px;
    box-shadow: inset 0 1px 0 rgba(255,255,255,0.06);
}

.workspace-metric-label {
    color: var(--text-dim) !important;
    font-size: 12px !important;
    font-weight: 700 !important;
}

.workspace-metric-value {
    color: var(--text-bright) !important;
    font-size: 22px !important;
    font-weight: 820 !important;
    margin-top: 2px;
}

.premium-nav-container {
    border: 1px solid var(--line) !important;
    border-radius: 999px !important;
    background: rgba(255,255,255,0.055) !important;
    padding: 7px !important;
    margin: 16px 0 26px !important;
    backdrop-filter: blur(16px);
    -webkit-backdrop-filter: blur(16px);
}

.nav-item {
    min-height: 48px;
    display: flex;
    align-items: center;
    justify-content: center;
    border-radius: 999px;
    color: var(--text-dim) !important;
    font-size: 14px !important;
    font-weight: 760 !important;
}

.nav-item.active {
    background: linear-gradient(135deg, rgba(240, 213, 139, 0.20), rgba(123, 211, 200, 0.12)) !important;
    color: var(--text-bright) !important;
    border: 1px solid rgba(240, 213, 139, 0.32) !important;
    box-shadow: inset 0 1px 0 rgba(255,255,255,0.12), 0 8px 24px rgba(0,0,0,0.22) !important;
}

.section-divider {
    height: 1px;
    background: linear-gradient(90deg, transparent, var(--line), transparent);
    margin: 18px 0 28px;
}

.section-title-box {
    border: 1px solid var(--line) !important;
    border-radius: var(--radius) !important;
    padding: 24px 26px !important;
    margin-bottom: 24px !important;
    background:
        linear-gradient(90deg, rgba(123, 211, 200, 0.08), transparent 42%),
        rgba(255,255,255,0.045) !important;
    box-shadow: none !important;
}

.section-step {
    display: inline-flex !important;
    color: var(--accent) !important;
    font-size: 12px !important;
    font-weight: 820 !important;
    letter-spacing: 0.12em !important;
    margin-bottom: 10px !important;
}

.section-title-box h2 {
    color: var(--text-bright) !important;
    font-size: clamp(24px, 3vw, 34px) !important;
    font-weight: 820 !important;
    margin: 0 0 8px !important;
}

.section-title-box p {
    color: var(--text-dim) !important;
    font-size: 15px !important;
    margin: 0 !important;
}

/* Cards */
.score-card,
.result-card,
.info-card,
.data-card,
.summary-hub,
.title-card,
.stat-box,
.empty-state,
.quick-start-card,
.onboarding-card,
.setup-step {
    border-radius: var(--radius) !important;
    border: 1px solid var(--line) !important;
    background: var(--surface) !important;
    box-shadow: 0 18px 44px rgba(0,0,0,0.22), inset 0 1px 0 rgba(255,255,255,0.06) !important;
    backdrop-filter: blur(14px);
    -webkit-backdrop-filter: blur(14px);
}

.score-card {
    padding: 34px 28px !important;
    text-align: center;
    background:
        linear-gradient(145deg, rgba(240,213,139,0.10), rgba(123,211,200,0.04)),
        rgba(12,14,19,0.86) !important;
}

.score-card::before,
.score-card::after,
.login-card::before,
.login-card::after {
    display: none !important;
}

.score-number {
    color: var(--gold-light) !important;
    -webkit-text-fill-color: var(--gold-light) !important;
    background: none !important;
    font-size: clamp(72px, 10vw, 128px) !important;
    font-weight: 820 !important;
    letter-spacing: 0 !important;
    filter: none !important;
}

.result-card,
.data-card,
.summary-hub {
    padding: 20px 22px !important;
}

.data-card {
    border-left: 3px solid var(--gold) !important;
}

.info-card {
    border-left: 3px solid var(--accent) !important;
    padding: 18px 20px !important;
}

.stat-box {
    padding: 22px 18px !important;
    text-align: left !important;
}

.stat-value {
    color: var(--text-bright) !important;
    font-size: 30px !important;
    font-weight: 820 !important;
    letter-spacing: 0 !important;
}

.stat-label {
    color: var(--text-dim) !important;
    font-size: 12px !important;
    font-weight: 700 !important;
    letter-spacing: 0 !important;
    text-transform: none !important;
}

.title-card {
    padding: 18px 20px !important;
}

.title-card:hover,
.data-card:hover,
.info-card:hover,
.stat-box:hover,
.summary-hub:hover {
    transform: translateY(-1px) !important;
    border-color: var(--line-strong) !important;
    background: var(--surface-2) !important;
}

.title-main {
    color: var(--text-bright) !important;
    font-size: 17px !important;
    font-weight: 780 !important;
}

.title-sub {
    color: var(--text-dim) !important;
    font-size: 13px !important;
}

.verdict-go,
.verdict-wait,
.verdict-no {
    display: inline-flex;
    align-items: center;
    justify-content: center;
    border-radius: 999px !important;
    padding: 7px 12px !important;
    font-size: 12px !important;
    font-weight: 820 !important;
    letter-spacing: 0 !important;
    text-transform: none !important;
}

.verdict-go {
    color: var(--success) !important;
    border-color: rgba(123, 201, 158, 0.5) !important;
    background: rgba(123, 201, 158, 0.10) !important;
}

.verdict-wait {
    color: var(--warning) !important;
    border-color: rgba(240, 198, 107, 0.5) !important;
    background: rgba(240, 198, 107, 0.10) !important;
}

.verdict-no {
    color: var(--danger) !important;
    border-color: rgba(238, 127, 115, 0.5) !important;
    background: rgba(238, 127, 115, 0.10) !important;
}

.empty-state {
    padding: 58px 24px !important;
    text-align: center !important;
    border-style: dashed !important;
}

.empty-state-text {
    color: var(--text-dim) !important;
    font-size: 15px !important;
}

.next-section {
    height: 14px;
}

/* Login and onboarding */
.login-card {
    max-width: 440px !important;
    margin: 88px auto 28px !important;
    padding: 52px 44px !important;
    border-radius: var(--radius) !important;
    border: 1px solid var(--line) !important;
    background:
        linear-gradient(135deg, rgba(240, 213, 139, 0.10), rgba(123, 211, 200, 0.06)),
        var(--surface) !important;
    box-shadow: var(--shadow) !important;
}

.login-title {
    color: var(--text-bright) !important;
    -webkit-text-fill-color: var(--text-bright) !important;
    background: none !important;
    font-size: 38px !important;
    font-weight: 860 !important;
    letter-spacing: 0.08em !important;
}

.login-subtitle {
    color: var(--text-dim) !important;
    font-size: 12px !important;
    font-weight: 780 !important;
    letter-spacing: 0.12em !important;
}

.onboarding-card {
    padding: 34px !important;
    margin: 16px 0 22px !important;
}

.onboarding-eyebrow {
    color: var(--accent) !important;
    font-size: 12px !important;
    font-weight: 820 !important;
    letter-spacing: 0.12em !important;
    text-transform: uppercase;
}

.onboarding-title {
    color: var(--text-bright) !important;
    font-size: clamp(28px, 4vw, 42px) !important;
    font-weight: 850 !important;
    line-height: 1.16 !important;
    margin: 8px 0 10px !important;
}

.onboarding-copy {
    color: var(--text-dim) !important;
    font-size: 15px !important;
    max-width: 760px;
}

.setup-grid {
    display: grid;
    grid-template-columns: repeat(3, minmax(0, 1fr));
    gap: 12px;
    margin: 18px 0 20px;
}

.setup-step {
    padding: 20px !important;
}

.setup-number {
    display: inline-flex;
    align-items: center;
    justify-content: center;
    width: 30px;
    height: 30px;
    border-radius: 999px;
    background: rgba(123, 211, 200, 0.12);
    color: var(--accent) !important;
    font-weight: 850;
    margin-bottom: 12px;
}

.setup-title {
    color: var(--text-bright) !important;
    font-size: 17px !important;
    font-weight: 800 !important;
    margin-bottom: 8px;
}

.setup-copy {
    color: var(--text-dim) !important;
    font-size: 13px !important;
}

.api-callout {
    border: 1px solid rgba(240, 213, 139, 0.34);
    border-radius: var(--radius);
    background: rgba(240, 213, 139, 0.10);
    padding: 18px 20px;
    margin: 16px 0;
}

.api-callout b {
    color: var(--gold-light) !important;
}

.quick-start-card {
    padding: 18px 20px !important;
    margin: 18px 0 !important;
    border-color: rgba(240, 213, 139, 0.30) !important;
    background:
        linear-gradient(135deg, rgba(240, 213, 139, 0.14), rgba(123, 211, 200, 0.08)),
        var(--surface) !important;
}

.quick-start-card .eyebrow {
    color: var(--gold-light) !important;
    font-size: 12px !important;
    font-weight: 820 !important;
    margin-bottom: 4px;
}

.quick-start-card .headline {
    color: var(--text-bright) !important;
    font-size: 20px !important;
    font-weight: 850 !important;
}

.app-footer {
    text-align: center;
    padding: 22px 20px;
    margin-top: 44px;
    border-top: 1px solid var(--line);
    color: var(--text-dim) !important;
    font-size: 13px !important;
}

.app-footer b {
    color: var(--gold-light) !important;
}

.sidebar-footer {
    text-align: center;
    padding: 14px 10px;
    border: 1px solid var(--line);
    border-radius: var(--radius);
    background: rgba(255,255,255,0.04);
    color: var(--text-dim) !important;
    font-size: 12px !important;
}

.sidebar-footer b {
    color: var(--gold-light) !important;
}

/* Cover/content preview */
.book-wrapper {
    border-radius: var(--radius) !important;
    background: rgba(255,255,255,0.045) !important;
}

.content-preview-box {
    border-radius: var(--radius) !important;
}

::-webkit-scrollbar {
    width: 8px;
    height: 8px;
}

::-webkit-scrollbar-track {
    background: #0c0e12;
}

::-webkit-scrollbar-thumb {
    background: rgba(247,244,236,0.22);
    border-radius: 999px;
}

@media (max-width: 900px) {
    .main .block-container {
        padding: 1.4rem 1rem 3rem !important;
    }

    .workspace-hero,
    .onboarding-card {
        padding: 22px !important;
    }

    .workspace-hero-top {
        flex-direction: column;
    }

    .workspace-metrics,
    .setup-grid {
        grid-template-columns: 1fr 1fr;
    }
}

@media (max-width: 640px) {
    .workspace-metrics,
    .setup-grid {
        grid-template-columns: 1fr;
    }

    .section-title-box {
        padding: 20px !important;
    }
}

@media (prefers-reduced-motion: reduce) {
    *, *::before, *::after {
        animation: none !important;
        transition: none !important;
    }
}
</style>
""", unsafe_allow_html=True)



# ==========================================
# 비밀번호 인증 (단순)
# ==========================================
if 'authenticated' not in st.session_state:
    st.session_state['authenticated'] = False

# 저장된 비밀번호로 자동 로그인 (매 실행마다 재확인)
# 쿠키는 첫 실행 뒤 재실행 시점에 늦게 로드될 수 있으므로, 미인증 상태에서는
# 매번 saved_password를 다시 확인해야 쿠키가 늦게 들어와도 자동 로그인이 된다.
if not st.session_state['authenticated']:
    if st.session_state.get('saved_password', '') == CORRECT_PASSWORD:
        st.session_state['authenticated'] = True

if not st.session_state['authenticated']:
    st.markdown("""
    <div class="login-card">
        <div class="login-subtitle">CASHMAKER</div>
        <div class="login-title">WRITEY</div>
        <div class="login-subtitle">Premium E-Book Studio</div>
    </div>
    """, unsafe_allow_html=True)

    col1, col2, col3 = st.columns([1, 1, 1])
    with col2:
        pw = st.text_input("비밀번호", type="password", key="pw_login", placeholder="Enter password...")
        remember_pw = st.checkbox("비밀번호 저장 (다음 접속 시 자동 로그인)", value=True, key="remember_pw")
        if st.button("입장", key="btn_login", use_container_width=True):
            if pw == CORRECT_PASSWORD:
                st.session_state['authenticated'] = True
                if remember_pw:
                    save_password_to_browser(pw)
                st.rerun()
            else:
                st.error("비밀번호가 올바르지 않습니다")
    st.stop()


# 세션 초기화
defaults = {
    'topic': '', 'target_persona': '', 'pain_points': '',
    'outline': [], 'chapters': {}, 'book_title': '', 'subtitle': '',
    'score_details': None, 'generated_titles': None, 'suggested_targets': None,
    'analyzed_pains': None, 'review_analysis': None, 'market_gaps': None,
    'knowledge_hub': [], 'study_summary': None, 'current_page': 0,
    'recommended_refs': None, 'generated_ideas': None,
    # 인터뷰 관련 변수
    'interview_completed': False,
    'interview_data': {},
    'author_name': '',
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

# 사이드바
with st.sidebar:
    # API 키 섹션 (접기/펼치기 가능)
    if 'show_api_section' not in st.session_state:
        st.session_state['show_api_section'] = True

    if 'api_key' not in st.session_state:
        # 쿠키에서 API 키 불러오기
        st.session_state['api_key'] = st.session_state.get('saved_api_key', '') or ''
    # api_key가 비어 있는데 쿠키에 저장된 키가 (늦게) 들어왔다면 복원
    elif not st.session_state['api_key'] and st.session_state.get('saved_api_key'):
        st.session_state['api_key'] = st.session_state['saved_api_key']

    # API 키가 입력되어 있으면 기본적으로 접힌 상태로
    api_key_exists = bool(st.session_state['api_key'])

    col_title, col_toggle = st.columns([4, 1])
    with col_title:
        st.markdown("### 🔑 Claude API 키")
    with col_toggle:
        toggle_label = "▼" if st.session_state['show_api_section'] else "▶"
        if st.button(toggle_label, key="toggle_api_section", help="접기/펼치기"):
            st.session_state['show_api_section'] = not st.session_state['show_api_section']
            st.rerun()

    if st.session_state['show_api_section']:
        api_key = st.text_input("키 입력", value=st.session_state['api_key'], type="password", key="api_sidebar", label_visibility="collapsed", placeholder="sk-ant-api03-... 형식")
        if api_key != st.session_state['api_key']:
            st.session_state['api_key'] = api_key
            # 쿠키에 저장
            if api_key:
                save_api_key_to_browser(api_key)
                # 비밀번호 흐름과 동일하게 즉시 재실행해 pending 값을 쿠키에 바로 기록
                # (재실행이 없으면 다음 상호작용 전까지 쿠키 기록이 미뤄져 저장이 누락될 수 있음)
                st.rerun()

        if api_key:
            st.success("✅ Claude 키 입력 완료!")
        else:
            st.error("⚠️ Claude API 키를 입력하세요")
    else:
        # 접힌 상태에서 간단한 상태 표시
        if st.session_state['api_key']:
            st.caption("✅ API 키 설정됨")
        else:
            st.caption("⚠️ API 키 필요")

    # 모델 선택
    st.markdown("### 🤖 모델 선택")
    if 'claude_model' not in st.session_state:
        st.session_state['claude_model'] = "claude-sonnet-4-5"

    model_options = {
        "Claude Sonnet 4.5 (추천)": "claude-sonnet-4-5",
        "Claude Opus 4.5 (최고 품질)": "claude-opus-4-5",
        "Claude Haiku 4.5 (저렴)": "claude-haiku-4-5"
    }
    selected_model = st.selectbox(
        "모델 선택",
        options=list(model_options.keys()),
        index=0,
        label_visibility="collapsed"
    )
    st.session_state['claude_model'] = model_options[selected_model]

    if "Haiku" in selected_model:
        st.info("💰 가장 저렴하고 빠름.\n📌 단, 목차·본문·프롤로그·에필로그·컨셉·제목 생성은 품질 보장을 위해 자동으로 Sonnet 4.5 사용")
    elif "Opus" in selected_model:
        st.info("💎 최고 품질, 단가가 가장 높음")
    else:
        st.info("⚡ 균형잡힌 품질/가격, 일반적으로 가장 추천")

    # API 키 발급 방법 안내
    with st.expander("📖 Claude API 키 발급 방법 (상세)", expanded=False):
        st.markdown("""
        ### 🟣 1단계: Anthropic 회원가입

        1. 아래 버튼을 클릭하세요
        2. **"Sign up"** 클릭
        3. Google 계정 또는 이메일로 가입
        """)
        st.link_button("🔗 Anthropic 가입 페이지", "https://console.anthropic.com/", use_container_width=True)

        st.markdown("""
        ---
        ### 💳 2단계: 결제 수단 등록

        1. 로그인 후 왼쪽 메뉴에서 **"Settings"** 클릭
        2. **"Billing"** 클릭
        3. **"Add payment method"** 클릭
        4. 카드 정보 입력 후 저장
        5. **"Add credits"**로 크레딧 충전 ($5~10 추천)
        """)
        st.link_button("🔗 Billing 페이지 바로가기", "https://console.anthropic.com/settings/billing", use_container_width=True)

        st.markdown("""
        ---
        ### 🔑 3단계: API 키 발급

        1. 왼쪽 메뉴에서 **"API Keys"** 클릭
        2. **"Create Key"** 버튼 클릭
        3. 이름 입력 (예: ebook)
        4. **"Create Key"** 클릭
        5. 생성된 키 **복사** (sk-ant-api03-... 형식)
        6. 위 입력창에 **붙여넣기**
        """)
        st.link_button("🔗 API Keys 페이지 바로가기", "https://console.anthropic.com/settings/keys", use_container_width=True)

        st.markdown("---")
        st.warning("⚠️ API 키는 한 번만 보여줍니다. 복사해두세요!")
        st.success("💰 예상 비용: 전자책 1권당 약 200~500원")

    st.markdown("---")
    st.markdown("### 📊 진행 상황")
    progress = sum([bool(st.session_state['topic']), bool(st.session_state['target_persona']), bool(st.session_state['outline']), len(st.session_state['chapters']) > 0]) / 4
    st.progress(progress)

    st.markdown("---")
    st.markdown("### 🚀 빠른 이동")
    sidebar_pages = ["① 주제", "② 목차", "③ 본문", "④ 완성"]
    sidebar_mapping = [0, 4, 5, 7]
    for i, p in enumerate(sidebar_pages):
        if st.button(p, key=f"sidebar_nav_{i}", use_container_width=True):
            st.session_state['current_page'] = sidebar_mapping[i]
            st.rerun()

    # 사이드바 하단 제작자 정보
    st.markdown("---")
    st.markdown("""
    <div class="sidebar-footer">
        <b>CASHMAKER</b><br>
        제작: 남현우 작가
    </div>
    """, unsafe_allow_html=True)

# ==========================================
# 헬퍼 함수
# ==========================================
def get_api_key():
    return st.session_state.get('api_key', '')

def clean_text(text):
    if not text:
        return ""
    text = re.sub(r'^#{1,6}\s*', '', text, flags=re.MULTILINE)
    text = re.sub(r'\*\*([^*]+)\*\*', r'「\1」', text)
    text = text.replace('**', '').replace('*', '').replace('###', '').replace('##', '').replace('#', '')
    return text.strip()

def clean_content(text, subtopic=None):
    if not text:
        return ""
    # HTML 테이블 및 모든 HTML 태그 제거
    text = re.sub(r'<table[^>]*>.*?</table>', '', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<[^>]+>', '', text)
    # 마크다운 제거
    text = re.sub(r'^#{1,6}\s*', '', text, flags=re.MULTILINE)
    # **굵은글씨** 패턴 완전 제거
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    text = text.replace('**', '').replace('*', '').replace('###', '').replace('##', '').replace('#', '')
    # 연속 줄바꿈 정리
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = text.strip()

    # 본문 첫 줄이 소제목과 동일하면 제거 (AI가 소제목을 본문 시작에 또 박는 경우)
    if subtopic:
        sub_clean = re.sub(r'[「」"\'\s\.\?!]+', '', subtopic).strip()
        lines = text.split('\n')
        if lines:
            first_line_clean = re.sub(r'[「」"\'\s\.\?!]+', '', lines[0]).strip()
            # 정확 일치 또는 거의 같은 경우 (90% 이상 매칭)
            if first_line_clean == sub_clean or (sub_clean and sub_clean in first_line_clean and len(first_line_clean) <= len(sub_clean) * 1.2):
                text = '\n'.join(lines[1:]).lstrip('\n').strip()
    return text

def parse_json(response):
    """JSON 파싱 - 개선된 에러 처리"""
    if not response:
        return None
    try:
        # 먼저 전체 응답에서 JSON 블록 찾기
        json_match = re.search(r'```json\s*([\s\S]*?)\s*```', response)
        if json_match:
            return json.loads(json_match.group(1))

        # JSON 블록이 없으면 중괄호로 시작하는 객체 찾기
        match = re.search(r'\{[\s\S]*\}', response)
        if match:
            json_str = match.group()
            # 불완전한 JSON 수정 시도
            json_str = re.sub(r',\s*}', '}', json_str)  # 마지막 쉼표 제거
            json_str = re.sub(r',\s*]', ']', json_str)  # 배열 마지막 쉼표 제거
            return json.loads(json_str)
    except json.JSONDecodeError as e:
        st.warning(f"JSON 파싱 경고: {str(e)[:50]}")
    except Exception as e:
        st.warning(f"파싱 오류: {str(e)[:50]}")
    return None

def ask_ai(prompt, temp=0.7, ensure_quality=False):
    """Claude API 호출

    ensure_quality=True 시 Haiku 선택해도 자동으로 Sonnet 4.5로 업그레이드.
    목차/본문/프롤로그/에필로그/컨셉 생성처럼 정교한 프롬프트를 따라야 하는 작업에 사용.
    """
    api_key = get_api_key()
    if not api_key:
        st.error("Claude API 키를 입력해주세요")
        return None

    if not CLAUDE_AVAILABLE:
        st.error("anthropic 패키지가 설치되지 않았습니다. pip install anthropic")
        return None

    # 선택된 모델 가져오기 (기본값: Sonnet 4.5)
    user_model = st.session_state.get('claude_model', 'claude-sonnet-4-5')

    # 핵심 생성 작업은 Haiku 자동 업그레이드 (품질 일관성 보장)
    if ensure_quality and 'haiku' in user_model.lower():
        model = 'claude-sonnet-4-5'
    else:
        model = user_model

    try:
        client = anthropic.Anthropic(api_key=api_key)
        message = client.messages.create(
            model=model,
            max_tokens=8000,
            temperature=temp,
            messages=[
                {"role": "user", "content": prompt}
            ]
        )
        return message.content[0].text
    except anthropic.AuthenticationError:
        st.error("API 키가 유효하지 않습니다. Claude API 키를 확인해주세요.")
        return None
    except anthropic.RateLimitError:
        st.error("API 할당량이 초과되었습니다. 잠시 후 다시 시도해주세요.")
        return None
    except anthropic.BadRequestError as e:
        try:
            err_msg = str(e).encode('utf-8', errors='ignore').decode('utf-8')[:100]
        except:
            err_msg = "요청 형식 오류"
        st.error(f"요청 오류: {err_msg}")
        return None
    except Exception as e:
        try:
            err_msg = str(e).encode('utf-8', errors='ignore').decode('utf-8')[:100]
        except:
            err_msg = "알 수 없는 오류"
        st.error(f"AI 오류: {err_msg}")
        return None

def generate_cover_image_gemini(title, subtitle, theme_keywords):
    """Google Gemini로 표지 배경 이미지 생성"""

    api_key = get_api_key()
    if not api_key:
        return None, "Gemini API 키가 필요합니다."

    if not IMAGEN_AVAILABLE:
        return None, "google-genai 패키지가 필요합니다: pip install google-genai"

    try:
        client = google_genai.Client(api_key=api_key)

        # 베스트셀러급 고급 표지 프롬프트 - 텍스트 절대 금지 강조
        prompt = f"""Create an ABSTRACT background image for a book cover.

Theme keywords: {theme_keywords}

STYLE: Dark, moody, cinematic atmosphere. Abstract shapes, gradients, smoke, light rays, or geometric patterns. Luxury aesthetic with gold/amber accent lighting on deep black background.

CRITICAL RULES:
- ONLY abstract visuals: smoke, light, shadows, gradients, textures
- NO objects, NO people, NO faces, NO hands
- NO text, NO letters, NO words, NO numbers, NO symbols, NO characters of ANY language
- NO Korean, NO English, NO Chinese, NO Japanese characters
- Pure abstract art only

OUTPUT: Dark dramatic background with subtle golden light accents, suitable for text overlay."""

        # Gemini 이미지 생성
        response = client.models.generate_content(
            model='gemini-2.0-flash-exp-image-generation',
            contents=prompt,
            config=genai_types.GenerateContentConfig(
                response_modalities=['IMAGE', 'TEXT']
            )
        )

        if response.candidates and response.candidates[0].content.parts:
            for part in response.candidates[0].content.parts:
                if hasattr(part, 'inline_data') and part.inline_data:
                    image_base64 = base64.b64encode(part.inline_data.data).decode('utf-8')
                    return image_base64, None

        return None, "이미지가 생성되지 않았습니다."

    except Exception as e:
        error_msg = str(e)
        if "quota" in error_msg.lower() or "limit" in error_msg.lower():
            return None, "API 할당량 초과. 잠시 후 다시 시도해주세요."
        elif "safety" in error_msg.lower():
            return None, "안전 필터에 의해 차단되었습니다. 다른 키워드로 시도해주세요."
        return None, f"이미지 생성 오류: {error_msg[:80]}"

def generate_cover_prompt_ai(title, subtitle, topic):
    """AI가 표지 디자인 컨셉과 이미지 프롬프트 생성"""
    prompt = f"""당신은 베스트셀러 책 표지 디자이너입니다.

책 제목: {title}
부제: {subtitle}
주제: {topic}

이 책의 표지 이미지를 위한 영문 프롬프트를 만들어주세요.

[요구사항]
1. 실제 베스트셀러 표지 스타일 분석 기반
2. 제목의 핵심 메시지를 시각적으로 표현
3. 고급스럽고 전문적인 느낌
4. 텍스트 오버레이를 위한 여백 고려
5. 추상적이거나 상징적인 이미지

[출력 형식]
IMAGE_PROMPT: (영문 이미지 생성 프롬프트, 50단어 이내)
COLOR_SCHEME: (추천 컬러 팔레트, 예: dark, gold, minimal)
STYLE: (디자인 스타일, 예: editorial, bold, elegant)

영문 프롬프트만 출력하세요. 한국어 설명 불필요."""

    result = ask_ai(prompt, temp=0.7)
    if result:
        # 파싱
        image_prompt = ""
        color_scheme = "dark"
        style = "editorial"

        for line in result.split('\n'):
            if 'IMAGE_PROMPT:' in line:
                image_prompt = line.split('IMAGE_PROMPT:')[-1].strip()
            elif 'COLOR_SCHEME:' in line:
                color_scheme = line.split('COLOR_SCHEME:')[-1].strip().lower()
            elif 'STYLE:' in line:
                style = line.split('STYLE:')[-1].strip().lower()

        return image_prompt, color_scheme, style
    return None, "dark", "editorial"


# ==========================================
# 고급 표지 렌더러 (외부 API 불필요, 벡터 SVG)
# ==========================================
COVER_TEMPLATES = {
    "bestseller_purple": "베스트셀러 브리핑 — 블랙 타이포 + 퍼플 띠지",
    "literary_watercolor": "문학 수채화 — 파스텔 풍경 + 수상작 무드",
    "signature_noir": "시그니처 누아르 — 블랙 재킷 + 금박 타이포",
    "modern_editorial": "모던 에디토리얼 — 아이보리 지면 + 강한 제목",
    "executive_teal": "비즈니스 딥틸 — 청록 포인트 + 페이퍼백",
    "archive_red": "아카이브 레드 — 서점형 강렬한 띠지",
    "paper_luxe": "페이퍼 럭스 — 미색 양장본 + 클래식 장식",
}


def _text_visual_len(text):
    """한글/영문 혼합 제목의 대략적인 시각 폭."""
    total = 0
    for ch in str(text or ""):
        if ch.isspace():
            total += 0.35
        elif ord(ch) < 128:
            total += 0.58
        else:
            total += 1
    return total


def _wrap_title_lines(title, max_chars=7, max_lines=3):
    """제목을 표지용으로 줄바꿈 (한글 글자수 기준 + 긴 단어 보정)."""
    title = re.sub(r'\s+', ' ', (title or "").strip())
    if not title:
        return ["제목"]
    words = title.split()
    lines, cur = [], ""
    for w in words:
        cand = (cur + " " + w).strip()
        if not cur or _text_visual_len(cand) <= max_chars:
            cur = cand
        else:
            lines.append(cur)
            cur = w
    if cur:
        lines.append(cur)
    # 띄어쓰기 없는 긴 단어는 강제 줄바꿈
    if len(lines) == 1 and _text_visual_len(lines[0]) > max_chars:
        s = lines[0]
        lines = [s[i:i + max_chars] for i in range(0, len(s), max_chars)]
    return lines[:max_lines]


def _wrap_svg_lines(text, max_chars=17, max_lines=2):
    text = re.sub(r'\s+', ' ', (text or "").strip())
    if not text:
        return []
    words = text.split()
    lines, cur = [], ""
    for w in words:
        cand = (cur + " " + w).strip()
        if not cur or _text_visual_len(cand) <= max_chars:
            cur = cand
        else:
            lines.append(cur)
            cur = w
        if len(lines) >= max_lines:
            break
    if cur and len(lines) < max_lines:
        lines.append(cur)
    return lines[:max_lines]


def _bestseller_title_lines(title):
    """강한 서점형 표지를 위해 짧은 제목은 의도적으로 큼직하게 분할."""
    clean = re.sub(r'\s+', ' ', (title or "").strip())
    if not clean:
        return ["제목"]
    words = clean.split()
    if len(words) == 2 and _text_visual_len(clean) <= 8.4:
        return words
    if len(words) == 3 and _text_visual_len(clean) <= 9.2:
        return [words[0], " ".join(words[1:])]
    return _wrap_title_lines(clean, max_chars=6, max_lines=4)


def _svg_tspans(lines, x, line_h):
    return "".join(
        f'<tspan x="{x}" dy="{0 if i == 0 else line_h}">{html.escape(str(line))}</tspan>'
        for i, line in enumerate(lines)
    )


def _candidate_cover_asset_dirs():
    asset_dirs = ("", "assets", "static", "media", "images", "covers")
    bases = []
    try:
        app_dir = Path(__file__).resolve().parent
        bases.extend(app_dir / name if name else app_dir for name in asset_dirs)
    except Exception:
        pass
    try:
        cwd = Path.cwd()
        bases.extend(cwd / name if name else cwd for name in asset_dirs)
    except Exception:
        pass
    bases.extend([Path.home() / "Downloads", Path.home() / "Desktop"])

    unique = []
    seen = set()
    for base in bases:
        key = str(base)
        if key not in seen:
            seen.add(key)
            unique.append(base)
    return unique


def _find_cover_photo():
    """베스트셀러형 하단 인물 사진 자동 탐색. 없으면 고급 실루엣으로 대체."""
    names = (
        "cover_photo.png", "cover_photo.jpg", "cover_photo.jpeg", "cover_photo.webp",
        "author_photo.png", "author_photo.jpg", "author_photo.jpeg", "author_photo.webp",
        "profile_photo.png", "profile_photo.jpg", "profile_photo.jpeg", "profile_photo.webp",
    )
    for base in _candidate_cover_asset_dirs():
        for name in names:
            p = base / name
            if p.exists():
                return p
    return None


def _image_data_uri(path):
    try:
        suffix = path.suffix.lower()
        mime = {
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".png": "image/png",
            ".webp": "image/webp",
        }.get(suffix)
        if not mime:
            return None
        with open(path, "rb") as f:
            return f"data:{mime};base64,{base64.b64encode(f.read()).decode()}"
    except Exception:
        return None


def _cover_photo_data_uri():
    p = _find_cover_photo()
    return _image_data_uri(p) if p else None


def build_cover_svg(template, title, subtitle, author):
    """종이책 베스트셀러 톤의 고급 표지 SVG 생성 (1600x2560, 무손실 벡터)."""
    template = template if template in COVER_TEMPLATES else "bestseller_purple"
    title = (title or "제목을 입력하세요").strip()
    subtitle = (subtitle or "").strip()
    author = (author or "저자").strip()

    title_lines = _wrap_title_lines(title, max_chars=7, max_lines=3)
    longest = max((_text_visual_len(line) for line in title_lines), default=4)
    if longest <= 4.2:
        title_size = 230
    elif longest <= 6.2:
        title_size = 190
    elif longest <= 8.2:
        title_size = 155
    else:
        title_size = 132
    if len(title_lines) >= 3:
        title_size = min(title_size, 155)
    title_line_h = int(title_size * 1.12)
    title_y = 1115 - int((len(title_lines) - 1) * title_line_h / 2)
    subtitle_lines = _wrap_svg_lines(subtitle, max_chars=19, max_lines=3)
    author_esc = html.escape(author)

    fonts = (
        '<style><![CDATA['
        "@import url('https://fonts.googleapis.com/css2?family=Noto+Serif+KR:wght@500;600;700;900&display=swap');"
        "@import url('https://cdn.jsdelivr.net/gh/orioncactus/pretendard/dist/web/static/pretendard.css');"
        ']]></style>'
    )
    serif = "'Noto Serif KR','Nanum Myeongjo',serif"
    sans = "'Pretendard','Apple SD Gothic Neo',sans-serif"
    base_defs = """
    <defs>
        <filter id="paperNoise" x="-20%" y="-20%" width="140%" height="140%">
            <feTurbulence type="fractalNoise" baseFrequency="0.85" numOctaves="4" seed="12" result="noise"/>
            <feColorMatrix in="noise" type="saturate" values="0"/>
            <feComponentTransfer><feFuncA type="table" tableValues="0 0.06"/></feComponentTransfer>
            <feBlend mode="multiply" in2="SourceGraphic"/>
        </filter>
        <linearGradient id="foil" x1="0%" y1="0%" x2="100%" y2="100%">
            <stop offset="0%" stop-color="#fff0b7"/>
            <stop offset="42%" stop-color="#c49436"/>
            <stop offset="67%" stop-color="#f3d68a"/>
            <stop offset="100%" stop-color="#926825"/>
        </linearGradient>
        <linearGradient id="softPaper" x1="0%" y1="0%" x2="100%" y2="100%">
            <stop offset="0%" stop-color="#fbf6e9"/>
            <stop offset="100%" stop-color="#ebe1cd"/>
        </linearGradient>
    </defs>
    """

    def subtitle_svg(x=800, y=1570, fill="#d8c7a5", anchor="middle", size=46, line_h=62, width=19):
        lines = _wrap_svg_lines(subtitle, max_chars=width, max_lines=3)
        if not lines:
            return ""
        return (
            f'<text x="{x}" y="{y}" text-anchor="{anchor}" font-family="{serif}" '
            f'font-size="{size}" font-weight="500" fill="{fill}" letter-spacing="2">'
            f'{_svg_tspans(lines, x, line_h)}</text>'
        )

    def title_svg(x=800, y=None, fill="#f8f1df", anchor="middle", weight=900, letter=0, size=None):
        fs = size or title_size
        lh = int(fs * 1.12)
        yy = y if y is not None else 1115 - int((len(title_lines) - 1) * lh / 2)
        return (
            f'<text x="{x}" y="{yy}" text-anchor="{anchor}" font-family="{serif}" '
            f'font-size="{fs}" font-weight="{weight}" fill="{fill}" letter-spacing="{letter}" '
            f'paint-order="stroke" stroke="rgba(0,0,0,0.10)" stroke-width="1">'
            f'{_svg_tspans(title_lines, x, lh)}</text>'
        )

    if template == "bestseller_purple":
        photo_uri = _cover_photo_data_uri()
        bestseller_lines = _bestseller_title_lines(title)
        bestseller_longest = max((_text_visual_len(line) for line in bestseller_lines), default=4)
        if bestseller_longest <= 4.2:
            bestseller_size = 312
        elif bestseller_longest <= 6.2:
            bestseller_size = 250
        elif bestseller_longest <= 8.2:
            bestseller_size = 184
        else:
            bestseller_size = 154
        if len(bestseller_lines) >= 4:
            bestseller_size = min(bestseller_size, 166)
        bestseller_lh = int(bestseller_size * 1.05)
        bestseller_y = 785 - int((len(bestseller_lines) - 1) * bestseller_lh / 2)
        top_copy = subtitle_lines or ["품격 있는 전자책을 위한"]
        top_copy_svg = _svg_tspans(top_copy[:2], 244, 72)
        photo_block = (
            f'<image href="{photo_uri}" x="72" y="1445" width="610" height="1115" '
            'preserveAspectRatio="xMidYMin slice" clip-path="url(#authorPhotoCut)"/>'
            '<rect x="72" y="1445" width="610" height="1115" fill="url(#photoShade)" '
            'clip-path="url(#authorPhotoCut)" opacity="0.24"/>'
        ) if photo_uri else """
        <g transform="translate(102 1510)">
            <ellipse cx="270" cy="958" rx="292" ry="86" fill="#d7d7d2" opacity="0.64"/>
            <path d="M72 1018 C122 790 206 690 286 690 C366 690 456 790 510 1018 Z" fill="#050505"/>
            <path d="M126 728 C148 560 202 476 286 476 C372 476 430 558 452 728 C385 792 194 792 126 728 Z" fill="#101010"/>
            <ellipse cx="286" cy="360" rx="135" ry="154" fill="#171717"/>
            <path d="M154 332 C178 214 248 158 338 188 C418 214 455 280 435 386 C378 327 265 310 154 332 Z" fill="#020202"/>
            <path d="M165 420 C212 496 355 498 410 420 C398 548 358 614 286 614 C215 614 176 548 165 420 Z" fill="#242424"/>
            <path d="M86 1018 H486" stroke="#6a21b8" stroke-width="12"/>
        </g>
        """
        bg = """
        <defs>
            <clipPath id="authorPhotoCut"><rect x="72" y="1445" width="610" height="1115"/></clipPath>
            <linearGradient id="photoShade" x1="0%" y1="0%" x2="0%" y2="100%">
                <stop offset="0%" stop-color="#ffffff" stop-opacity="0"/>
                <stop offset="72%" stop-color="#000000" stop-opacity="0.12"/>
                <stop offset="100%" stop-color="#000000" stop-opacity="0.42"/>
            </linearGradient>
        </defs>
        <rect width="1600" height="2560" fill="#ffffff"/>
        <rect width="1600" height="1535" fill="#040404" filter="url(#paperNoise)"/>
        <rect x="0" y="1535" width="1600" height="1025" fill="#f7f7f4"/>
        <path d="M158 130 H1320 M158 130 V1518 H998 L918 1636 L1284 1518 H1450" fill="none" stroke="#6220a5" stroke-width="30" stroke-linejoin="miter"/>
        <path d="M158 130 H1320 M158 130 V1518 H998 L918 1636 L1284 1518 H1450" fill="none" stroke="#7b35c4" stroke-width="12" stroke-linejoin="miter"/>
        <rect x="0" y="1530" width="1600" height="10" fill="#ecebe7"/>
        <rect x="100" y="2384" width="470" height="10" fill="#6a21b8"/>
        <rect x="625" y="2384" width="330" height="10" fill="#6a21b8"/>
        <rect x="1020" y="2384" width="420" height="10" fill="#6a21b8"/>
        """
        kicker = (
            f'<text x="1412" y="165" text-anchor="end" font-family="{sans}" font-size="45" '
            f'font-weight="900" fill="#f6f6f2" letter-spacing="1">{author_esc} 지음</text>'
            f'<text x="244" y="305" font-family="{serif}" font-size="72" font-weight="800" '
            f'fill="#f6f6f2" letter-spacing="-1">{top_copy_svg}</text>'
            f'<text x="1295" y="575" text-anchor="middle" font-family="{sans}" font-size="43" '
            f'font-weight="900" fill="#a9a9a9" letter-spacing="3">INSIGHT</text>'
            f'<text x="1295" y="645" text-anchor="middle" font-family="{sans}" font-size="43" '
            f'font-weight="900" fill="#a9a9a9" letter-spacing="3">STRATEGY</text>'
            f'<text x="1295" y="715" text-anchor="middle" font-family="{sans}" font-size="43" '
            f'font-weight="900" fill="#a9a9a9" letter-spacing="3">MONEY</text>'
            f'<text x="1295" y="785" text-anchor="middle" font-family="{sans}" font-size="43" '
            f'font-weight="900" fill="#a9a9a9" letter-spacing="3">BRANDING</text>'
        )
        title_block = (
            f'<text x="238" y="{bestseller_y}" text-anchor="start" font-family="{serif}" '
            f'font-size="{bestseller_size}" font-weight="900" fill="#ffffff" '
            f'letter-spacing="-6" paint-order="stroke" stroke="#111111" stroke-width="2">'
            f'{_svg_tspans(bestseller_lines, 238, bestseller_lh)}</text>'
        )
        sub_block = (
            f'{photo_block}'
            f'<text x="1100" y="1772" text-anchor="middle" font-family="{sans}" '
            f'font-size="48" font-weight="900" fill="#161616">매일 읽히는 한 권의</text>'
            f'<text x="1100" y="1880" text-anchor="middle" font-family="{serif}" '
            f'font-size="78" font-weight="900" fill="#6a21b8" letter-spacing="-2">프리미엄 브리핑</text>'
            f'<text x="1100" y="1988" text-anchor="middle" font-family="{sans}" '
            f'font-size="47" font-weight="900" fill="#161616">독자의 판단을 움직이는</text>'
            f'<text x="1100" y="2094" text-anchor="middle" font-family="{sans}" '
            f'font-size="64" font-weight="950" fill="#6a21b8">진짜 지식</text>'
            f'<text x="292" y="2268" text-anchor="middle" font-family="{sans}" font-size="48" '
            f'font-weight="900" fill="#ffffff">★★★★★</text>'
            f'<text x="292" y="2326" text-anchor="middle" font-family="{sans}" font-size="38" '
            f'font-weight="900" fill="#ffffff">베스트셀러 감성</text>'
            f'<text x="1165" y="2268" text-anchor="middle" font-family="{sans}" font-size="48" '
            f'font-weight="900" fill="#6a21b8">★★★★★</text>'
            f'<text x="1165" y="2326" text-anchor="middle" font-family="{sans}" font-size="38" '
            f'font-weight="900" fill="#2a2a2a">서점형 표지</text>'
        )
        author_block = (
            f'<rect x="72" y="2192" width="610" height="190" fill="#050505" opacity="0.94"/>'
            f'<text x="376" y="2258" text-anchor="middle" font-family="{sans}" font-size="34" '
            f'font-weight="900" fill="#ffffff" letter-spacing="2">{author_esc}</text>'
            f'<text x="376" y="2317" text-anchor="middle" font-family="{sans}" font-size="26" '
            f'font-weight="800" fill="#cfcfcf">AUTHOR EDITION</text>'
        )
    elif template == "literary_watercolor":
        literary_lines = _wrap_title_lines(title, max_chars=13, max_lines=2)
        literary_longest = max((_text_visual_len(line) for line in literary_lines), default=4)
        literary_size = 78 if literary_longest <= 11 else 64
        literary_lh = int(literary_size * 1.22)
        literary_y = 410 - int((len(literary_lines) - 1) * literary_lh / 2)
        lit_sub = _wrap_svg_lines(subtitle, max_chars=26, max_lines=2)
        bg = """
        <defs>
            <linearGradient id="litSky" x1="0%" y1="0%" x2="0%" y2="100%">
                <stop offset="0%" stop-color="#fae1ce"/>
                <stop offset="36%" stop-color="#f4d7d5"/>
                <stop offset="62%" stop-color="#cfe4ef"/>
                <stop offset="100%" stop-color="#eef1df"/>
            </linearGradient>
            <linearGradient id="litWater" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#d8e4f1"/>
                <stop offset="55%" stop-color="#eef2e4"/>
                <stop offset="100%" stop-color="#fff2ce"/>
            </linearGradient>
            <filter id="softBlur"><feGaussianBlur stdDeviation="18"/></filter>
        </defs>
        <rect width="1600" height="2560" fill="#fbf5df" filter="url(#paperNoise)"/>
        <rect x="0" y="0" width="1600" height="1682" fill="url(#litSky)"/>
        <ellipse cx="330" cy="210" rx="460" ry="86" fill="#fff3df" opacity="0.7" filter="url(#softBlur)"/>
        <ellipse cx="1030" cy="275" rx="500" ry="105" fill="#fff8e6" opacity="0.58" filter="url(#softBlur)"/>
        <ellipse cx="730" cy="438" rx="660" ry="130" fill="#f5d2c8" opacity="0.34" filter="url(#softBlur)"/>
        <path d="M-80 1510 C210 1080 430 935 678 1210 C858 960 1092 700 1710 1320 V1682 H-80 Z" fill="#9bb6d6" opacity="0.48"/>
        <path d="M-90 1610 C210 1240 440 1088 704 1338 C878 1130 1110 935 1705 1428 V1682 H-90 Z" fill="#708fc4" opacity="0.38"/>
        <path d="M-20 1715 C270 1500 535 1460 806 1585 C1040 1420 1326 1415 1650 1618 V1810 H-20 Z" fill="url(#litWater)" opacity="0.92"/>
        <path d="M905 1365 C920 1330 950 1330 961 1368 L970 1505 H892 Z" fill="#25304f" opacity="0.78"/>
        <line x1="931" y1="1372" x2="931" y2="1298" stroke="#25304f" stroke-width="7" opacity="0.55"/>
        <rect x="0" y="1682" width="1600" height="878" fill="#fbf5df"/>
        <path d="M0 1808 L800 1960 L1600 1808" fill="none" stroke="#d7cfb6" stroke-width="5"/>
        <rect x="92" y="92" width="1416" height="2376" fill="none" stroke="#ffffff" stroke-width="10" opacity="0.48"/>
        """
        kicker = (
            f'<text x="800" y="{literary_y}" text-anchor="middle" font-family="{sans}" '
            f'font-size="{literary_size}" font-weight="800" fill="#5260b3" letter-spacing="5">'
            f'{_svg_tspans(literary_lines, 800, literary_lh)}</text>'
        )
        title_block = ""
        sub_block = (
            f'<text x="800" y="555" text-anchor="middle" font-family="{serif}" '
            f'font-size="34" font-weight="500" fill="#7a6ca7" letter-spacing="2">'
            f'{_svg_tspans(lit_sub or ["A quiet premium edition"], 800, 46)}</text>'
            f'<text x="800" y="755" text-anchor="middle" font-family="{sans}" font-size="34" '
            f'font-weight="800" fill="#5962ad" letter-spacing="3">{author_esc}</text>'
            f'<text x="800" y="1870" text-anchor="middle" font-family="{serif}" font-size="42" '
            f'font-weight="700" fill="#6a5ca8">“읽는 순간, 마음에 오래 남는 한 권.”</text>'
            f'<text x="360" y="2200" text-anchor="middle" font-family="{sans}" font-size="92" fill="#5b4ca1">‹</text>'
            f'<text x="535" y="2200" text-anchor="middle" font-family="{sans}" font-size="92" fill="#5b4ca1">›</text>'
            f'<text x="448" y="2140" text-anchor="middle" font-family="{sans}" font-size="31" font-weight="900" fill="#5b4ca1">감성문장</text>'
            f'<text x="448" y="2190" text-anchor="middle" font-family="{sans}" font-size="28" font-weight="800" fill="#5b4ca1">프리미엄 에디션</text>'
            f'<text x="710" y="2200" text-anchor="middle" font-family="{sans}" font-size="92" fill="#5b4ca1">‹</text>'
            f'<text x="890" y="2200" text-anchor="middle" font-family="{sans}" font-size="92" fill="#5b4ca1">›</text>'
            f'<text x="800" y="2140" text-anchor="middle" font-family="{sans}" font-size="31" font-weight="900" fill="#5b4ca1">수채화 표지</text>'
            f'<text x="800" y="2190" text-anchor="middle" font-family="{sans}" font-size="28" font-weight="800" fill="#5b4ca1">서점 진열형</text>'
            f'<text x="1068" y="2200" text-anchor="middle" font-family="{sans}" font-size="92" fill="#5b4ca1">‹</text>'
            f'<text x="1248" y="2200" text-anchor="middle" font-family="{sans}" font-size="92" fill="#5b4ca1">›</text>'
            f'<text x="1158" y="2140" text-anchor="middle" font-family="{sans}" font-size="31" font-weight="900" fill="#5b4ca1">독자 추천</text>'
            f'<text x="1158" y="2190" text-anchor="middle" font-family="{sans}" font-size="28" font-weight="800" fill="#5b4ca1">소장 가치</text>'
        )
        author_block = (
            f'<text x="800" y="2398" text-anchor="middle" font-family="{serif}" '
            f'font-size="38" font-weight="600" fill="#786aa7" letter-spacing="5">CASHMAKER LIBRARY</text>'
        )
    elif template == "modern_editorial":
        bg = """
        <rect width="1600" height="2560" fill="url(#softPaper)" filter="url(#paperNoise)"/>
        <rect x="94" y="94" width="1412" height="2372" fill="none" stroke="#161616" stroke-width="3"/>
        <rect x="128" y="128" width="1344" height="2304" fill="none" stroke="#b98b42" stroke-width="1.4"/>
        <rect x="128" y="340" width="1344" height="8" fill="#1d1d1d"/>
        <rect x="128" y="2018" width="1344" height="116" fill="#191919"/>
        <circle cx="1260" cy="505" r="116" fill="none" stroke="#a46d43" stroke-width="14"/>
        <line x1="210" y1="1770" x2="1390" y2="1770" stroke="#1d1d1d" stroke-width="3"/>
        """
        kicker = (
            f'<text x="210" y="260" font-family="{sans}" font-size="34" font-weight="800" '
            f'fill="#1d1d1d" letter-spacing="6">BESTSELLER EDITION</text>'
        )
        title_block = title_svg(x=210, y=840, fill="#111111", anchor="start", weight=900, letter=-1, size=min(title_size + 10, 215))
        sub_block = subtitle_svg(x=210, y=1505, fill="#6e5634", anchor="start", size=43, line_h=60, width=21)
        author_block = (
            f'<text x="800" y="2088" text-anchor="middle" font-family="{sans}" '
            f'font-size="44" font-weight="800" fill="#f7f1df" letter-spacing="8">{author_esc}</text>'
        )
    elif template == "executive_teal":
        bg = """
        <linearGradient id="tealBg" x1="0%" y1="0%" x2="100%" y2="100%">
            <stop offset="0%" stop-color="#0f292b"/>
            <stop offset="55%" stop-color="#102024"/>
            <stop offset="100%" stop-color="#071113"/>
        </linearGradient>
        <rect width="1600" height="2560" fill="url(#tealBg)" filter="url(#paperNoise)"/>
        <rect x="116" y="116" width="1368" height="2328" fill="#f4ead8"/>
        <rect x="156" y="156" width="1288" height="2248" fill="none" stroke="#0f292b" stroke-width="2"/>
        <rect x="156" y="156" width="290" height="2248" fill="#12383b"/>
        <rect x="446" y="156" width="998" height="2248" fill="#f4ead8"/>
        <rect x="446" y="156" width="18" height="2248" fill="#c58a41"/>
        <path d="M146 2040 C390 1950 510 2070 740 1995 S1150 1880 1454 1994" fill="none" stroke="#c58a41" stroke-width="7" opacity="0.45"/>
        """
        kicker = (
            f'<text x="300" y="330" text-anchor="middle" font-family="{sans}" font-size="28" '
            f'font-weight="900" fill="#f4ead8" letter-spacing="5" transform="rotate(-90 300 330)">PRACTICAL GUIDE</text>'
        )
        exec_lines = _wrap_title_lines(title, max_chars=5, max_lines=4)
        exec_longest = max((_text_visual_len(line) for line in exec_lines), default=4)
        exec_size = 172 if exec_longest <= 4.2 else 146
        exec_line_h = int(exec_size * 1.1)
        exec_y = 835 - int((len(exec_lines) - 1) * exec_line_h / 2)
        title_block = (
            f'<text x="545" y="{exec_y}" text-anchor="start" font-family="{serif}" '
            f'font-size="{exec_size}" font-weight="900" fill="#111817" letter-spacing="-2">'
            f'{_svg_tspans(exec_lines, 545, exec_line_h)}</text>'
        )
        sub_block = subtitle_svg(x=545, y=1540, fill="#405f5a", anchor="start", size=43, line_h=61, width=20)
        author_block = (
            f'<text x="545" y="2190" font-family="{sans}" font-size="42" font-weight="850" '
            f'fill="#12383b" letter-spacing="6">{author_esc}</text>'
            f'<text x="545" y="2254" font-family="{sans}" font-size="24" font-weight="700" '
            f'fill="#8c6b38" letter-spacing="3">CASHMAKER BOOKS</text>'
        )
    elif template == "archive_red":
        bg = """
        <linearGradient id="archiveBg" x1="0%" y1="0%" x2="0%" y2="100%">
            <stop offset="0%" stop-color="#efe8d9"/>
            <stop offset="100%" stop-color="#d9cab0"/>
        </linearGradient>
        <rect width="1600" height="2560" fill="url(#archiveBg)" filter="url(#paperNoise)"/>
        <rect x="0" y="0" width="1600" height="565" fill="#9d2f24"/>
        <rect x="0" y="565" width="1600" height="34" fill="#2b2420"/>
        <rect x="140" y="700" width="1320" height="1" stroke="#b49d78" stroke-width="4"/>
        <rect x="140" y="1880" width="1320" height="1" stroke="#b49d78" stroke-width="4"/>
        <rect x="118" y="118" width="1364" height="2324" fill="none" stroke="#2b2420" stroke-width="2"/>
        <rect x="142" y="142" width="1316" height="2276" fill="none" stroke="#9d2f24" stroke-width="1.4" opacity="0.75"/>
        <circle cx="800" cy="2052" r="86" fill="#2b2420"/>
        <circle cx="800" cy="2052" r="62" fill="none" stroke="#d7b86a" stroke-width="4"/>
        """
        kicker = (
            f'<text x="800" y="240" text-anchor="middle" font-family="{sans}" font-size="34" '
            f'font-weight="900" fill="#fff4df" letter-spacing="10">NEW CLASSIC</text>'
            f'<text x="800" y="317" text-anchor="middle" font-family="{serif}" font-size="54" '
            f'font-weight="700" fill="#fff4df">읽히는 책의 기준</text>'
        )
        title_block = title_svg(x=800, y=1025, fill="#211914", anchor="middle", weight=900, letter=-1, size=min(title_size + 8, 210))
        sub_block = subtitle_svg(x=800, y=1620, fill="#6e3a2e", anchor="middle", size=43, line_h=60, width=22)
        author_block = (
            f'<text x="800" y="2225" text-anchor="middle" font-family="{sans}" font-size="42" '
            f'font-weight="850" fill="#2b2420" letter-spacing="9">{author_esc}</text>'
        )
    elif template == "paper_luxe":
        bg = """
        <rect width="1600" height="2560" fill="#f6f0e2" filter="url(#paperNoise)"/>
        <rect x="78" y="78" width="1444" height="2404" rx="0" fill="none" stroke="#b58a38" stroke-width="4"/>
        <rect x="110" y="110" width="1380" height="2340" fill="none" stroke="#342c23" stroke-width="1.4"/>
        <path d="M800 250 C910 250 995 335 995 445 C995 555 910 640 800 640 C690 640 605 555 605 445 C605 335 690 250 800 250 Z" fill="none" stroke="#b58a38" stroke-width="5"/>
        <path d="M590 1260 H1010 M650 1300 H950 M705 1340 H895" stroke="#b58a38" stroke-width="4"/>
        <path d="M210 360 C330 330 405 330 525 360 M1075 360 C1195 330 1270 330 1390 360" fill="none" stroke="#342c23" stroke-width="2"/>
        <path d="M210 2200 C330 2230 405 2230 525 2200 M1075 2200 C1195 2230 1270 2230 1390 2200" fill="none" stroke="#342c23" stroke-width="2"/>
        """
        kicker = (
            f'<text x="800" y="470" text-anchor="middle" font-family="{sans}" font-size="28" '
            f'font-weight="850" fill="#342c23" letter-spacing="7">CASHMAKER LIBRARY</text>'
        )
        title_block = title_svg(x=800, y=900, fill="#221b14", anchor="middle", weight=900, letter=-1, size=min(title_size + 5, 205))
        sub_block = subtitle_svg(x=800, y=1525, fill="#755c35", anchor="middle", size=42, line_h=59, width=22)
        author_block = (
            f'<text x="800" y="2255" text-anchor="middle" font-family="{serif}" font-size="44" '
            f'font-weight="700" fill="#342c23" letter-spacing="7">{author_esc}</text>'
        )
    else:  # signature_noir
        bg = """
        <radialGradient id="noirGlow" cx="50%" cy="38%" r="72%">
            <stop offset="0%" stop-color="#28231b"/>
            <stop offset="58%" stop-color="#11100e"/>
            <stop offset="100%" stop-color="#050505"/>
        </radialGradient>
        <rect width="1600" height="2560" fill="url(#noirGlow)" filter="url(#paperNoise)"/>
        <rect x="72" y="72" width="1456" height="2416" fill="none" stroke="url(#foil)" stroke-width="3" opacity="0.78"/>
        <rect x="112" y="112" width="1376" height="2336" fill="none" stroke="#6d5120" stroke-width="1" opacity="0.75"/>
        <rect x="0" y="0" width="210" height="2560" fill="#0b0a08" opacity="0.88"/>
        <rect x="210" y="0" width="10" height="2560" fill="url(#foil)" opacity="0.9"/>
        <path d="M358 398 H1240 M358 2112 H1240" stroke="url(#foil)" stroke-width="5"/>
        <path d="M620 1725 C735 1665 865 1665 980 1725" fill="none" stroke="url(#foil)" stroke-width="5" opacity="0.65"/>
        """
        kicker = (
            f'<text x="800" y="315" text-anchor="middle" font-family="{sans}" font-size="31" '
            f'font-weight="850" fill="#d7b86a" letter-spacing="9">PREMIUM PAPERBACK</text>'
        )
        title_block = title_svg(x=800, y=845, fill="url(#foil)", anchor="middle", weight=900, letter=-1, size=min(title_size + 12, 218))
        sub_block = subtitle_svg(x=800, y=1555, fill="#d8c7a5", anchor="middle", size=44, line_h=61, width=21)
        author_block = (
            f'<text x="800" y="2190" text-anchor="middle" font-family="{sans}" font-size="42" '
            f'font-weight="850" fill="#f7f0dc" letter-spacing="10">{author_esc}</text>'
            f'<text x="800" y="2260" text-anchor="middle" font-family="{sans}" font-size="24" '
            f'font-weight="700" fill="#a89162" letter-spacing="5">CASHMAKER</text>'
        )

    return (
        '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 1600 2560" '
        'width="100%" preserveAspectRatio="xMidYMid meet">'
        f'{fonts}{base_defs}{bg}{kicker}{title_block}{sub_block}{author_block}</svg>'
    )


def estimate_docx_pages(chapters, outline):
    """워드(A5) 실제 출력 구조를 그대로 모사한 페이지 수 추정.

    워드는 본문 10.5pt / 줄간격 1.85 / A5 + 소제목마다 새 페이지로 나뉘고,
    표지·판권·프롤로그·에필로그·저자소개·챕터 오프너가 더해진다.
    기존의 '글자수//500'은 이 구조를 반영하지 않아 워드 페이지 수와 크게 어긋났다.
    """
    import math
    CPP = 330  # A5 11.5pt/1.85 + 소제목별 페이지 나눔 실측(페이지당 약 320자) 반영
    FRONT = 2  # 표지 + 판권
    PROLOGUE = 2
    EPILOGUE = 2
    AUTHOR = 1

    pages = 0
    chapter_with_content = 0
    for ch in (outline or []):
        cd = (chapters or {}).get(ch)
        if not cd:
            continue
        sub_pages = 0
        has_content = False
        for s in cd.get('subtopics', []):
            content = cd.get('subtopic_data', {}).get(s, {}).get('content', '')
            if content:
                has_content = True
                chars = len(content.replace(' ', '').replace('\n', ''))
                sub_pages += max(1, math.ceil(chars / CPP))  # 소제목은 새 페이지에서 시작
        if has_content:
            chapter_with_content += 1
            pages += 1          # 챕터 오프너 페이지
            pages += sub_pages

    if chapter_with_content == 0:
        return 0
    return FRONT + PROLOGUE + pages + EPILOGUE + AUTHOR


def extract_video_id(url):
    """YouTube URL에서 video ID 추출"""
    patterns = [
        r'(?:youtube\.com\/watch\?v=|youtu\.be\/|youtube\.com\/embed\/)([^&\n?#]+)',
        r'youtube\.com\/watch\?.*v=([^&\n?#]+)',
    ]
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    return None

def get_youtube_transcript(video_id):
    """YouTube 자막 가져오기"""
    if not YOUTUBE_TRANSCRIPT_AVAILABLE:
        return None, "youtube-transcript-api가 설치되지 않았습니다. pip install youtube-transcript-api"

    try:
        # 한국어 자막 우선, 없으면 영어, 없으면 자동생성
        transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)

        transcript = None
        # 수동 자막 먼저 시도
        for lang in ['ko', 'en']:
            try:
                transcript = transcript_list.find_transcript([lang])
                break
            except:
                continue

        # 수동 자막 없으면 자동 생성 자막
        if not transcript:
            try:
                transcript = transcript_list.find_generated_transcript(['ko', 'en'])
            except:
                # 아무 자막이나 가져오기
                for t in transcript_list:
                    transcript = t
                    break

        if transcript:
            fetched = transcript.fetch()
            full_text = ' '.join([item['text'] for item in fetched])
            return full_text, None
        else:
            return None, "자막을 찾을 수 없습니다"

    except Exception as e:
        return None, f"자막 추출 오류: {str(e)[:100]}"

def analyze_youtube_video_direct(url):
    """YouTube 영상 자막 기반 분석 (빠르고 정확)"""
    api_key = get_api_key()
    if not api_key:
        st.error("API 키를 입력해주세요")
        return None

    # 1. Video ID 추출
    video_id = extract_video_id(url)
    if not video_id:
        st.error("올바른 YouTube URL이 아닙니다")
        return None

    # 2. 자막 가져오기
    transcript, error = get_youtube_transcript(video_id)
    if error:
        st.warning(f"자막 추출 실패: {error}")
        st.info("자막이 없는 영상입니다. 텍스트 입력으로 직접 내용을 입력해주세요.")
        return None

    if not transcript or len(transcript) < 50:
        st.warning("자막 내용이 너무 짧습니다")
        return None

    # 3. 자막 기반 분석
    prompt = f"""다음은 YouTube 영상의 자막입니다. 이 내용을 분석해주세요.

[자막 내용]
{transcript[:15000]}

[분석 요청]
위 자막 내용을 바탕으로 분석해주세요. 자막에 없는 내용은 추측하지 마세요.

JSON 형식으로 응답:
{{
    "title": "영상의 핵심 주제 (자막 기반 추론)",
    "creator": "알 수 없음",
    "main_topic": "메인 주제 한 줄 요약",
    "key_points": ["핵심 포인트 1", "핵심 포인트 2", "핵심 포인트 3", "핵심 포인트 4", "핵심 포인트 5"],
    "detailed_notes": ["상세 내용 1", "상세 내용 2", "상세 내용 3"],
    "actionable_tips": ["실천 팁 1", "실천 팁 2", "실천 팁 3"],
    "quotes": ["인상적인 문장 1", "인상적인 문장 2"],
    "vocabulary": [{{"term": "용어", "definition": "설명"}}],
    "study_questions": ["학습 질문 1", "학습 질문 2"],
    "summary": "전체 내용 5-7문장 요약"
}}"""

    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-2.0-flash')
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        st.error(f"분석 오류: {str(e)[:150]}")
        return None

def get_full_content():
    full = ""
    for ch in st.session_state.get('outline', []):
        if ch in st.session_state.get('chapters', {}):
            ch_data = st.session_state['chapters'][ch]
            ch_content = ""
            for s in ch_data.get('subtopics', []):
                c = ch_data.get('subtopic_data', {}).get(s, {}).get('content', '')
                if c:
                    ch_content += f"\n\n【{s}】\n\n{clean_content(c)}"
            if ch_content:
                full += f"\n\n{'='*50}\n{ch}\n{'='*50}{ch_content}"
    return full.strip()

# ==========================================
# 전자책 워드 디자인 폰트 (출판물 톤)
# KoPub/Noto 계열이 있으면 종이책 톤으로 보이고, 없으면 Word가 기본 명조 계열로 대체한다.
# ==========================================
EBOOK_SERIF_KR = 'KoPubWorldBatang Medium'
EBOOK_SERIF_LATIN = 'Georgia'
EBOOK_SANS_KR = 'Pretendard'
EBOOK_SANS_LATIN = 'Arial'

# 출판 톤 색상 팔레트
_INK = (30, 28, 24)        # 본문/제목 (따뜻한 먹색)
_SOFT = (95, 86, 74)       # 부제/캡션
_FAINT = (154, 140, 118)   # 라벨/번호
_HAIR = (217, 207, 190)    # 가는 선
_GOLD = (178, 135, 52)     # 차분한 금박 포인트
_PAPER = (247, 242, 231)   # 본문 종이톤
_DEEP = (31, 28, 24)       # 표지/챕터 짙은 색
_TERRACOTTA = (151, 72, 48)


def create_ebook_docx(title, subtitle, author, chapters_data, outline, interview_data=None):
    """베스트셀러 출판물 스타일 워드 문서 생성 (프리미엄 에디토리얼)."""
    if not DOCX_AVAILABLE:
        return None, "python-docx 패키지가 필요합니다: pip install python-docx"

    try:
        BRAND = "CASHMAKER"
        doc = Document()

        # 기본 문단 간격 0 (빈 문단이 멋대로 커지는 것 방지)
        normal = doc.styles['Normal']
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(0)
        normal.paragraph_format.line_spacing = 1.0

        def _apply_section(sec):
            sec.page_width = Cm(14.8)
            sec.page_height = Cm(21)
            sec.left_margin = Cm(1.85)
            sec.right_margin = Cm(1.85)
            sec.top_margin = Cm(2.05)
            sec.bottom_margin = Cm(2.05)

        def _apply_cover_section(sec):
            sec.page_width = Cm(14.8)
            sec.page_height = Cm(21)
            sec.left_margin = Cm(0.72)
            sec.right_margin = Cm(0.72)
            sec.top_margin = Cm(0.72)
            sec.bottom_margin = Cm(0.72)
            sec.different_first_page_header_footer = True

        _apply_section(doc.sections[0])

        # ── 폰트/런 ──
        def set_font(run, size, bold=False, color=None, italic=False,
                     serif=False, track=None):
            run.font.size = Pt(size)
            fam_latin = EBOOK_SERIF_LATIN if serif else EBOOK_SANS_LATIN
            fam_kr = EBOOK_SERIF_KR if serif else EBOOK_SANS_KR
            run.font.name = fam_latin
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.find(qn('w:rFonts'))
            if rfonts is None:
                rfonts = OxmlElement('w:rFonts')
                rpr.append(rfonts)
            rfonts.set(qn('w:ascii'), fam_latin)
            rfonts.set(qn('w:hAnsi'), fam_latin)
            rfonts.set(qn('w:eastAsia'), fam_kr)
            run.bold = bold
            run.italic = italic
            if color:
                run.font.color.rgb = RGBColor(*color)
            if track is not None:
                sp = OxmlElement('w:spacing')
                sp.set(qn('w:val'), str(int(track * 20)))
                rpr.append(sp)

        def vspace(pts):
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_before = Pt(0); pf.space_after = Pt(0)
            pf.line_spacing = Pt(pts)
            return p

        def track_text(text, n=1):
            return (" " * n).join(list(str(text)))

        def hairline(align=WD_ALIGN_PARAGRAPH.CENTER, width_cm=None, color=_HAIR,
                     space_before=0, space_after=12, size=6):
            p = doc.add_paragraph()
            p.alignment = align
            pf = p.paragraph_format
            pf.space_before = Pt(space_before); pf.space_after = Pt(space_after)
            if width_cm and align == WD_ALIGN_PARAGRAPH.CENTER:
                side = (10.8 - width_cm) / 2
                if side > 0:
                    pf.left_indent = Cm(side); pf.right_indent = Cm(side)
            pPr = p._p.get_or_add_pPr()
            pbdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), str(size))
            bottom.set(qn('w:space'), '1')
            _col = color if isinstance(color, str) else '{:02X}{:02X}{:02X}'.format(*color)
            bottom.set(qn('w:color'), _col)
            pbdr.append(bottom)
            pPr.append(pbdr)
            return p

        def ornament(glyph="✦", color=_FAINT, size=11, space=22):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(space)
            p.paragraph_format.space_after = Pt(space)
            set_font(p.add_run(glyph), size, serif=False, color=color)
            return p

        def _shade_para(p, hex_fill):
            """문단 배경색 칠하기."""
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_fill)
            pPr.append(shd)

        def _box_border(p, color_hex, sides=('top', 'bottom', 'left', 'right'), sz=6):
            """문단 사면(또는 일부) 테두리 — 박스 효과."""
            pPr = p._p.get_or_add_pPr()
            pbdr = OxmlElement('w:pBdr')
            for s in sides:
                el = OxmlElement(f'w:{s}')
                el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), str(sz))
                el.set(qn('w:space'), '6'); el.set(qn('w:color'), color_hex)
                pbdr.append(el)
            pPr.append(pbdr)

        def key_summary_box(items, chapter_no):
            """챕터 끝 '핵심 정리' 요약 박스 (베스트셀러식 정리 코너)."""
            # 제목 줄 (다크 배경 + 골드 글자)
            title_p = doc.add_paragraph()
            title_p.paragraph_format.space_before = Pt(20); title_p.paragraph_format.space_after = Pt(0)
            title_p.paragraph_format.left_indent = Pt(2); title_p.paragraph_format.right_indent = Pt(2)
            _shade_para(title_p, '211C17')
            r = title_p.add_run(f"  CHAPTER {chapter_no} · 핵심 정리  ")
            set_font(r, 11, bold=True, serif=False, color=(214, 180, 101), track=0.5)
            title_p.paragraph_format.space_after = Pt(0)
            # 항목들 (연한 배경 박스 안)
            for k, item in enumerate(items):
                ip = doc.add_paragraph()
                ip.paragraph_format.left_indent = Cm(0.2); ip.paragraph_format.right_indent = Cm(0.2)
                ip.paragraph_format.space_before = Pt(0)
                ip.paragraph_format.space_after = Pt(2 if k < len(items) - 1 else 0)
                ip.paragraph_format.line_spacing = 1.5
                _shade_para(ip, 'F8F1E4')
                num_r = ip.add_run(f"  {k+1}  ")
                set_font(num_r, 10.5, bold=True, serif=False, color=_TERRACOTTA)
                txt_r = ip.add_run(item.strip() + "  ")
                set_font(txt_r, 10.5, serif=True, color=_INK)
            # 박스 하단 골드 라인
            hairline(align=WD_ALIGN_PARAGRAPH.LEFT, color=_GOLD, space_before=0, space_after=16, size=10)

        def body_paragraph(text, first=False, size=11.2, ls=1.68):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_font(p.add_run(text), size, serif=True, color=_INK)
            pf = p.paragraph_format
            pf.line_spacing = ls
            pf.space_after = Pt(10)
            if not first:
                pf.first_line_indent = Cm(0.48)
            return p

        def drop_cap(letter, rest, lines=3, size=11.2, ls=1.68):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf = p.paragraph_format
            pf.line_spacing = ls
            pf.space_after = Pt(12)
            if letter:
                r1 = p.add_run(letter)
                set_font(r1, 23, serif=True, bold=True, color=_GOLD)
            if rest:
                r2 = p.add_run(rest)
                set_font(r2, size, serif=True, color=_INK)
            return p

        def _cell_shade(cell, hex_fill):
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_fill)
            cell._tc.get_or_add_tcPr().append(shd)

        def _cell_para(cell, text, size, *, serif=True, bold=False, italic=False,
                       color=(255, 255, 255), align=WD_ALIGN_PARAGRAPH.CENTER,
                       track=None, space_before=0, space_after=0, ls=1.2, first=True):
            p = cell.add_paragraph() if not first else cell.paragraphs[0]
            p.alignment = align
            pf = p.paragraph_format
            pf.space_before = Pt(space_before); pf.space_after = Pt(space_after); pf.line_spacing = ls
            if text:
                set_font(p.add_run(text), size, serif=serif, bold=bold, italic=italic, color=color, track=track)
            return p

        def _no_cell_borders(cell):
            tcPr = cell._tc.get_or_add_tcPr()
            tcb = OxmlElement('w:tcBorders')
            for s in ('top', 'bottom', 'left', 'right'):
                e = OxmlElement(f'w:{s}'); e.set(qn('w:val'), 'nil'); tcb.append(e)
            tcPr.append(tcb)

        def page_number_footer(section):
            footer = section.footer
            footer.is_linked_to_previous = False
            p = footer.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            set_font(run, 9, serif=True, color=_SOFT)
            f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
            instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
            f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
            run._r.append(f1); run._r.append(instr); run._r.append(f2)

        def full_page_cover(title, subtitle, author, brand):
            """A5 한 페이지를 꽉 쓰는 베스트셀러 브리핑형 종이책 표지."""
            _apply_cover_section(doc.sections[-1])
            tbl = doc.add_table(rows=2, cols=1)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

            top = tbl.rows[0].cells[0]
            bottom = tbl.rows[1].cells[0]
            for row, height in [(tbl.rows[0], 11.55), (tbl.rows[1], 7.8)]:
                row.height = Cm(height)
                row.height_rule = 2  # EXACT (한 페이지 고정)
                row.cells[0].width = Cm(13.36)
                va = OxmlElement('w:vAlign'); va.set(qn('w:val'), 'center')
                row.cells[0]._tc.get_or_add_tcPr().append(va)
                _no_cell_borders(row.cells[0])

            _cell_shade(top, '050505')
            _cell_shade(bottom, 'F7F7F4')

            for cell, margins in [
                (top, [('top', '430'), ('left', '500'), ('bottom', '300'), ('right', '500')]),
                (bottom, [('top', '300'), ('left', '500'), ('bottom', '260'), ('right', '500')]),
            ]:
                tcMar = OxmlElement('w:tcMar')
                for mn, mv in margins:
                    m = OxmlElement(f'w:{mn}'); m.set(qn('w:w'), mv); m.set(qn('w:type'), 'dxa'); tcMar.append(m)
                cell._tc.get_or_add_tcPr().append(tcMar)

            PURPLE = (106, 33, 184)
            CREAM = (255, 255, 255)
            DIM = (190, 190, 190)
            INK = (18, 18, 18)

            # 상단: 실제 베스트셀러형 블랙 재킷 + 보라색 프레임 감성
            _cell_para(top, "━━━━━━━━━━━━━━━━━━━━", 9, serif=False, bold=True,
                       color=PURPLE, space_before=0, space_after=14, first=True)
            if subtitle:
                sp = top.add_paragraph()
                sp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                sp.paragraph_format.space_before = Pt(0); sp.paragraph_format.space_after = Pt(22)
                sp.paragraph_format.line_spacing = 1.22
                set_font(sp.add_run(subtitle), 15.8, serif=True, bold=True, color=CREAM)
            else:
                _cell_para(top, "품격 있는 전자책을 위한", 15.8, serif=True, bold=True,
                           color=CREAM, space_after=22, first=False, align=WD_ALIGN_PARAGRAPH.LEFT)

            # 제목: 레퍼런스처럼 크게, 타이트하게, 흰색 명조 중심
            tlines = _bestseller_title_lines(title)
            longest = max((_text_visual_len(line) for line in tlines), default=4)
            if longest <= 4.2:
                cover_title_size = 39
            elif longest <= 6.2:
                cover_title_size = 34
            else:
                cover_title_size = 28
            if len(tlines) >= 4:
                cover_title_size = min(cover_title_size, 28)
            for line in tlines:
                _cell_para(top, line, cover_title_size, serif=True, bold=True, color=CREAM,
                           ls=1.0, space_after=0, first=False, align=WD_ALIGN_PARAGRAPH.LEFT)

            tagline = top.add_paragraph()
            tagline.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            tagline.paragraph_format.space_before = Pt(24); tagline.paragraph_format.space_after = Pt(0)
            set_font(tagline.add_run("INSIGHT · STRATEGY · MONEY · BRANDING"), 8.6,
                     serif=False, bold=True, color=DIM, track=0.7)

            # 하단: 흰색 띠지 영역 + 보라색 판매 문구
            _cell_para(bottom, "매일 읽히는 한 권의 프리미엄 브리핑", 12.4,
                       serif=False, bold=True, color=INK, space_before=2, space_after=8, first=True)
            _cell_para(bottom, "지금, 독자의 판단을 움직이는", 14.4,
                       serif=False, bold=True, color=INK, space_after=2, first=False)
            _cell_para(bottom, "진짜 지식", 22.5,
                       serif=False, bold=True, color=PURPLE, space_after=22, first=False)

            badge = bottom.add_paragraph()
            badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
            badge.paragraph_format.space_before = Pt(0); badge.paragraph_format.space_after = Pt(18)
            _shade_para(badge, '050505')
            set_font(badge.add_run("  ★★★★★  종이책 베스트셀러형 표지  "), 9.8,
                     serif=False, bold=True, color=CREAM, track=0.4)

            _cell_para(bottom, f"{author or '저자'} 지음", 12.8, serif=False, bold=True,
                       color=INK, space_after=0, first=False)
            _cell_para(bottom, "CASHMAKER PREMIUM EDITION", 8.2, serif=False, bold=True,
                       color=PURPLE, space_after=0, first=False, track=0.8)
            next_sec = doc.add_section(WD_SECTION.NEW_PAGE)
            _apply_section(next_sec)



        def running_header(section, text):
            header = section.header
            header.is_linked_to_previous = False
            p = header.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(p.add_run(track_text(text, 1)), 8, serif=False, color=_FAINT)

        def add_bookmark(paragraph, name):
            clean = re.sub(r'[^\w가-힣]', '_', name)[:40]
            bs = OxmlElement('w:bookmarkStart')
            bs.set(qn('w:id'), str(abs(hash(clean)) % 100000))
            bs.set(qn('w:name'), clean)
            be = OxmlElement('w:bookmarkEnd')
            be.set(qn('w:id'), str(abs(hash(clean)) % 100000))
            paragraph._p.insert(0, bs)
            paragraph._p.append(be)

        def add_hyperlink(paragraph, text, name, size=10, bold=False, color=_SOFT, serif=True):
            clean = re.sub(r'[^\w가-힣]', '_', name)[:40]
            hl = OxmlElement('w:hyperlink'); hl.set(qn('w:anchor'), clean)
            r = OxmlElement('w:r'); rPr = OxmlElement('w:rPr')
            rf = OxmlElement('w:rFonts')
            fam_latin = EBOOK_SERIF_LATIN if serif else EBOOK_SANS_LATIN
            fam_kr = EBOOK_SERIF_KR if serif else EBOOK_SANS_KR
            rf.set(qn('w:ascii'), fam_latin); rf.set(qn('w:hAnsi'), fam_latin); rf.set(qn('w:eastAsia'), fam_kr)
            rPr.append(rf)
            sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(int(size * 2))); rPr.append(sz)
            if bold:
                rPr.append(OxmlElement('w:b'))
            c = OxmlElement('w:color'); c.set(qn('w:val'), '{:02X}{:02X}{:02X}'.format(*color)); rPr.append(c)
            r.append(rPr)
            t = OxmlElement('w:t'); t.text = text; r.append(t)
            hl.append(r); paragraph._p.append(hl)

        # ── 표 헬퍼 (기존 로직 보존) ──
        def parse_table_data(text):
            lines = text.strip().split('\n'); table_data = []

            def _is_separator_cells(cells):
                # 모든 셀이 ---, :---:, --- 같은 구분선 기호로만 이루어졌으면 구분선 행
                if not cells:
                    return False
                for c in cells:
                    cc = c.strip()
                    if cc == '' or re.fullmatch(r'[:\-\s─━]+', cc) is None:
                        return False
                return True

            for line in lines:
                line = line.strip()
                if not line:
                    continue
                if re.match(r'^\|[\s\-:]+\|$', line):
                    continue
                if re.match(r'^[─━┌┬┐├┼┤└┴┘│┃]+$', line):
                    continue
                if re.match(r'^[\s\-]+$', line) and len(line.replace(' ', '').replace('-', '')) == 0:
                    continue
                if line.startswith('|') and line.endswith('|'):
                    cells = [c.strip() for c in line.split('|')]; cells = [c for c in cells if c]
                    if cells and not _is_separator_cells(cells):
                        table_data.append(cells)
                elif '|' in line and not line.startswith('|'):
                    cells = [c.strip() for c in line.split('|')]; cells = [c for c in cells if c]
                    if cells and not _is_separator_cells(cells):
                        table_data.append(cells)
                elif '\t' in line:
                    cells = [c.strip() for c in line.split('\t')]; cells = [c for c in cells if c]
                    if len(cells) >= 2 and not _is_separator_cells(cells):
                        table_data.append(cells)
                elif ':' in line and not line.startswith('http'):
                    parts = line.split(':', 1)
                    if len(parts) == 2 and len(parts[0]) < 30:
                        table_data.append([parts[0].strip(), parts[1].strip()])
            return table_data

        def add_premium_table(table_data):
            if not table_data or len(table_data) < 1:
                return None
            rows = len(table_data); cols = max(len(r) for r in table_data)
            is_comp = cols == 2 and rows >= 2
            table = doc.add_table(rows=rows, cols=cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            for ci in range(cols):
                for row in table.rows:
                    if ci < len(row.cells):
                        width = (Cm(3) if ci == 0 else Cm(7.8)) if is_comp else Cm(10.8 / cols)
                        row.cells[ci].width = width
            for i, row_data in enumerate(table_data):
                row = table.rows[i]
                for j, txt in enumerate(row_data):
                    if j < cols:
                        cell = row.cells[j]; cell.text = ''; para = cell.paragraphs[0]
                        if i == 0:
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            set_font(para.add_run(str(txt)), 9.2, bold=True, color=(255, 248, 235), serif=False)
                            shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), '211C17')
                            cell._tc.get_or_add_tcPr().append(shd)
                        elif is_comp and j == 0:
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            set_font(para.add_run(str(txt)), 9.2, bold=True, color=(86, 60, 38), serif=False)
                            shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), 'EFE6D5')
                            cell._tc.get_or_add_tcPr().append(shd)
                        else:
                            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            para.paragraph_format.left_indent = Pt(6)
                            para.paragraph_format.line_spacing = 1.35
                            set_font(para.add_run(str(txt)), 9.3, color=(48, 44, 39), serif=True)
                            if i % 2 == 0:
                                shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), 'FBF7EF')
                                cell._tc.get_or_add_tcPr().append(shd)
                        tcPr = cell._tc.get_or_add_tcPr()
                        tcMar = OxmlElement('w:tcMar')
                        for mn, mv in [('top', '115'), ('left', '145'), ('bottom', '115'), ('right', '145')]:
                            m = OxmlElement(f'w:{mn}'); m.set(qn('w:w'), mv); m.set(qn('w:type'), 'dxa'); tcMar.append(m)
                        tcPr.append(tcMar)
                        va = OxmlElement('w:vAlign'); va.set(qn('w:val'), 'center'); tcPr.append(va)
            tbl = table._tbl
            tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
            tblBorders = OxmlElement('w:tblBorders')
            for bn in ['top', 'bottom']:
                b = OxmlElement(f'w:{bn}'); b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '10'); b.set(qn('w:color'), '211C17'); tblBorders.append(b)
            for bn in ['left', 'right']:
                b = OxmlElement(f'w:{bn}'); b.set(qn('w:val'), 'nil'); tblBorders.append(b)
            for bn in ['insideH', 'insideV']:
                b = OxmlElement(f'w:{bn}'); b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '4'); b.set(qn('w:color'), 'E6D8BF'); tblBorders.append(b)
            tblPr.append(tblBorders)
            sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(14)
            return table

        def process_content_with_tables(text):
            blocks = []; current = []; tbuf = []
            lines = text.split('\n'); i = 0

            def is_start(line, nxt=None):
                s = line.strip()
                if s.startswith('|') and s.endswith('|') and s.count('|') >= 2:
                    return True
                if '|' in s and len(s.split('|')) >= 2:
                    if any(p.strip() and not re.match(r'^[\s\-:]+$', p) for p in s.split('|')):
                        return True
                if nxt and ':' in s and ':' in nxt:
                    if len(s.split(':')[0].strip()) < 30 and len(nxt.split(':')[0].strip()) < 30:
                        return True
                return False

            def is_cont(line):
                s = line.strip()
                if not s:
                    return False
                if s.startswith('|') and s.endswith('|'):
                    return True
                if re.match(r'^\|[\s\-:]+\|$', s):
                    return True
                if '|' in s:
                    return True
                if ':' in s and len(s.split(':')[0].strip()) < 30:
                    return True
                return False

            while i < len(lines):
                line = lines[i]
                nxt = lines[i + 1] if i + 1 < len(lines) else None
                if is_start(line, nxt):
                    if current:
                        blocks.append(('text', '\n'.join(current))); current = []
                    tbuf = [line]; i += 1
                    while i < len(lines) and is_cont(lines[i]):
                        tbuf.append(lines[i]); i += 1
                    blocks.append(('table', '\n'.join(tbuf)))
                else:
                    current.append(line); i += 1
            if current:
                blocks.append(('text', '\n'.join(current)))
            return blocks

        # ════════════════════════ 표지 (풀 페이지 짙은 배경) ════════════════════════
        full_page_cover(title, subtitle, author, BRAND)

        # ════════════════════════ 판권 ════════════════════════
        vspace(255)
        cp_title = doc.add_paragraph(); cp_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(cp_title.add_run(title or "전자책"), 15, bold=True, color=_INK, serif=True)
        cp_title.paragraph_format.space_after = Pt(8)
        hairline(width_cm=1.2, color=_GOLD, space_after=18, size=8)
        cr = [
            f"지은이  {author or '저자'}",
            f"펴낸곳  {BRAND}",
            datetime.now().strftime("발행일  %Y.%m.%d"),
            "",
            "이 책의 저작권은 저자에게 있습니다.",
            "무단 전재와 복제를 금합니다.",
        ]
        for line in cr:
            cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if line:
                set_font(cp.add_run(line), 9.2, color=_SOFT, serif=True)
            cp.paragraph_format.space_after = Pt(4)
        doc.add_page_break()

        # ════════════════════════ 프롤로그 ════════════════════════
        vspace(36)
        pl = doc.add_paragraph(); pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(pl.add_run(track_text("PROLOGUE", 2)), 10, serif=False, color=_FAINT)
        pl.paragraph_format.space_after = Pt(8)
        hairline(width_cm=1.4, color=_GOLD, space_after=30, size=8)

        prologue_text = None
        if interview_data:
            prologue_prompt = f"""당신은 한국 자기계발 베스트셀러 작가입니다. 독자가 첫 문장에서 "이거 내 얘기야"라고 무릎 치고, 마지막 문장에서 "다음 페이지가 너무 궁금해"라며 책장을 넘기게 만드는 프롤로그를 작성하세요.

[저자 정보 - 참고용, 그대로 복사하지 말 것]
- 분야: {interview_data.get('field', '')}
- 경력: {interview_data.get('experience_years', '')}
- 책 주제: {interview_data.get('topic', '')}
- 타겟 독자: {interview_data.get('target_reader', '')}
- 독자의 고민: {interview_data.get('target_problem', '')}
- 집필 동기: {interview_data.get('why_write', '')}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🎯 프롤로그의 단 하나의 미션
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

독자가 첫 문장부터 마지막 문장까지 한숨에 읽고, 자기 이야기처럼 공감하면서, 본문이 미치도록 궁금해서 1장으로 넘어갈 수밖에 없게 만들 것.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
✍️ 톤: 공감 후킹 + 스토리텔링 + 호기심 갭
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[5막 구조 - 베스트셀러 프롤로그 공식]

1막) 누구나 겪는 장면 (3~4문장) - 공감 끌어올리기
   - 독자가 "이거 내 얘기야"라고 즉시 떠올리는 구체적 장면
   - 시간·장소·감정의 디테일 (예: "수요일 밤 11시. 또 인스타를 켰다. 친구 OO이 부자가 돼 있었다.")
   - 일반론 절대 금지. 손에 잡히는 장면 하나

2막) 저자의 고백 (3~4문장) - "저도 그랬습니다"
   - 구체적 실패담 + 그때의 감정 (수치심, 막막함, 분노 등)
   - 진심 어린 톤 (위에서 내려다보지 말 것)
   - 약점 노출이 신뢰를 만든다

3막) 결정적 전환 (2~3문장) - 그러던 어느 날
   - 발견 / 사건 / 만남의 순간
   - "그날 알게 된 한 가지가 모든 걸 바꿨다"
   - 그게 뭔지 다 말하지 말고 살짝 가리기

4막) 약속 + 변화 (2문장) - 이 책이 줄 것
   - 본문이 다룰 핵심 변화를 구체적으로 (숫자/기간 포함)
   - 추상 X, 구체 O (예: "이 책은 그 90초가 어떻게 작동하는지 단계별로 풀어냅니다.")

5막) 본문 미끼 (1~2문장) - 페이지 넘기게 하기
   - 본문 1장이 다룰 가장 강한 장면 또는 통찰의 일부만 흘리기
   - "그런데 그 출발점은 의외의 한 가지였습니다." 같은 호기심 갭
   - "다음 페이지부터 시작됩니다" 같은 직접 안내는 절대 X

[문체]
- 합쇼체 기본 + 구어체 자연스럽게 ("~거든요", "~더라고요", "~잖아요")
- 짧은 문장 위주, 가끔 긴 문장으로 호흡 변화
- 현재형/과거형 혼용으로 생생한 장면감
- 디테일이 살아 있는 묘사 (시간, 숫자, 표정, 사물, 장소)

[분량] 600~800자

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🚫 절대 금지
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
❌ 마크다운 문법 일체 (# ## ### **굵게** > 인용 - 글머리) — 제목/라벨 출력 금지, 본문 텍스트만
❌ "프롤로그", "Prologue", "Prologue.", "들어가며" 같은 제목/라벨 출력 금지 (이미 본문 위에 'Prologue' 표시됨)
❌ 위 저자 정보를 그대로 복사 붙여넣기
❌ 자청, 역행자, 자의식 해체, 유전자 역행, 원시인, 추월차선 (특정 작가 고유 표현)
❌ 교과서 표현: "여정", "발걸음", "함께 하시길 바랍니다", "진정한", "이 책의 여정"
❌ AI 어휘: "~의 중요성", "다양한", "효과적인", "~를 통해", "이 책을 통해"
❌ 과장: "놀라운", "혁신적인", "충격적인", "어마어마한", "기적의"
❌ 뻔한 말: "포기하지 마세요", "꾸준히 하세요", "당신도 할 수 있습니다", "함께 가요"
❌ 직접 호칭: 여러분, 당신, 독자님
❌ 위에서 내려다보는 어조 (당신은 이래서 안 됩니다 X)

본문 텍스트만 출력하세요. 어떤 마크다운 헤더(#)나 '프롤로그' 라벨도 출력하지 마세요. 첫 문장부터 바로 본문이 시작되어야 합니다."""
            gen = ask_ai(prologue_prompt, 0.7, ensure_quality=True)
            if gen:
                prologue_text = gen
        if not prologue_text:
            prologue_text = ("수요일 밤 열한 시였습니다. 또 휴대폰을 들었습니다. 피드 속 친구는 벌써 두 번째 "
                             "집을 샀다고 했습니다. 저는 통장을 열어 보지도 못했습니다.\n\n"
                             "그때는 몰랐습니다. 문제는 의지가 아니라 구조였다는 걸요.\n\n"
                             "어느 날 한 줄의 계산을 마주하고서야 알았습니다. 같은 돈을 벌어도 누구는 쌓고 "
                             "누구는 흘려보내는 이유가 따로 있다는 것을요.\n\n"
                             "이 책은 그 단순한 차이를 처음부터 끝까지 풀어냅니다.")
        prologue_text = re.sub(r'^\s*#+\s*(프롤로그|Prologue|들어가며|머리말)\s*\.?\s*$', '', prologue_text, flags=re.MULTILINE | re.IGNORECASE)
        prologue_text = prologue_text.lstrip('\n').strip()
        paras = [x for x in prologue_text.split('\n\n') if x.strip()]
        for i, t in enumerate(paras):
            if i == 0 and len(t.strip()) > 1:
                drop_cap(t.strip()[0], t.strip()[1:])
            else:
                body_paragraph(t.strip(), first=(i == 0))
        doc.add_page_break()

        # ════════════════════════ 목차 ════════════════════════
        vspace(44)
        ct = doc.add_paragraph(); ct.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(ct.add_run(track_text("CONTENTS", 4)), 17, serif=True, bold=True, color=_INK)
        ct.paragraph_format.space_after = Pt(5)
        cts = doc.add_paragraph(); cts.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(cts.add_run("차 례"), 10, serif=False, color=_FAINT, track=3)
        cts.paragraph_format.space_after = Pt(12)
        hairline(width_cm=1.6, color=_GOLD, space_after=34, size=10)

        # 목차 전용 색: 종이책 목차처럼 먹색 박스 + 금박 라인
        TOC_PURPLE = '211C17'
        TOC_PURPLE_RGB = (151, 72, 48)
        TOC_TEAL = _GOLD

        # 분야 라벨(있으면) 추출용
        field_label = ''
        if interview_data:
            field_label = (interview_data.get('field', '') or '').strip()

        # 챕터별 시작 페이지 추정 (앞 구조 페이지 + 누적)
        import math as _math
        _CPP = 330
        page_cursor = 9  # 표지·판권·프롤로그·목차 등 앞부분 근사
        chapter_start_pages = {}
        sub_pages_map = {}
        for cidx, chap in enumerate(outline):
            chapter_start_pages[cidx] = page_cursor
            page_cursor += 1  # 챕터 오프너
            if chap in chapters_data:
                for sname in chapters_data[chap].get('subtopics', []):
                    sub_pages_map[(cidx, sname)] = page_cursor
                    c = chapters_data[chap].get('subtopic_data', {}).get(sname, {}).get('content', '')
                    chars = len((c or '').replace(' ', '').replace('\n', ''))
                    page_cursor += max(1, _math.ceil(chars / _CPP))

        def _toc_part_chip(idx, clean):
            """PART 박스 + 제목을 한 줄 표로 배치."""
            t = doc.add_table(rows=1, cols=2)
            t.alignment = WD_TABLE_ALIGNMENT.LEFT
            t.autofit = False
            cbox, ctitle = t.rows[0].cells
            cbox.width = Cm(1.7); ctitle.width = Cm(9.1)
            # PART 박스
            _cell_shade(cbox, TOC_PURPLE)
            _no_cell_borders(cbox)
            bp1 = cbox.paragraphs[0]; bp1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            bp1.paragraph_format.space_before = Pt(5); bp1.paragraph_format.space_after = Pt(0); bp1.paragraph_format.line_spacing = 1.0
            set_font(bp1.add_run(f"PART {idx+1}"), 8, serif=False, bold=True, color=(246, 235, 214), track=0.3)
            if field_label:
                bp2 = cbox.add_paragraph(); bp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                bp2.paragraph_format.space_before = Pt(1); bp2.paragraph_format.space_after = Pt(5); bp2.paragraph_format.line_spacing = 1.0
                set_font(bp2.add_run(field_label[:6]), 10, serif=False, bold=True, color=(214, 180, 101), track=1)
            else:
                bp1.paragraph_format.space_after = Pt(5)
            va = OxmlElement('w:vAlign'); va.set(qn('w:val'), 'center'); cbox._tc.get_or_add_tcPr().append(va)
            # 제목 셀
            _no_cell_borders(ctitle)
            tp = ctitle.paragraphs[0]
            tp.paragraph_format.left_indent = Cm(0.35)
            tp.paragraph_format.space_before = Pt(2); tp.paragraph_format.space_after = Pt(0); tp.paragraph_format.line_spacing = 1.2
            add_hyperlink(tp, clean, f"chapter_{idx+1}", size=14.5, bold=True, color=_INK, serif=True)
            va2 = OxmlElement('w:vAlign'); va2.set(qn('w:val'), 'center'); ctitle._tc.get_or_add_tcPr().append(va2)

        for idx, chapter in enumerate(outline):
            clean = chapter
            for pre in [f"PART {idx+1}.", f"PART{idx+1}.", f"PART {idx+1} ", f"{idx+1}.", f"{idx+1})"]:
                clean = clean.replace(pre, "").strip()
            if idx > 0:
                vspace(20)
            # 청록 점선 (상단)
            hairline(align=WD_ALIGN_PARAGRAPH.LEFT, color=TOC_TEAL, space_before=0, space_after=8, size=6)
            # 보라 PART 박스 + 제목
            _toc_part_chip(idx, clean)
            # 청록 점선 (하단)
            hairline(align=WD_ALIGN_PARAGRAPH.LEFT, color=TOC_TEAL, space_before=8, space_after=14, size=6)
            # 소제목 + 점선 리더 + 페이지 번호
            if chapter in chapters_data:
                for sub in chapters_data[chapter].get('subtopics', []):
                    si = chapters_data[chapter]['subtopics'].index(sub)
                    pg = sub_pages_map.get((idx, sub), '')
                    sr = doc.add_paragraph()
                    sr.paragraph_format.left_indent = Cm(0.3); sr.paragraph_format.space_after = Pt(9)
                    sr.paragraph_format.line_spacing = 1.3
                    sr.paragraph_format.tab_stops.add_tab_stop(
                        Cm(10.6), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
                    add_hyperlink(sr, sub, f"subtopic_{idx+1}_{si+1}", size=10.5, bold=False, color=(58, 58, 64), serif=True)
                    # 탭 + 페이지번호
                    tabrun = sr.add_run(f"\t{pg:03d}" if isinstance(pg, int) else "")
                    set_font(tabrun, 10, serif=False, color=TOC_PURPLE_RGB)
        doc.add_page_break()

        # ════════════════════════ 본문 (장마다 섹션) ════════════════════════
        for idx, chapter in enumerate(outline):
            clean = chapter
            for pre in [f"PART {idx+1}.", f"PART{idx+1}.", f"PART {idx+1} ", f"{idx+1}.", f"{idx+1})"]:
                clean = clean.replace(pre, "").strip()
            sec = doc.add_section(WD_SECTION.NEW_PAGE)
            _apply_section(sec)
            sec.different_first_page_header_footer = True
            running_header(sec, clean)
            page_number_footer(sec)

            vspace(72)
            lab = doc.add_paragraph(); lab.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(lab.add_run(track_text("CHAPTER", 2)), 9.5, serif=False, color=_FAINT, track=1)
            lab.paragraph_format.space_after = Pt(2)
            num = doc.add_paragraph(); num.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(num.add_run(f"{idx+1:02d}"), 82, serif=True, bold=True, color=_GOLD)
            num.paragraph_format.space_after = Pt(4)
            hairline(width_cm=1.5, color='211C17', space_after=10, size=8)
            hairline(width_cm=0.55, color=_GOLD, space_before=0, space_after=22, size=9)
            cn = doc.add_paragraph(); cn.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_bookmark(cn, f"chapter_{idx+1}")
            set_font(cn.add_run(clean), 22, bold=True, color=_INK, serif=True)
            cn.paragraph_format.line_spacing = 1.24; cn.paragraph_format.space_after = Pt(52)

            if chapter not in chapters_data:
                continue
            subs = chapters_data[chapter].get('subtopics', [])
            for si, sub in enumerate(subs):
                content = chapters_data[chapter].get('subtopic_data', {}).get(sub, {}).get('content', '')
                if not content:
                    continue
                if si > 0:
                    doc.add_page_break()
                # 소제목 번호 칩 (골드 배경 박스)
                chip = doc.add_paragraph()
                chip.paragraph_format.space_before = Pt(6); chip.paragraph_format.space_after = Pt(5)
                _shade_para(chip, '211C17')
                cr = chip.add_run(f"  {idx+1}.{si+1}  ")
                set_font(cr, 10.5, bold=True, serif=False, color=(214, 180, 101), track=1)
                # 소제목 제목 (크고 굵게)
                ht = doc.add_paragraph()
                add_bookmark(ht, f"subtopic_{idx+1}_{si+1}")
                set_font(ht.add_run(sub), 17, bold=True, color=_INK, serif=True)
                ht.paragraph_format.space_after = Pt(7); ht.paragraph_format.line_spacing = 1.22
                # 진한 더블 밑줄(챕터 오프너와 확실히 구분)
                hairline(align=WD_ALIGN_PARAGRAPH.LEFT, color='211C17', space_before=0, space_after=2, size=10)
                hairline(align=WD_ALIGN_PARAGRAPH.LEFT, color=_GOLD, space_before=0, space_after=20, size=4)

                cleaned = clean_content(content, subtopic=sub)
                first_para_done = False
                for btype, bcontent in process_content_with_tables(cleaned):
                    if btype == 'table':
                        td = parse_table_data(bcontent)
                        if td and len(td) >= 2:
                            sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(8)
                            add_premium_table(td)
                            first_para_done = True
                    else:
                        for pt in [p for p in bcontent.split('\n\n') if p.strip()]:
                            # 챕터 첫 소제목의 첫 문단만 드롭캡, 이후엔 일반
                            if (not first_para_done) and si == 0 and len(pt.strip()) > 1:
                                drop_cap(pt.strip()[0], pt.strip()[1:])
                            else:
                                body_paragraph(pt.strip(), first=(not first_para_done))
                            first_para_done = True
                if si < len(subs) - 1:
                    ornament()

            # 챕터 끝 핵심 정리 박스
            ch_summary = chapters_data[chapter].get('key_points', [])
            if ch_summary:
                key_summary_box(ch_summary, idx + 1)

        # ════════════════════════ 에필로그 ════════════════════════
        ep_sec = doc.add_section(WD_SECTION.NEW_PAGE)
        _apply_section(ep_sec)
        ep_sec.different_first_page_header_footer = True
        page_number_footer(ep_sec)
        vspace(36)
        el = doc.add_paragraph(); el.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(el.add_run(track_text("EPILOGUE", 2)), 10, serif=False, color=_FAINT)
        el.paragraph_format.space_after = Pt(8)
        hairline(width_cm=1.4, color=_GOLD, space_after=30, size=8)
        epilogue_text = None
        if interview_data:
            epilogue_prompt = f"""당신은 한국 자기계발 베스트셀러 작가입니다. 마지막 페이지를 덮은 독자가 한 번 더 처음으로 돌아가게 만드는 에필로그를 작성하세요.

[저자 정보 - 참고용, 그대로 복사하지 말 것]
- 분야: {interview_data.get('field', '')}
- 경력 기간: {interview_data.get('experience_years', '')}
- 책 주제: {interview_data.get('topic', '')}
- 타겟 독자: {interview_data.get('target_reader', '')}
- 독자에게 전하고 싶은 말: {interview_data.get('final_message', '')}
- 작가 경력/경험: {interview_data.get('author_career', '')}
- 어려움/실패 경험: {interview_data.get('struggle_story', '')}
- 극복 스토리: {interview_data.get('breakthrough', '')}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

## 에필로그 작성 원칙 (스토리텔링)

### 1. 나의 스토리로 시작 (3-4문장)
- 작가 경력/경험을 자연스럽게 녹여서
- "저는 ~했습니다" 형식으로 간결하게
- 구체적 숫자나 사실 포함

### 2. 왜 이 책을 썼는지 (2-3문장)
- 내가 겪은 어려움 + 극복 과정 힌트
- 독자를 위해 책을 쓴 진심

### 3. 독자에게 한마디 (2-3문장)
- 지금 당장 할 수 있는 구체적 행동 하나
- 진심 어린 마무리 (근데 뻔하지 않게)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[문체]
- 짧은 문장, 짧은 문단 (2-4문장)
- 구어체 + 합쇼체 ("~거든요", "~잖아요" OK)

[분량] 400-500자

[금지 - 절대 쓰지 말 것]
- 저자 정보를 그대로 복사 붙여넣기
- 자청, 역행자, 자의식 해체, 유전자 역행, 원시인, 추월차선 (특정 작가 고유 표현)
- 교과서 표현: "여정", "발걸음", "함께 하시길 바랍니다", "진정한"
- AI 표현: "~의 중요성", "다양한", "효과적인", "~를 통해"
- 과장: "놀라운", "혁신적인", "충격적인"
- 뻔한 말: "포기하지 마세요", "꾸준히 하세요", "화이팅"
- 직접 호칭: 여러분, 당신, 독자님
- 마크다운 문법

에필로그만 출력하세요."""
            gen = ask_ai(epilogue_prompt, 0.7, ensure_quality=True)
            if gen:
                epilogue_text = gen
        if not epilogue_text:
            epilogue_text = ("여기까지 읽어주셔서 고맙습니다.\n\n완벽하지 않아도 괜찮습니다. 오늘 할 수 "
                             "있는 한 가지만 시작해 보세요. 작은 시작이 가장 멀리 갑니다.")
        epilogue_text = re.sub(r'^\s*#+\s*(에필로그|Epilogue|마치며|맺음말)\s*\.?\s*$', '', epilogue_text, flags=re.MULTILINE | re.IGNORECASE)
        epilogue_text = epilogue_text.lstrip('\n').strip()
        for i, t in enumerate([x for x in epilogue_text.split('\n\n') if x.strip()]):
            body_paragraph(t.strip(), first=(i == 0))
        vspace(30)
        sg = doc.add_paragraph(); sg.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_font(sg.add_run(author or "저자"), 11, italic=True, color=_SOFT, serif=True)

        # ════════════════════════ 저자 소개 ════════════════════════
        doc.add_page_break()
        vspace(96)
        al = doc.add_paragraph(); al.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(al.add_run(track_text("ABOUT THE AUTHOR", 1)), 9, serif=False, color=_FAINT)
        al.paragraph_format.space_after = Pt(10)
        hairline(width_cm=1.4, color=_GOLD, space_after=22, size=8)
        an = doc.add_paragraph(); an.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(an.add_run(author or "저자"), 16, bold=True, color=_INK, serif=True)
        an.paragraph_format.space_after = Pt(18)
        if interview_data:
            field = interview_data.get('field', ''); exp = interview_data.get('experience_years', '')
            career = interview_data.get('author_career', ''); method = interview_data.get('core_method', '')
            if career:
                author_bio = f"{field} 분야에서 {exp} 활동해온 실전가.\n\n{career}\n\n{method[:100] if method else ''}"
            else:
                author_bio = f"{field} 분야에서 {exp} 활동해온 실전가.\n\n{method}"
        else:
            author_bio = "실전에서 직접 부딪히며 쌓은 노하우를 독자와 나누고자 이 책을 썼다."
        for t in [x for x in author_bio.split('\n\n') if x.strip()]:
            bp = doc.add_paragraph(); bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(bp.add_run(t.strip()), 10, color=_SOFT, serif=True)
            bp.paragraph_format.line_spacing = 1.7; bp.paragraph_format.space_after = Pt(10)
        vspace(40)
        em = doc.add_paragraph(); em.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(em.add_run("✦"), 11, color=_FAINT, serif=False)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue(), None

    except Exception as e:
        return None, f"문서 생성 오류: {str(e)}"
# ==========================================
# 기법(메서드) 이름 — 영문 통일 + 책마다 고유
# ==========================================
def load_used_method_names():
    """이전 전자책들에서 쓴 기법 이름 목록 (겹침 방지용)"""
    names = load_config().get('used_method_names', [])
    return names if isinstance(names, list) else []


def remember_method_name(name):
    if not name:
        return
    names = load_used_method_names()
    if name not in names:
        names.append(name)
        save_config({'used_method_names': names[-200:]})  # 최근 200개만 보관


def _fallback_method_name(topic):
    import random
    # 발음 가능한 2~3자 영문 약자 + 자연스러운 한글 접미사 (예: "ERP 공식")
    acronym = "".join(random.sample("ABCDEFGHJKLMNPRSTVWXYZ", k=random.choice([2, 3])))
    suffix = random.choice(["공식", "법칙", "시스템", "전략", "루틴", "구조", "사이클"])
    return f"{acronym} {suffix}"


def get_or_create_method_name(topic, interview_data=None, force_new=False):
    """이 책의 기법 이름을 만들거나 가져온다.

    형식: '영문 약자(2~4자) + 한글 접미사'  예) "ERP 공식", "DPS 법칙"
    - force_new=True: 새 전자책 → 이전 책들과 겹치지 않는 새 이름
    - force_new=False: 같은 책 안의 재생성 → 이미 정해진 이름 재사용
    """
    if not force_new and st.session_state.get('method_name'):
        return st.session_state['method_name']

    used = load_used_method_names()
    used_str = ", ".join(used[-60:]) if used else "(아직 없음)"
    seed = uuid.uuid4().hex[:6]
    core = (interview_data or {}).get('core_method', '') if interview_data else ''

    prompt = f"""이 전자책의 '기법 이름' 딱 하나를 짓는다.

주제: {topic}
핵심 방법: {core}

형식 (반드시 지킬 것):
- "영문 대문자 약자(2~4자) + 한글 접미사" 한 덩어리.  예) "ERP 공식", "DPS 법칙", "ARC 전략", "PMR 루틴"
- 약자는 발음 가능하고, 각 글자가 주제와 관련된 영어 단어의 첫 글자여야 한다(의미 있는 약자).
  예) ERP = Earn-Reinvest-Profit, DPS = Discover-Plug-Scale
- 한글 접미사는 다음 중에서만: 공식 / 법칙 / 시스템 / 전략 / 루틴 / 구조 / 사이클
- 약자만 영어이고, 영어 '단어'를 길게 늘어놓지 말 것.

❌ 절대 금지 (어색한 예):
- "eBook Revenue Loop", "Cashflow Ladder" 같은 영어 다단어 구절 (부자연스러움)
- 한글 약자, 의미 없는 자음 나열, 발음 불가능한 약자
- 아래 이미 쓴 이름과 약자·접미사가 겹치는 것:
  {used_str}

다양성 시드 {seed} 참고해 매번 다르게.

JSON만 출력:
{{"acronym": "영문 대문자 2~4자", "expansion": "약자 풀이(영문, 하이픈 연결)", "suffix": "한글 접미사", "method_name": "약자 + 공백 + 접미사"}}"""

    name = None
    expansion = ""
    try:
        res = ask_ai(prompt, 0.9)
        parsed = parse_json(res) if res else None
        if parsed:
            ac = (parsed.get('acronym') or '').strip().upper()
            sf = (parsed.get('suffix') or '').strip()
            mn = (parsed.get('method_name') or '').strip()
            expansion = (parsed.get('expansion') or '').strip()
            # method_name 우선, 없으면 약자+접미사 조합
            name = mn if mn else (f"{ac} {sf}" if ac and sf else '')
    except Exception:
        name = None

    # 형식 검증: 영문 대문자 2~4자 + 공백 + 허용된 한글 접미사
    suffix_re = "(공식|법칙|시스템|전략|루틴|구조|사이클)"
    pattern = r"^[A-Z]{2,4}\s" + suffix_re + r"$"

    def _valid(n):
        return bool(n) and re.fullmatch(pattern, (n or "").strip()) is not None and n.strip() not in used

    name = (name or "").strip()
    if not _valid(name):
        name = _fallback_method_name(topic)
        expansion = ""  # 폴백 이름은 의미 있는 약자 풀이가 없음
        tries = 0
        while name in used and tries < 15:
            name = _fallback_method_name(topic)
            tries += 1

    st.session_state['method_name'] = name
    st.session_state['method_expansion'] = expansion
    remember_method_name(name)
    return name


def method_lock_rule(name):
    """모든 생성 단계에서 동일하게 끼워 넣는 '기법 이름 고정' 규칙 블록"""
    return f"""━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔒 기법 이름 고정 (이 프롬프트의 다른 어떤 작명/금지 규칙보다 우선)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
이 책의 기법 이름은 정확히 "{name}" 으로 이미 확정되어 있다. (형식: 영문 약자 + 한글 접미사)
- 모든 곳에서 "{name}" 을 글자 그대로, 띄어쓰기까지 똑같이 사용한다.
- 약자를 영어 단어로 풀어쓰거나, 다른 약자/이름으로 바꾸거나, 접미사를 바꾸는 것을 절대 금지한다.
- ○○○ / [컨셉명] / [시스템명] 자리에는 전부 "{name}" 을 넣는다.
- "{name}" 에 또 다른 한글 접미사(시스템/공식 등)를 덧붙이지 않는다. 위 이름 그대로만 쓴다.
"""


def generate_outline_only(interview_data, progress_placeholder):
    """인터뷰 데이터를 기반으로 목차까지만 생성 (본문 제외)"""
    try:
        topic = interview_data.get('topic', '')
        if not topic:
            return False

        # 1. 타겟 자동 설정
        progress_placeholder.info("🎯 1/4 타겟 독자 분석 중...")
        target = f"{interview_data.get('target_reader', '')} - {interview_data.get('target_problem', '')}"
        st.session_state['target_persona'] = target

        # 새 전자책 → 이전 책들과 겹치지 않는 영문 기법 이름을 새로 확정
        method_name = get_or_create_method_name(topic, interview_data, force_new=True)
        method_expansion = st.session_state.get('method_expansion', '')

        # 2. 책 고유 컨셉 생성 (가장 중요!)
        progress_placeholder.info("💡 2/4 책 고유 컨셉 설계 중...")
        concept_prompt = f"""당신은 크몽/클래스101 베스트셀러 전자책 기획자입니다.
이 책의 기법 이름은 이미 "{method_name}" (영문 약자 + 한글 접미사 형식)으로 확정되어 있습니다.
당신의 일은 새 이름을 짓는 게 아니라, 이 기법을 중심으로 책의 관점과 메시지를 설계하는 것입니다.

{method_lock_rule(method_name)}

이 기법의 약자 풀이(참고): {method_expansion if method_expansion else "(풀이 없음)"}
약자가 의미 있는 머리글자라면, 핵심 관점·메시지에 그 단계적 흐름이 자연스럽게 묻어나게 하라. 단 "Y:~" 식 사전 나열은 금지.

[저자 정보]
주제: {topic}
핵심 방법: {interview_data.get('core_method', '')}
저자만의 차별점: {interview_data.get('unique_point', '')}
타겟의 고민: {interview_data.get('target_problem', '')}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
✅ 만들어야 할 것
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

1. 핵심 관점 — 이 주제를 "{method_name}"(이)라는 기법으로 어떻게 새롭게 보는가? 남들과 다른 접근법.
2. 핵심 메시지 — "{method_name}(으)로 ~할 수 있다" 형식의 한 문장.

❌ 절대 금지:
- 기법 이름을 새로 짓거나 한글로 바꾸는 것 (이미 "{method_name}"으로 고정)
- "제국을 건설/왕좌/전설의/세계 최초" 같은 과장
- 황금·보물·비밀·마법·연금술 같은 유치한 단어

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📝 출력 형식
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[고유 시스템/공식 이름]
{method_name}

[핵심 관점]
(2~3문장, 자연스럽게)

[핵심 메시지]
(한 문장, "{method_name}(으)로 ~하는 방법")

[목차에서 반복할 키워드]
{method_name}"""

        book_concept = ask_ai(concept_prompt, 0.8, ensure_quality=True)
        st.session_state['book_concept'] = book_concept

        # 3. 제목 생성
        progress_placeholder.info("📝 3/4 제목 생성 중...")
        title_prompt = f"""당신은 크몽/클래스101 베스트셀러 전자책 제목을 만드는 전문가입니다.
결제 버튼을 누르게 만드는 제목을 써주세요.

[이 책의 컨셉]
{book_concept}

이 책의 기법 이름은 "{method_name}"으로 확정됨. 제목에 기법 이름을 넣을 경우 반드시 이 이름 그대로 사용(약자 풀어쓰기·재작명 금지).

[주제]
{topic}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔥 실제 잘 팔리는 전자책 제목 분석
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[컨셉 중심형 - 짧은 신조어 + 부제]
• 단단한 돈 - 잃지 않는 사람의 7가지 원칙
• 돈의 속성 - 최소한 이것만은 알아야 할
• 1억 모으는 통장 - 30대 직장인의 5단계 공식

[신사임당/클래스101 스타일 - 결과 중심]
• 퇴사 후 월 1000만원 버는 글쓰기
• 블로그로 월 300 만드는 현실적인 방법
• 투잡러의 시간관리 비법

[크몽 베스트셀러 - 구체적 약속]
• 30일 만에 첫 수익 내는 스마트스토어
• 3개월 안에 월 100 만드는 전자책 공식
• 회사 다니면서 월 200 추가 수입 만들기

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
✅ 좋은 제목의 공식
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[공식 1] 컨셉형 (2~4단어)
저자만의 프레임워크/용어가 들어간 제목
예: 짧고 강렬한 신조어/약어 (3~5자), 동사의 명사화, 은유적 압축어

[공식 2] 결과형 (구체적 숫자 포함)
기간 + 결과가 명확한 제목
예: "3개월 만에 월 300", "100일 글쓰기"

[공식 3] 타겟형 (누구를 위한)
특정 대상의 고민을 건드리는 제목
예: "퇴사 준비생의 월급 독립기", "직장인의 두 번째 월급"

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🚫 절대 금지
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

❌ 유치한 비유:
나침반, 지도, 열쇠, 보물, 황금, 마법, 연금술

❌ AI스러운 제목:
"~의 이해", "~가이드", "~완벽 정복"
"효과적인 ~", "성공적인 ~"

❌ 너무 추상적:
의미를 알 수 없는 신조어
무슨 내용인지 전혀 감이 안 오는 제목

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📝 출력
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

제목: 컨셉이 드러나면서도 무슨 책인지 알 수 있게
부제: 구체적인 결과/약속을 담아 15~25자

JSON만 출력:
{{
    "title": "제목 (컨셉+내용이 드러나게, 3~8단어)",
    "subtitle": "부제 (구체적 결과/약속, 15~25자)"
}}"""

        title_result = ask_ai(title_prompt, 0.4)
        title_data = parse_json(title_result)
        if title_data:
            st.session_state['book_title'] = title_data.get('title', topic)
            st.session_state['subtitle'] = title_data.get('subtitle', '')

        # 4. 목차 생성 (책 컨셉 기반)
        progress_placeholder.info("📋 4/4 목차 설계 중...")
        outline_prompt = f"""당신은 한국 자기계발 분야 톱 0.1% 기획자입니다. 서점에서 단 5초간 목차만 본 사람이 책을 손에서 못 놓게 만드는 5장짜리 목차를 씁니다.

목차의 단 하나의 목적: 독자가 "이 책을 안 읽으면 평생 손해"라고 느끼게 만드는 것.
정보 전달은 본문이 한다. 목차는 100% 구매심리만 다룬다.

[이 책의 시그니처 컨셉/시스템]
{book_concept}

{method_lock_rule(method_name)}

[주제]: {topic}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🛒 구매 결정 5초 룰 (모든 규칙 중 1순위)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

독자는 목차를 5초만 본다. 그 5초 안에 다음 3가지가 동시에 작동해야 결제한다.

[1] 정체성 변화 발견 — "이걸 읽으면 나는 OO한 사람이 된다"
   → 5개 챕터 제목을 이어 읽으면 한 사람의 결정적 변화가 보여야 한다.
   ✅ "단단해진 멘탈은 인생을 통째로 바꾼다" (변화 서사 O)
   ❌ "멘탈 관리의 다양한 기법" (정체성 변화 X — 즉시 폐기)

[2] 손실회피 작동 — "이걸 모르면 평생 OO한다"
   → 통념 박살(인지부조화) 챕터/소제목이 최소 3개 들어가야 한다.
   ✅ "의지로 버틴 사람일수록 더 크게 무너진다"
   ❌ "멘탈 관리의 중요성" (잃을 게 안 보임)

[3] 구체성 — 추상 명사 1개당 구체적 숫자/장면 1개
   → "많은 사람" 금지, "월급 280만원짜리 7년차 회사원" 가능
   → 시간(90초, 47일), 금액(34만원, 1억), 비율(99%, 8할) 적극 사용

25줄(챕터 5 + 소제목 20) 중 한 줄이라도 "그냥 정보"가 섞이면 그 목차는 평이해진다.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
💎 기법 이름 사용 (위 고정 규칙을 그대로 따른다)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

기법 이름은 위에서 "{method_name}"(영문 약자+한글 접미사)으로 이미 확정되었다.
- 새로 작명하지 말고, 약자를 풀어쓰지 말고, 글자 그대로 쓴다.
- 아래 예시의 ○○○ 자리에는 전부 "{method_name}" 을 넣는다.
- "{method_name}" 에 또 다른 접미사를 덧붙이지 않는다.

[약자의 의미를 목차에 '간접적으로' 녹여라]
이 기법의 약자 풀이: {method_expansion if method_expansion else "(풀이 없음 — 약자가 무엇의 머리글자인지 책 흐름으로 자연스럽게 암시)"}
- 약자가 의미 있는 머리글자(예: ERP = Earn-Reinvest-Profit)라면, 그 각 단계가 PART 1~5의 흐름이나 소제목에 '자연스럽게' 드러나도록 목차를 설계하라. 독자가 목차만 읽어도 "아, 이 책이 이런 단계를 다루는구나"를 감 잡게.
- 단, "Y: ~, M: ~" 처럼 글자별로 사전식 나열은 절대 금지. 풀이를 직접 받아쓰지 말 것.
- 어디까지나 자연스러운 호기심 자극 제목 안에 의미가 스며들게 한다. 예) 약자 첫 단계가 'Earn(벌기)'이면 PART 1을 버는 단계의 통념 박살로 구성하는 식.
- 풀이가 없으면 억지로 짜맞추지 말고, 평소처럼 호기심 중심으로 쓰되 책 전체가 하나의 방법론을 단계적으로 다룬다는 느낌만 유지하라.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🧠 마케팅 뇌과학 8대 트리거 (목차 전체에 골고루 박을 것)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

1. **손해회피 + 충격 통계** — "1년 안에 99%가 다시 무너진다", "1억치 강의 들어도 망한다"
2. **인지부조화/통념 박살** — "의지로 버틴 사람일수록 더 크게 무너진다", "노력할수록 가난해진다"
3. **권위 어휘 (과학/임상)** — 뇌, 신경회로, N주 후, 임상, 데이터, 알고리즘 (절대 비유로 남용 금지, 사실 진술로만)
4. **임박감 + 절대성** — "이 90초를 놓치면 며칠 걸린다", "다시는 ~하지 않는다"
5. **정체성 전환 약속** — "회복한 뇌는 다시 무너지지 않는다", "단단해진 사람은 ~한다"
6. **인그룹 사회증명** — "○○를 익힌 사람들의 5년 뒤", "상위 1%만 도달하는"
7. **이중/삼중 보상** — "통장과 인간관계가 함께 변한다", "돈도 사람도 따라온다"
8. **호기심 갭** — "두 달 안에 가장 먼저 끊은 한 가지", "정확히 어디부터 멈추는가"

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🧲 호기심 갭(Information Gap) 강화 — 결제를 부르는 가장 강력한 무기
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

조지 로웬스타인의 정보격차 이론: 사람은 자기가 모르는 것이 '뭔지'는 알지만 '내용'은 모를 때 가장 강하게 끌린다. 25줄 중 최소 10줄에 이 갭을 박아야 목차만 보고 결제한다.

[호기심 갭 6대 공식 - 결과만 보이고 방법/이유는 본문으로 미루기]

1. **숫자 + 미공개 결과** — '정확히/딱'으로 시작해 결과만 보여주고 내용은 숨기기
   ✅ "정확히 47일째에 통장이 처음 뒤집힌 그 순간"
   ✅ "27만원짜리 첫 정산서가 알려준 단 한 가지"
   ❌ "47일 만에 돈을 버는 방법" (방법을 다 보여줘버림 → 결제 안 함)

2. **이미 벌어진 사건 + 원인 숨김** — '왜?'를 유발하는 결과만
   ✅ "3년 차 베테랑이 신입에게 6개월 만에 따라잡힌 단 하나의 이유"
   ✅ "월 1,000을 찍은 사람들이 가장 먼저 끊은 습관 한 가지"

3. **묘하게 구체적인 행동/대상 + 이유 숨김** — 디테일이 호기심을 폭발시킨다
   ✅ "성공한 부업러가 매일 밤 11시에 반드시 끄는 것"
   ✅ "1년 만에 1억 모은 사람들이 절대 안 쓰는 5단어"
   ✅ "월 500 넘긴 사람들 카톡 프로필에서 사라진 한 단어"

4. **반대 결과 미스터리** — 통념과 정반대 결과만 던지고 메커니즘은 본문
   ✅ "더 열심히 할수록 더 가난해진 7년의 비밀"
   ✅ "잠을 늘렸더니 매출이 2배가 된 이상한 메커니즘"

5. **'딱 하나' 절대성** — 수많은 변수 중 단 하나만 보여주기
   ✅ "월 100 / 월 500을 가르는 단 한 줄의 차이"
   ✅ "결국 모든 게 무너지는 사람들의 공통점 단 하나"

6. **시간 압축 미스터리** — 짧은 시간에 큰 일이 일어났는데 그 사이를 숨기기
   ✅ "퇴근 후 90분이 1년 뒤 인생을 갈라놓는다"
   ✅ "주말 4시간이 5년치 월급을 바꾼 그 과정"

[호기심 갭 만들 때 절대 어기지 말 것]
• 답을 같은 줄에 다 보여주지 마라. "왜 ~한가" "어떻게 ~하는가"로 끝나면 본문을 사야 알 수 있게.
• "방법" "비법" "노하우" 같은 말로 끝내면 갭이 닫힘 → 결제 안 함.
• 결과/사건/디테일은 보이고, 원리/메커니즘/순서는 숨겨라.
• 한 줄 안에 "장면 + 의문"이 같이 있어야 호기심이 작동한다.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🎯 단 하나의 미션
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

5개 챕터 제목만 빠르게 이어 읽었을 때 한 사람의 변화 이야기가 보이고, 한 줄 한 줄이 다음 챕터를 못 참게 만들어야 한다.

[컨셉명 등장 규칙 - 3번, 자연스럽게 분산]
컨셉명은 챕터 제목에는 등장 X. 소제목에만 정확히 3번 자연스럽게 박는다.

⚠️ 컨셉명 글자 수/단어 수를 단정하는 표현 절대 금지 (가장 자주 어기는 실수!)
   ❌ "○○○ 세 글자에서 시작한다" — 컨셉명이 3글자 아니면 거짓말이 됨
   ❌ "○○○ 네 글자만 기억하라" — 글자 수 단정 금지
   ❌ "단 한 단어가 모든 걸 바꾼다" — 단어 수 단정 금지
   ✅ 컨셉명이 'DPS'(3자)든 '단단한 매출 구조'(8자)든 '복리 자산 공식'(7자)이든 모두 자연스럽게 작동하는 문장만 사용

[PART 1 마지막 소제목 - 도입] (아래 풀에서 1개 선택, 책에 가장 어울리는 것)
   • "결국 모든 답은 '○○○' 안에 있었다"
   • "이 책의 모든 페이지는 '○○○' 하나를 향해 간다"
   • "지금부터 '○○○' 단 하나만 기억하면 된다"
   • "여기서부터 진짜 이야기, '○○○'가 시작된다"
   • "마지막에 도달하는 곳은 결국 '○○○'다"
   • "'○○○'를 만나기 전과 후는 완전히 다른 게임이다"
   • "이 모든 혼란을 한 줄로 정리하는 '○○○'"

[PART 3 또는 PART 4 안 1곳 - 작동·전환점]
   • "○○○가 본격 작동하기 시작하는 4가지 신호"
   • "○○○를 처음 적용한 사람들이 가장 먼저 느낀 변화"
   • "○○○가 통장에 처음 흔적을 남기는 순간"
   • "○○○ 한 달 차에 가장 먼저 무너지는 한 가지"

[PART 5 마지막 소제목 - 확장·사회증명]
   • "○○○를 익힌 사람들의 5년 뒤가 완전히 다른 이유"
   • "○○○로 자리잡은 사람들이 다시는 돌아가지 않는 이유"
   • "○○○ 이후, 1년 만에 가장 크게 달라지는 단 한 가지"

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📐 챕터 제목 형식 (가장 중요)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[형식]
- 13~18자의 자연스러운 한국어 한 문장
- 단어 한두 개짜리 라벨 금지, 라벨 분리(— ㅣ :) 금지
- 평서문 또는 단언형. "~이유"로 끝나는 설명체는 한 PART에만 사용
- 명사 엔딩과 동사 엔딩을 챕터별로 섞어라 (5개 모두 명사 엔딩 금지)
- 5개를 이으면 [좌절 → 통념 박살 → 첫 사건 → 안정화 → 도약]의 5막

[좋은 예 - 성공적인 멘탈 관리 비결]
PART 1. 1년 안에 99%가 다시 무너지는 결정적 이유
PART 2. 의지로 잡으려는 순간 뇌는 반대로 움직인다
PART 3. 회복의 8할이 결판나는 폭발 직후 90초
PART 4. 한 번 회복한 뇌는 다시는 무너지지 않는다
PART 5. 단단해진 멘탈은 인생을 통째로 바꾼다

[좋은 예 - 30대 직장인 N잡 월 500]
PART 1. 직장인 99%가 부업 30일을 못 버티는 이유
PART 2. 노력보다 자리가 먼저다
PART 3. 첫 30만원이 통장에 찍힌 그날
PART 4. 새벽 3시에도 매출이 들어온다
PART 5. 월 500 다음, 억대 수익으로 가는 길

[나쁜 예 - 즉시 폐기]
- "발굴", "폭로", "전환" 같은 한두 단어짜리
- "발굴 — 부업의 90%는 첫 단추에서 망한다" (라벨 + 대시)
- "DPS의 첫 관문, 노력 없이도 돈이 따라오는 자리를 찾는 법" (시스템명 라벨화 + 너무 김)
- "이제 무너지는 게 더 이상 사건이 아니다" (사건이 아니다 ← 말이 안 됨)
- "한 번 흔들려도 다음 날엔 흔적도 없다" (AI식 과장)
- "월수도 시스템의 첫 설계" (의미 불명 + 설계라는 설명체 어휘)
- "MDS 파이프라인" (파이프라인이라는 영어 외래어를 시스템 접미사로)
- "주가 -12% 떨어져도 매도 안 하는 뇌 회로가 박혔다" (뇌 회로가 박히다 = 어법 어색, 비유 남용)
- "신경회로가 새로 깔리고 있다는 신호" 류 (한 번까진 OK, 같은 비유 두 번 X)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
✍️ 소제목 톤: 한국 자기계발 베스트셀러 + 마케팅 뇌과학
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

각 PART당 소제목 4개. 모두 다른 패턴 + 명사·동사 엔딩 섞기 (4개 모두 같은 엔딩 금지).

[좋은 패턴 풀 - 매번 다르게]
1. 통계 충격형 — "1년 안에 99%가 다시 무너진다", "한 달 차에 90%가 다시 무너지는 정확한 이유"
2. 인지부조화형 — "의지로 버틴 사람일수록 더 크게 무너진다"
3. 뇌과학 권위형 — "감정이 폭발할 때 뇌는 정확히 어디부터 멈추는가", "신경회로가 새로 깔리고 있다는 4가지 신호"
4. 임박감/손해회피형 — "이 90초를 놓치면 다시 일어서는 데 며칠이 걸린다"
5. 정체성 전환형 — "한 번 회복한 뇌는 다시는 무너지지 않는다"
6. 호기심 갭형 — "단번에 회복한 사람들이 모두 무의식적으로 하는 행동"
7. 이중/삼중 보상형 — "회복 후 6개월 만에 통장과 인간관계가 함께 변한다"
8. 인그룹 사회증명형 — "○○를 익힌 사람들의 5년 뒤가 완전히 다른 이유"

[★ 가장 중요 — 목차만 보고 결제하게 만들어라 (설명문 금지)]
이 목차의 단 하나의 목적: 서점에서 5초 훑은 사람이 "이건 사야 해"라고 결제하게 만드는 것.
정보 전달은 본문이 한다. 목차는 100% 구매 욕구만 자극한다.

지금 가장 흔한 실패 = "설명문처럼 평이함". 아래 셋 중 하나라도 걸리면 그 줄은 죽은 줄이니 다시 써라:
  (1) 정보 전달형("~하는 법", "~의 중요성", "~란 무엇인가") — 교과서 목차
  (2) 답을 다 말해버림 — 본문을 살 이유가 사라짐
  (3) 자극이 없음 — 심장이 안 뛰면 결제 안 한다

[자청식 후킹 7가지 무기 — PART마다 최소 2개 이상 섞어라]
1. 통념 정면 박살: "열심히 할수록 가난해지는 이유", "절약이 당신을 평생 가난하게 만든다"
2. 자기기만 적발(뜨끔하게): "공부만 하다 1년을 통째로 날리는 사람들의 5가지 변명"
3. 단정 선언(반박 불가 톤): "결국 전부 이거 하나에서 갈렸다", "답은 처음부터 정해져 있었다"
4. 날 선 대비: "버는 사람과 버는 척하는 사람", "3년 버틴 사람과 3개월에 접은 사람의 통장"
5. 충격 숫자: "10명 중 8명이 1년 안에 무너지는 진짜 이유", "상위 3%만 아는 한 가지"
6. 공포·손실 회피: "지금 이걸 모르면 5년 뒤 똑같은 자리에 있다", "당신이 놓치는 사이 벌어지는 일"
7. 미스터리 갭(정체를 가림): "퇴사 3개월 만에 다시 돌아온 회사원이 깨달은 단 하나"

[좋은 변환 — 평이 → 자극, 어법은 자연스럽게]
- "배당주 고르는 법" → "10년 배당을 받고도 한 푼도 못 쓴 사람들의 공통점"
- "복리의 중요성" → "같은 돈을 넣었는데 7년 뒤 잔고가 두 배로 갈린 이유"
- "감정 관리가 필요하다" → "계좌를 들여다본 횟수가 수익률을 갉아먹은 증거"
- "분산 투자를 하자" → "한 종목에 몰빵한 사람이 그해 오히려 덜 잃은 까닭"
- "꾸준함이 답이다" → "3년을 버틴 사람과 3개월에 그만둔 사람의 통장이 비슷했던 이유"

[규칙]
- 소제목 절반 이상에 구체 숫자(금액·기간·비율·인원)를 박는다.
- 통념을 뒤집거나 자기기만을 찌르는 줄을 PART마다 최소 1개.
- 답을 같은 줄에서 다 말하지 마라. "왜/어떻게/무엇이"의 정체는 본문에 숨긴다.
- 사람을 모욕하지 마라(욕설·인신공격 금지). 통념과 '행동'을 때리되 독자를 적으로 만들지 않는다.
- 자청·역행자 등 특정 작가의 '단어'는 쓰지 말 것. 톤만 가져온다.
- 자극을 위해 억지 비유나 말 안 되는 조합을 만들지 마라. 아래 [자연스러움 원칙]이 항상 우선.

[자연스러움 원칙 - 절대 어기지 말 것]
⚠️ 어법이 1순위다. 호기심보다 "말이 되는가"가 먼저다. 어법이 어색하면 호기심을 줄여서라도 자연스러운 문장으로 바꾼다.
- 모든 소제목/챕터 제목은 한국어 원어민 편집자가 손대지 않고 통과시킬, 문법적으로 완결된 자연스러운 문장이어야 한다.
- 단어를 억지로 조합한 "그럴듯해 보이지만 뜻이 안 통하는" 표현은 즉시 폐기. 특히 은유·비유를 무리하게 끼워 맞추지 마라.
  ❌ "배당 사이 파이프 굵기 차이는 정확히 얼마인가" (파이프 굵기? 배당 사이? — 무슨 말인지 알 수 없음)
  ❌ "수익이 자라는 토양의 산도를 맞추는 법" (억지 비유, 어법 어색)
  ✅ "같은 배당주를 사도 누구는 월세처럼 받고 누구는 못 받는 이유" (자연스럽고 호기심도 유발)
- 어법 검사: 주어와 서술어가 자연스럽게 연결되는가? ("회복이 굴러간다" X — 회복은 굴러가지 않음)
- 추상 개념 + 기계·물리·생물 동사 금지: 계좌·통장·재투자·수익·복리·멘탈·습관·시스템 같은 추상/사물에 "켜다/끄다/돌리다/감다/조이다/풀다/꽂다/심다/굴린다/얹는다/멈춘다/자란다/살아난다/숨쉰다/깨어난다" 같은 동사를 붙이지 마라. (계좌는 자라지 않고, 시스템은 살아나지 않는다)
  ❌ "재투자를 켠 사람과 끈 사람" / "수익률 화면을 끄니 계좌가 살아났다" / "내버려 두면 계좌가 자란다"
  ✅ "수익을 다시 넣은 사람과 빼서 쓴 사람의 3년 뒤" / "계좌를 덜 들여다본 해에 수익률이 더 높았던 이유" (대비·호기심은 살리되 주어-서술어가 말이 되게)
- 과장 형용사 금지: "흔적도 없다", "통째로", "완전히" (꼭 필요할 때만)
- 추상 X, 구체 O: "많은 사람" → "월급 280만원짜리 7년차 회사원"
- 도구/플랫폼명 적극: 네이버, 카카오, 노션, 카톡, 캘린더, 구글 시트
- 출력 직전, 모든 줄을 소리내어 읽어라. 원어민이 한 번에 이해 못 하거나 "이게 무슨 말이지?" 싶은 줄은 전부 다시 쓴다.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🚫 즉시 폐기 표현
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

특정 작가 고유어 (절대 금지): 자청 / 역행자 / 자의식 해체 / 유전자 역행 / 원시인 / 추월차선 / 아토믹 해빗 / 언카피어블
AI 클리셰: 졸업 / 정체 / 마지막 한 수 / 다른 차원 / 결정적 시그널 / 진짜 게임 / 흔적도 없다 / 사건이 아니다 / 회로가 박혔다
시스템 의인화: "○○가 멈춘 날", "○○ 위에 얹다", "○○를 졸업한", "회복이 굴러가다"
뇌과학 비유 남용: "뇌 회로가 박혔다", "뇌 회로가 새로 깔린다" (전체 목차에 뇌·신경회로는 사실 진술로 1~2회만, 비유 남용 X)
밍밍: 효과적인 / 성공적인 / ~의 모든 것 / ~하는 방법 / 알아야 할 / 의 중요성
유치 비유: 나침반 / 열쇠 / 보물 / 황금 / 마법 / 파이프라인 / 엔진 / 톱니바퀴 / 사이클 / 눈덩이
참고서: 첫걸음 / 완벽가이드 / 핵심정리 / 기초/중급/고급 / 첫 설계 / 첫 셋업
챕터 제목 라벨: "발굴 —", "1단계:", "STEP 1." 같은 분리 형식
의문문 문어체: "왜 ~는 ~하지 못하는가" 식의 한 PART에 1개까지만
콜론(:) — 단 한 번도 쓰지 마라
숫자 중복 금지: 전체 목차에서 같은 숫자(예: 3시간 + 3개월) 두 번 등장 금지
숫자 표기 - 부호 금지: "-12%" "+30%" 같은 부호 사용 X. "12% 폭락에도", "30% 상승하면" 식으로

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📝 출력 형식 (이 형식 외 어떤 텍스트도 출력 금지)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

PART 1. [13~18자, 좌절 + 통계 충격]
- [통계 충격 또는 결정적 함정형]
- [인지부조화/통념 박살형]
- [패턴 - 위 8개 중 다른 것]
- [컨셉명 첫 등장: 자연스러운 도입]

PART 2. [13~18자, 통념 박살 + 뇌과학 권위]
- [패턴]
- [패턴 - 다른 것]
- [패턴 - 또 다른 것]
- [패턴 - 또 다른 것]

PART 3. [13~18자, 첫 사건/결정적 순간]
- [패턴]
- [임박감/손해회피형]
- [패턴]
- [컨셉명 등장 가능: "○○○를 처음 적용한 사람들이 가장 먼저 느낀 변화" — PART 4에 넣을 거면 여기는 일반 패턴]

PART 4. [13~18자, 정체성 전환 선언]
- [컨셉명 등장 가능: "○○○가 본격 작동하는 4가지 신호" — PART 3에 안 넣었다면 여기에]
- [뇌과학 권위형 또는 통계형]
- [패턴]
- [패턴]

PART 5. [13~18자, 도약/이중 보상]
- [호기심 갭형]
- [이중/삼중 보상형]
- [패턴 - 또 다른 것]
- [컨셉명 세 번째 등장: 인그룹 사회증명]

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔍 출력 전 자가 점검 (반드시 통과)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

출력 직전에 5가지 모두 통과해야 한다. 하나라도 No면 다시 써라.

체크 1. 5개 챕터 제목만 이어 읽었을 때 "X였던 사람이 Y로 바뀐다"는 변화 서사가 또렷한가?
체크 2. 통념 박살(인지부조화) 패턴이 5개 챕터+20개 소제목 안에 3개 이상 박혀있는가?
체크 3. 구체적 숫자(시간/금액/비율)가 8개 이상 등장하는가?
체크 4. 평이한 표현("~의 방법", "~의 모든 것", "~의 중요성", "효과적인", "성공적인")이 0개인가?
체크 5. 5초간 훑어본 가상 독자가 "이건 안 사면 손해"라고 느낄 만한 손실회피 트리거가 챕터 제목 5개 중 2개 이상에 있는가?
체크 6. 호기심 갭(결과만 보이고 방법/이유는 숨김)이 20개 소제목 중 10개 이상에 박혀 있는가? — "결과만 보이는데 본문을 사야 알 수 있는 한 줄"이 절반 이상이어야 결제 전환됨.
체크 7. 컨셉명 글자 수를 단정하는 표현("세 글자", "네 글자", "단 한 단어")이 단 하나도 없는가? — 하나라도 있으면 즉시 전체 다시 쓰기.
체크 8. (어법 - 가장 중요) 모든 챕터 제목·소제목을 소리내어 읽었을 때, 원어민이 한 번에 이해되고 어법이 자연스러운가? "배당 사이 파이프 굵기 차이는…" 같은 억지 조합·뜻 모를 비유가 단 하나라도 있으면 그 줄을 자연스럽게 다시 써라.
체크 9. (어법) 추상 개념에 기계·물리 동사를 붙인 줄("재투자를 켠/끈", "수익을 돌린다" 류)이 하나도 없는가? 있으면 즉시 자연스러운 표현으로 교체.
체크 10. (자극) 통념을 정면으로 뒤집거나 자기기만을 찌르는 줄이 PART마다 최소 1개 있는가? 전부 정보 전달형이면 평이한 것 — 다시 써라.

목차만 출력. 콜론 금지. 매 소제목 다른 패턴. 명사·동사 엔딩 섞기. 어법 어색한 표현 즉시 폐기. 각 PART는 정확히 소제목 4개. 컨셉명 글자수 단정 금지. 자가점검 결과는 출력하지 말 것."""

        chapters = []
        subtopics = {}
        for _outline_attempt in range(2):
            outline_result = ask_ai(outline_prompt, 0.85, ensure_quality=True)
            if not outline_result:
                continue

            chapters = []
            subtopics = {}
            current_ch = None

            lines = outline_result.split('\n')
            for i, orig_line in enumerate(lines):
                line = orig_line.strip()
                if not line:
                    continue

                # 챕터 감지 (PART, 파트, Chapter, 1., 2. 등 다양한 형식)
                is_chapter = False
                ch_name = None

                # 마크다운 강조/헤더 기호를 벗겨낸 감지용 라인
                # (신형 Claude 모델이 **PART 1. ...**, ## PART 1, > 등으로 헤더를 감싸 출력해도 인식)
                detect_line = re.sub(r'^[\s>#\*_`~]+', '', line)
                detect_line = re.sub(r'[\*_`~]+$', '', detect_line).strip()

                # PART 1. 제목 형식
                if re.match(r'^(PART|파트|Part)\s*\d+[\.\s]', detect_line, re.IGNORECASE):
                    is_chapter = True
                    ch_name = detect_line
                # Chapter 1. 제목 형식
                elif re.match(r'^(Chapter|챕터)\s*\d+[\.\s]', detect_line, re.IGNORECASE):
                    is_chapter = True
                    ch_name = detect_line
                # 마크다운 헤더 형식
                elif re.match(r'^#+\s*(PART|파트|Chapter|챕터|\d+)', line, re.IGNORECASE):
                    is_chapter = True
                    ch_name = re.sub(r'^#+\s*', '', line)
                # 1. 제목 형식 (숫자로 시작, 들여쓰기 없음)
                elif re.match(r'^\d+[\.\)]\s', detect_line) and not orig_line.lstrip('*_`~ ').startswith(' '):
                    is_chapter = True
                    ch_name = detect_line
                # 【1부】 형식
                elif re.match(r'^[【\[]?\s*\d+\s*(부|장|편)[】\]]?', detect_line):
                    is_chapter = True
                    ch_name = detect_line

                if is_chapter and ch_name:
                    ch_name = re.sub(r'^[#\*\-\s]+', '', ch_name)
                    ch_name = ch_name.replace('**', '').replace('*', '').replace('#', '').strip()
                    if ch_name and len(ch_name) > 3:
                        current_ch = ch_name
                        if current_ch not in chapters:
                            chapters.append(current_ch)
                            subtopics[current_ch] = []
                elif current_ch:
                    # 소제목 감지
                    is_subtopic = False
                    st_name = None

                    # - 소제목 형식
                    if re.match(r'^[\-\•\·\*\→\▶]\s*', line):
                        is_subtopic = True
                        st_name = re.sub(r'^[\-\•\·\*\→\▶]\s*', '', line)
                    # 1) 소제목, a) 소제목 형식
                    elif re.match(r'^[a-z\d][\)\.\:]\s', line, re.IGNORECASE):
                        is_subtopic = True
                        st_name = re.sub(r'^[a-z\d][\)\.\:]\s*', '', line, flags=re.IGNORECASE)
                    # 들여쓰기된 라인
                    elif orig_line.startswith('  ') or orig_line.startswith('\t'):
                        is_subtopic = True
                        st_name = line.lstrip('- •·*→▶0123456789.):\t ')
                    # 챕터가 아닌 일반 텍스트 (이전이 챕터였고, 현재가 짧은 문장이면 소제목으로 간주)
                    elif len(chapters) > 0 and not re.match(r'^(PART|파트|Part|Chapter|챕터|\d+[\.\)])', line, re.IGNORECASE):
                        if len(line) > 5 and len(line) < 100:
                            is_subtopic = True
                            st_name = line.lstrip('- •·*→▶0123456789.):\t ')

                    if is_subtopic and st_name:
                        st_name = st_name.replace('**', '').replace('*', '').replace('#', '').strip()
                        st_name = re.sub(r'^\d+[\.\)\:]\s*', '', st_name)  # 앞 숫자 제거
                        if st_name and len(st_name) > 3 and len(subtopics[current_ch]) < 4:
                            # 챕터 이름과 동일하면 스킵
                            if st_name.lower() != current_ch.lower() and st_name not in subtopics[current_ch]:
                                subtopics[current_ch].append(st_name)

            # 챕터가 충분히 파싱됐으면 재시도 중단, 부족하면 한 번 더 생성 시도
            if len(chapters) >= 5:
                break

        if chapters:
            st.session_state['outline'] = chapters
            st.session_state['chapters'] = {}
            for ch in chapters:
                st.session_state['chapters'][ch] = {
                    'subtopics': subtopics.get(ch, []),
                    'subtopic_data': {s: {'questions': [], 'answers': [], 'content': ''} for s in subtopics.get(ch, [])}
                }

        # 목차가 생성되지 않았으면 기본 목차 생성
        if not st.session_state.get('outline'):
            progress_placeholder.warning("⚠️ AI 목차 생성/파싱에 실패해 기본 목차를 사용합니다. (API 키·모델 설정을 확인하거나 다시 시도해 주세요)")
            default_chapters = [
                "PART 1. 왜 지금인가",
                "PART 2. 진짜 비밀",
                "PART 3. 실전 공식",
                "PART 4. 수익화",
                "PART 5. 다음 단계"
            ]
            default_subtopics = {
                default_chapters[0]: [f"90%가 {topic}에 실패하는 이유", "아무도 말해주지 않는 진실", "지금 시작해야 하는 3가지 이유"],
                default_chapters[1]: ["전문가들이 숨기는 핵심 원칙", f"{topic}의 본질을 꿰뚫는 법", "이것만 알면 절반은 성공"],
                default_chapters[2]: ["바로 써먹는 5단계 공식", "실패 없이 시작하는 체크리스트", "첫 달에 결과 내는 비법"],
                default_chapters[3]: ["월 100만원 만드는 구조", "자동화로 시간 벌기", "확장 전략 A to Z"],
                default_chapters[4]: ["1년 후 당신의 모습", "다음 레벨로 가는 로드맵", "지금 바로 해야 할 첫 번째 행동"]
            }
            st.session_state['outline'] = default_chapters
            st.session_state['chapters'] = {}
            for ch in default_chapters:
                st.session_state['chapters'][ch] = {
                    'subtopics': default_subtopics.get(ch, []),
                    'subtopic_data': {s: {'questions': [], 'answers': [], 'content': ''} for s in default_subtopics.get(ch, [])}
                }

        # 저자명 및 인터뷰 데이터 저장
        st.session_state['author_name'] = interview_data.get('author_name', '')
        st.session_state['interview_data'] = interview_data
        st.session_state['topic'] = topic

        progress_placeholder.success("✅ 목차 생성 완료! 목차를 확인하고 수정할 수 있습니다.")
        return True

    except Exception as e:
        progress_placeholder.error(f"오류 발생: {str(e)}")
        return False

def regenerate_single_subtopic(chapter_name, subtopic_index, existing_subtopics):
    """개별 소제목 AI 재생성 - 자기계발 베스트셀러 톤"""
    topic = st.session_state.get('topic', '')
    book_concept = st.session_state.get('book_concept', '')

    # 기존 소제목들 (중복 방지용)
    other_subtopics = [s for i, s in enumerate(existing_subtopics) if i != subtopic_index]

    prompt = f"""당신은 한국 자기계발 베스트셀러 편집자입니다. 목차만 보고 결제하게 만드는 소제목 하나를 써주세요.

[책 컨셉]
{book_concept}

[챕터]: {chapter_name}
[주제]: {topic}

[기존 소제목들 - 이것들과 완전히 다르게]
{chr(10).join(f'- {s}' for s in other_subtopics)}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔥 한국 자기계발 베스트셀러 톤 예시 (참고만, 그대로 베끼지 말 것)
- "5분만에 돈 되는 사업 아이템 찾는 비법"
- "월급 280만원이 월 1,000만원으로 바뀌기까지 47일"
- "회사 몰래 부업하다 잘리는 직장인의 진짜 이유"
- "노력이 결과로 안 바뀌는 결정적 함정"
- "첫 정산 받고 인생이 달라진 그날의 기록"

✅ 형식 (하나 선택, 매번 다른 패턴):
- 비법/법형: "5분만에 ~하는 비법", "~하는 정확한 방법"
- 충격형: "~는 거짓말이다", "~하면 오히려 망한다"
- 간증형: "~받고 인생이 달라졌다", "~를 깨닫기까지"
- 도발형: "~은 필요 없다", "~만 있으면 된다"
- 질문형: "왜 ~은 실패하는가"
- 숫자형: "정확히 47일 만에 일어난 일"
- 호기심 갭형(가장 강력): "월 500을 찍은 사람들이 가장 먼저 끊은 습관 한 가지" (결과만 보이고 방법은 숨김)

🧲 호기심 갭 우선 — 결과/사건/디테일은 보이고, 원리/방법/순서는 숨기면 결제 전환률 폭증

⚠️ 컨셉명을 쓰는 경우, 글자 수 단정 표현 절대 금지 ("세 글자", "네 글자", "단 한 단어" 등 — 컨셉명 길이와 안 맞으면 즉시 폐기)

❌ 금지:
- 자청, 역행자, 유전자, 원시인, 추월차선 등 특정 작가 고유 표현
- 설명서 표현: "~의 이해", "~하는 방법"
- 유치한 비유: 나침반, 열쇠, 마법, 톱니바퀴, 파이프라인
- AI 어휘: 효과적인, 다양한, ~를 통해
- 기존 소제목과 비슷한 패턴
- 컨셉명 글자 수 단정 ("세 글자에서 시작한다" 류)

소제목 하나만 (15~30자, 기호 없이):"""

    result = ask_ai(prompt, 0.9)
    if result:
        return result.strip().strip('"').strip("'").strip('-').strip()
    return None

def regenerate_chapter_subtopics(chapter_name, chapter_index):
    """챕터의 모든 소제목 AI 재생성 - 자기계발 베스트셀러 톤"""
    topic = st.session_state.get('topic', '')
    book_concept = st.session_state.get('book_concept', '')
    outline = st.session_state.get('outline', [])

    # 다른 챕터들의 소제목들 (중복 방지)
    other_chapter_subtopics = []
    for ch in outline:
        if ch != chapter_name:
            ch_data = st.session_state['chapters'].get(ch, {})
            other_chapter_subtopics.extend(ch_data.get('subtopics', []))

    # 챕터별 역할 정의
    chapter_roles = {
        0: "착각/각성 - 독자가 몰랐던 불편한 진실을 폭로",
        1: "해체 - 기존 상식과 믿음을 완전히 무너뜨림",
        2: "구조/재구축 - 저자만의 새로운 방법론 제시",
        3: "실전 - 구체적이고 따라할 수 있는 방법",
        4: "도약 - 변화된 미래와 행동 촉구"
    }
    current_role = chapter_roles.get(chapter_index, "핵심 내용 전달")

    prompt = f"""당신은 한국 자기계발 베스트셀러 편집자입니다. 이 챕터의 소제목 4개를 결제하고 싶게 써주세요.

[책 컨셉]
{book_concept}

[주제]: {topic}
[챕터]: {chapter_name}
[이 챕터의 역할]: {current_role}

[다른 챕터 소제목들 - 완전히 다르게 써야 함]
{chr(10).join(f'- {s}' for s in other_chapter_subtopics[:8])}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔥 한국 자기계발 베스트셀러 톤 예시 (참고만, 그대로 베끼지 말 것)
- "5분만에 돈 되는 사업 아이템 찾는 비법"
- "월급 280만원이 월 1,000으로 바뀌기까지 47일"
- "회사 몰래 부업하다 잘리는 직장인의 진짜 이유"
- "첫 정산 받고 인생이 달라진 그날의 기록"
- "노력이 결과로 안 바뀌는 결정적 함정"
- "통장 잔고 23만원에서 시작한 한 가지 시도"

✅ 4개 소제목 모두 다른 형식으로 (4개 다 같은 패턴 금지):
1번: 통계 충격/숫자형 (예: "정확히 47일째에 달라진 한 가지", "1년 안에 99%가 다시 무너지는 결정적 이유")
2번: 간증/사건형 (예: "27만원짜리 첫 정산서가 알려준 한 가지", "~받고 인생이 달라진 그날")
3번: 충격/통념 박살형 (예: "노력할수록 가난해진다", "~만 있으면 된다")
4번: 호기심 갭형 (결과만 보이고 방법은 숨김 — 예: "월 500을 찍은 사람들이 가장 먼저 끊은 습관 한 가지", "성공한 부업러가 매일 밤 11시에 반드시 끄는 것")

🧲 호기심 갭 규칙 (반드시 4개 중 1개 이상 포함)
   - 결과/사건/디테일은 보이고, 원리/방법/순서는 숨겨라
   - "방법", "비법", "노하우"로 끝내면 갭이 닫혀 결제 안 함
   - 한 줄에 "묘하게 구체적인 디테일 + 의문"이 같이 있어야 작동

⚠️ 컨셉명을 쓰는 경우, 글자 수 단정 표현 절대 금지 ("세 글자", "네 글자", "단 한 단어" 등)

❌ 금지:
- 자청, 역행자, 유전자, 원시인, 추월차선 등 특정 작가 고유 표현
- 설명서 표현: "~의 이해", "~하는 방법", "효과적인", "다양한"
- 유치한 비유: 나침반, 열쇠, 마법, 톱니바퀴, 파이프라인
- 같은 패턴 반복

소제목 정확히 4개만 출력 (줄바꿈으로 구분, 기호/번호 없이, 각 줄 15~30자):"""

    result = ask_ai(prompt, 0.8)
    if result:
        lines = [line.strip().strip('"').strip("'").strip('-').strip() for line in result.strip().split('\n') if line.strip() and len(line.strip()) > 5]
        return lines[:4] if lines else None
    return None

def generate_body_from_outline(interview_data, progress_placeholder):
    """생성된 목차를 기반으로 본문만 생성"""
    try:
        topic = interview_data.get('topic', '')
        book_concept = st.session_state.get('book_concept', '')

        if not st.session_state.get('outline') or not st.session_state.get('chapters'):
            progress_placeholder.error("먼저 목차를 생성해주세요.")
            return False

        # 본문 생성
        total_subtopics = sum(len(st.session_state['chapters'][ch]['subtopics']) for ch in st.session_state['outline'])
        done = 0

        for ch in st.session_state['outline']:
            ch_data = st.session_state['chapters'][ch]
            for sub in ch_data['subtopics']:
                done += 1
                progress_placeholder.info(f"✍️ 본문 작성 중... ({done}/{total_subtopics}) - {sub[:20]}...")

                # 이전 소제목들의 내용 요약 (중복 방지용)
                prev_contents = []
                for prev_ch in st.session_state['outline']:
                    if prev_ch == ch:
                        break
                    prev_ch_data = st.session_state['chapters'].get(prev_ch, {})
                    for prev_sub in prev_ch_data.get('subtopics', []):
                        prev_content = prev_ch_data.get('subtopic_data', {}).get(prev_sub, {}).get('content', '')
                        if prev_content:
                            prev_contents.append(f"- {prev_sub}: {prev_content[:100]}...")

                # 현재 챕터의 이전 소제목들
                current_ch_prev = []
                for prev_sub in ch_data['subtopics']:
                    if prev_sub == sub:
                        break
                    prev_content = ch_data.get('subtopic_data', {}).get(prev_sub, {}).get('content', '')
                    if prev_content:
                        current_ch_prev.append(f"- {prev_sub}: {prev_content[:100]}...")

                prev_summary = "\n".join(prev_contents[-5:] + current_ch_prev) if (prev_contents or current_ch_prev) else "없음"

                # 소제목 인덱스에 따라 다른 시작 스타일 선택
                # 더 다양한 시작 스타일 (15가지)
                hook_styles = [
                    "도발적 질문 (예: '왜 99%는 이걸 모를까요?')",
                    "충격적 고백 (예: '저도 3년간 완전히 잘못하고 있었습니다.')",
                    "반전 사실 (예: '사실 정반대였습니다.')",
                    "구체적 숫자 (예: '정확히 47일 만에 달라졌습니다.')",
                    "생생한 에피소드 (예: '그날 카페에서 노트북을 열었을 때였습니다.')",
                    "단호한 선언 (예: '결론부터 말씀드리겠습니다.')",
                    "대화체 시작 (예: '\"이게 진짜 돼요?\" 처음 들었을 때 저도 그랬습니다.')",
                    "before/after (예: '6개월 전만 해도 저는 완전히 다른 사람이었습니다.')",
                    "상식 뒤집기 (예: '노력하면 된다? 완전히 틀렸습니다.')",
                    "비유로 시작 (예: '이건 마치 고장난 네비게이션을 따라가는 것과 같습니다.')",
                    "독자 공감 (예: '혹시 이런 경험 있으신가요?')",
                    "미래 제시 (예: '3개월 후, 완전히 다른 결과를 보게 될 겁니다.')",
                    "실패담 (예: '처음엔 완전히 망했습니다.')",
                    "발견의 순간 (예: '그때 깨달았습니다. 방법이 틀렸던 거였죠.')",
                    "핵심 한 줄 (예: '핵심은 딱 하나입니다.')",
                ]
                current_hook_style = hook_styles[done % len(hook_styles)]

                # 저자 이름 가져오기
                author_name = interview_data.get('author_name', '') or st.session_state.get('author_name', '') or '저자'

                content_prompt = f"""당신은 한국 자기계발 베스트셀러 작가입니다. 첫 문장으로 독자를 멈춰 세우고, 권석천 기자 칼럼처럼 정확한 디테일과 통찰로 끝까지 끌고 가는 본문을 씁니다.

🚨🚨🚨 최우선 규칙 🚨🚨🚨
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
1. 첫 문장이 가장 중요! 반드시 이 스타일로 시작:
   👉 {current_hook_style}

2. 이전 내용과 절대 중복 금지!
   아래 내용은 이미 썼으니 완전히 다른 이야기를 해라:
   {prev_summary}

3. 독자 직접 호칭 금지
   ❌ "여러분", "당신", "독자님", "~하시는 분들"
   ✅ "저는", "우리는", "제가"

4. 특정 작가 고유 표현 절대 금지 (저작권/표절 위험)
   ❌ "자청", "자청아", "자청씨"
   ❌ "역행자", "유전자 역행", "유전자 오작동", "자의식 해체"
   ❌ "원시인", "추월차선", "아토믹 해빗", "언카피어블"
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[집필 정보]
주제: {topic}
챕터: {ch}
소제목: {sub}
핵심 방법론: {interview_data.get('core_method', '')}

[책 컨셉]
{book_concept}

이 책의 기법 이름은 "{st.session_state.get('method_name', '')}"으로 확정됨. 본문에서 기법을 언급할 때 반드시 이 이름 그대로 사용(약자 풀어쓰기·재작명 금지).

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
✍️ 본문 톤: 자기계발 후킹 + 권석천 칼럼 깊이
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[문체]
- 합쇼체 기본 ("~입니다", "~합니다") + 가끔 구어체 ("~거든요", "~더라고요")
- 현재 시제로 장면을 그리듯
- 짧은 문장과 긴 문장을 교차해 리듬을 만듦
- 추상보다 구체. "많은 사람" → "월급 280만원짜리 7년차 회사원"
- 결론부터, 분석은 그 다음

[권석천식 깊이 - 칼럼처럼 인과를 추적]
- 사실 → 분석 → 통찰 순서로 전개
- 가설 검증식 흐름: "왜 그럴까. 이유를 되짚어봤습니다"
- 사회적 맥락이나 통계, 책/논문 인용 자연스럽게 섞기
- 결론을 강요하지 말고, 독자가 스스로 깨닫게 단서를 깔기

[자기계발식 후킹 - 뇌를 멈춰 세우는 첫 문장]
- 사건/숫자/대사/의문 중 하나로 시작 (일반론 절대 금지)
- 본문 중간에 작은 반전 1회 ("그런데 진짜 흥미로운 건 그 다음이었습니다")
- 마지막 문장은 발견의 결과로서의 통찰 한 줄

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🎯 구체성을 끝까지 밀어붙여라
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

✅ 숫자: "많이" X → "월 340만원, 정확히 47일" O
✅ 도구/플랫폼: 네이버, 카카오, 노션, 카톡, 구글 시트 등 실제 이름
✅ 실행 순서: "무엇을 → 어디서 → 어떻게" 명시

✅ 사례:
- 내 경험: "제가 처음 시작했을 때", "그날 회의실에서"
- 타인 사례: 수강생, 지인 (이름은 가끔만, 매번 X)
- 사례 없어도 OK. 일반적 원리·논리 전개로 충분

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🚫 절대 금지
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
❌ 자청, 역행자, 자의식 해체, 유전자 역행, 원시인, 추월차선, 아토믹 해빗 (특정 작가/책 고유 표현)
❌ 같은 이름 반복 (민준, 지수가 계속 나오면 안 됨)
❌ "김씨", "이씨" 같은 성씨 호칭
❌ 유치한 표현: 후다닥, 짜잔, 대박, ㅋㅋ, 어마어마한
❌ 유치한 비유: 마법, 황금열쇠, 나침반, 로켓, 눈덩이
❌ 억지 메타포: 순환법, 엔진, 고리, 파이프라인, 톱니바퀴
❌ AI스러운: 중요합니다, 따라서, 결론적으로, ~를 통해, 다양한, 효과적인, 진정한
❌ 형식: 1. 2. 첫째, 둘째, 글머리 기호, 이모지

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📊 표는 '필요할 때만' (강제 아님)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
- 비교(전/후, A안/B안), 단계별 정리, 수치 묶음처럼 표로 보여주면 확 명확해지는 내용이 있을 때만 표 1개를 넣어라.
- 그런 내용이 없으면 표를 넣지 마라. 모든 소제목에 표가 있을 필요 없다(대략 3개 중 1개 정도가 자연스럽다).
- 표 형식은 반드시 아래 마크다운 파이프 형식. 첫 줄이 헤더다. (HTML 태그 금지)
  | 구분 | 기존 방식 | 새로운 방식 |
  | 비용 | 월 30만원 | 월 0원 |
  | 시간 | 하루 3시간 | 하루 30분 |
- 표 앞뒤로는 반드시 설명 문단을 둔다. 표만 툭 던지지 마라.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔥 몰입·충격 어조 (자청식 몰입감 — 특정 작가 고유어는 제외)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

- 도입부터 통념을 정면으로 깬다. 독자가 "어, 내가 알던 거랑 반대네" 하고 멈칫하게.
- 글 중간마다 '작은 충격'을 심어라: 예상과 반대되는 사실, 뒤집히는 데이터, 의외의 결론.
- 한 호흡에 읽히는 리듬 — 짧은 단정문으로 치고, 긴 문장으로 풀고, 다시 끊는다.
- "왜?"를 계속 만들어 다음 문단을 안 읽고는 못 배기게 한다(궁금증 갭).
- 충격 뒤엔 반드시 '그래서 무엇을 어떻게'의 실질이 온다. 겁만 주고 끝내지 않는다.
- 자청·역행자 등 특정 작가의 '단어'는 절대 쓰지 않는다. 몰입감과 태도만 가져온다.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🎓 전문성 (아마추어 글과 가르는 지점)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

- 추상 주장 1개당 근거 1개(데이터·연구·구체 사례·숫자)를 반드시 붙인다.
- 메커니즘을 설명하라. "왜 그렇게 되는가"의 원리와 과정을 단계로 풀어라.
- 바로 따라 할 수 있게: 무엇을·어디서·어떤 순서로·얼마나, 구체적으로.
- 흔한 조언의 한계를 짚고, 더 정확한 기준·예외를 제시한다(깊이의 증거).

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📏 분량: 2400~3000자 (최소 2400자 이상, 전문성 있는 밀도로)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

⚠️ 길이는 사례·데이터·단계별 설명·메커니즘으로 채운다. 같은 말 반복·미사여구로 늘리면 즉시 실패.

'{sub}' 본문 작성.
- 시작: {current_hook_style}
- 이전 내용과 완전히 다른 새로운 이야기
- 권석천 칼럼처럼 사실 → 분석 → 통찰 순서로 인과를 추적
- 마지막 한 줄에 발견된 통찰 하나
- 표가 정말 도움이 되는 내용이면 위 마크다운 파이프 형식으로 1개까지 넣어도 됨(아니면 넣지 말 것)

⛔ 절대 금지: 본문 첫 줄에 소제목('{sub}')을 다시 쓰지 마라.
   소제목은 위에 이미 표시되므로, 본문은 곧장 첫 후킹 문장으로 시작한다.
   ❌ 잘못: "{sub}\\n\\n그날 새벽 두 시였습니다..." (소제목 반복)
   ✅ 올바름: "그날 새벽 두 시였습니다..." (바로 본문 시작)"""

                content = ask_ai(content_prompt, 0.7, ensure_quality=True)
                if content:
                    content = clean_content(content, subtopic=sub)  # 이모티콘/마크다운 제거 + 소제목 중복 제거
                    ch_data['subtopic_data'][sub]['content'] = content

            # ── 챕터 끝: 핵심 정리(key_points) 생성 ──
            try:
                chapter_text = ""
                for s in ch_data['subtopics']:
                    c = ch_data.get('subtopic_data', {}).get(s, {}).get('content', '')
                    if c:
                        chapter_text += f"\n[{s}]\n{c[:600]}\n"
                if chapter_text.strip():
                    summary_prompt = f"""다음은 책 한 챕터의 본문이다. 이 챕터에서 독자가 반드시 기억해야 할 핵심을 3~4개로 정리하라.

[챕터 제목]: {ch}
[본문 발췌]
{chapter_text[:3000]}

규칙:
- 각 항목은 한 문장(공백 포함 45자 이내), 자연스러운 한국어 평서문.
- 추상적 구호 금지. 본문에서 실제로 다룬 구체적 통찰·행동을 요약.
- 번호·기호·이모지 없이 문장만.

JSON만 출력: {{"points": ["...", "...", "..."]}}"""
                    sres = ask_ai(summary_prompt, 0.5)
                    sparsed = parse_json(sres) if sres else None
                    pts = (sparsed or {}).get('points', []) if sparsed else []
                    pts = [p.strip() for p in pts if isinstance(p, str) and p.strip()][:4]
                    if pts:
                        ch_data['key_points'] = pts
            except Exception:
                pass

        # 완료 처리
        st.session_state['interview_completed'] = True
        progress_placeholder.success("✅ 본문 생성 완료!")
        return True

    except Exception as e:
        progress_placeholder.error(f"오류 발생: {str(e)}")
        return False

def go_next():
    if st.session_state['current_page'] < 7:
        st.session_state['current_page'] += 1

def go_prev():
    if st.session_state['current_page'] > 0:
        st.session_state['current_page'] -= 1

def auto_generate_all(topic, progress_placeholder):
    """주제만 입력하면 목차+본문까지 자동 생성"""
    try:
        # 1. 타겟 자동 생성
        progress_placeholder.info("🎯 1/4 타겟 분석 중...")
        target_result = suggest_targets(topic)
        targets = parse_json(target_result)
        if targets and targets.get('targets'):
            first_target = targets['targets'][0]
            persona = f"{first_target.get('name', '')} - {first_target.get('description', '')}"
            st.session_state['target_persona'] = persona

            # 페인포인트 분석
            pain_result = analyze_pains_deep(topic, persona)
            pain_data = parse_json(pain_result)
            if pain_data:
                st.session_state['pains'] = pain_data.get('pains', [])

        # 2. 목차 자동 생성
        progress_placeholder.info("📋 2/4 목차 생성 중...")
        outline_result = generate_outline(
            topic,
            st.session_state.get('target_persona', ''),
            st.session_state.get('pains', [])
        )

        # 목차 텍스트 파싱 (PAGE 4와 동일한 방식)
        if outline_result:
            chapters = []
            subtopics = {}
            current_ch = None

            for line in outline_result.split('\n'):
                orig_line = line
                line = line.strip()
                if not line:
                    continue

                # 챕터 감지 (PART, 1., 2. 등)
                is_chapter = False
                ch_name = None

                if re.match(r'^(PART|파트)\s*\d+', line, re.IGNORECASE):
                    is_chapter = True
                    ch_name = line
                elif re.match(r'^\d+[\.\)]\s', line):
                    is_chapter = True
                    ch_name = line
                elif re.match(r'^#+\s*(PART|파트|\d+)', line, re.IGNORECASE):
                    is_chapter = True
                    ch_name = re.sub(r'^#+\s*', '', line)

                if is_chapter and ch_name:
                    ch_name = ch_name.replace('**', '').replace('*', '').replace('#', '').strip()
                    if ch_name and len(ch_name) > 3:
                        current_ch = ch_name
                        if current_ch not in chapters:
                            chapters.append(current_ch)
                            subtopics[current_ch] = []
                elif current_ch:
                    # 소제목 감지
                    is_subtopic = False
                    st_name = None

                    if line.startswith('-') or line.startswith('•') or line.startswith('·'):
                        is_subtopic = True
                        st_name = line.strip().lstrip('-•· ')
                    elif re.match(r'^\s+[\da-z][\)\.]', orig_line):
                        is_subtopic = True
                        st_name = re.sub(r'^[\s\da-z\)\.\-]+', '', line).strip()

                    if is_subtopic and st_name:
                        st_name = st_name.replace('**', '').replace('*', '').replace('#', '').strip()
                        if st_name and len(st_name) > 3 and not re.match(r'^(PART|파트|Chapter|챕터)', st_name, re.IGNORECASE):
                            subtopics[current_ch].append(st_name)

            if chapters:
                st.session_state['outline'] = chapters
                st.session_state['chapters'] = {}
                for ch in chapters:
                    st.session_state['chapters'][ch] = {
                        'subtopics': subtopics.get(ch, []),
                        'subtopic_data': {s: {'questions': [], 'answers': [], 'content': ''} for s in subtopics.get(ch, [])}
                    }

        # 3. 본문 자동 생성
        progress_placeholder.info("✍️ 3/4 본문 작성 중...")
        if st.session_state.get('outline') and st.session_state.get('chapters'):
            total_subtopics = sum(len(st.session_state['chapters'][ch]['subtopics']) for ch in st.session_state['outline'])
            done = 0

            for ch in st.session_state['outline']:
                ch_data = st.session_state['chapters'][ch]
                for sub in ch_data['subtopics']:
                    done += 1
                    progress_placeholder.info(f"✍️ 본문 작성 중... ({done}/{total_subtopics})")

                    content = generate_content_premium(sub, ch, [], [], topic, st.session_state.get('target_persona', ''))
                    if content:
                        ch_data['subtopic_data'][sub]['content'] = content
                        ch_data['subtopic_data'][sub]['formatted'] = format_content_html(content)

        # 4. 완료
        progress_placeholder.success("✅ 완료! 본문 페이지로 이동합니다...")
        return True

    except Exception as e:
        progress_placeholder.error(f"오류 발생: {str(e)}")
        return False


# ==========================================
# AI 함수들
# ==========================================
def analyze_market_deep(topic):
    prompt = f"""주제: {topic}

이 주제로 전자책 시장을 분석해주세요.

[중요] 모든 답변은 반드시 한국어로만 작성하세요.

JSON:
{{
    "verdict": "강력 추천/추천/보류/비추천 중 하나",
    "verdict_reason": "판정 이유 한국어로",
    "total_score": 85,
    "search_data": {{
        "naver_monthly": "네이버 월간 검색량 예시: 12,000회",
        "google_monthly": "구글 월간 검색량 예시: 8,500회",
        "naver_blog_posts": "블로그 게시물 수",
        "youtube_videos": "유튜브 영상 수",
        "search_trend": "상승 또는 유지 또는 하락"
    }},
    "market_size": {{
        "score": 85,
        "level": "매우 큼/큼/보통/작음 중 하나",
        "analysis": "분석 2문장 한국어로"
    }},
    "competition": {{
        "score": 70,
        "level": "치열함/보통/낮음 중 하나",
        "your_opportunity": "차별화 기회 한국어로"
    }},
    "profit": {{
        "score": 80,
        "price_range": "권장 가격대",
        "monthly_revenue": "예상 월 수익"
    }},
    "popular_ebooks": [
        {{
            "title": "이 주제 관련 인기 전자책 제목",
            "platform": "크몽/탈잉/클래스101/리디북스/yes24 중 하나",
            "url": "해당 전자책 실제 URL (예: https://kmong.com/xxx)",
            "price": "가격"
        }},
        {{
            "title": "두번째 인기 전자책",
            "platform": "플랫폼명",
            "url": "URL",
            "price": "가격"
        }},
        {{
            "title": "세번째 인기 전자책",
            "platform": "플랫폼명",
            "url": "URL",
            "price": "가격"
        }}
    ],
    "recommendation": "최종 권장 2문장 한국어로"
}}"""
    return ask_ai(prompt, 0.5)


def suggest_targets(topic):
    prompt = f"""주제: {topic}

이 주제의 전자책을 가장 많이 구매할 것 같은 핵심 타겟 3개만 추천해주세요.
가장 적합하고 구매 가능성이 높은 타겟만 엄선해서 3개만 알려주세요.

[중요] 모든 답변은 반드시 한국어로만 작성하세요.

JSON:
{{
    "personas": [
        {{
            "name": "타겟 이름 (구체적으로)",
            "demographics": "연령대, 직업",
            "needs": "이 타겟이 이 책을 사는 이유",
            "pain_points": ["핵심 고민1", "고민2", "고민3", "고민4", "고민5"]
        }}
    ]
}}"""
    return ask_ai(prompt, 0.7)


def analyze_pains_deep(topic, persona):
    prompt = f"""주제: {topic}
타겟: {persona}

이 타겟의 고민을 아주 깊이 분석해주세요.

[중요] 모든 답변은 반드시 한국어로만 작성하세요. 외국어 사용 금지.

JSON:
{{
    "surface_pains": {{
        "pains": ["표면적 고민1", "고민2", "고민3", "고민4", "고민5"],
        "description": "표면적 고민 설명 3문장"
    }},
    "hidden_pains": {{
        "pains": ["숨겨진 진짜 고민1", "고민2", "고민3", "고민4"],
        "description": "숨겨진 고민 설명 3문장"
    }},
    "emotional_pains": {{
        "pains": ["감정적 고통1", "고통2", "고통3"],
        "description": "감정적 고통 설명 2문장"
    }},
    "failed_attempts": {{
        "attempts": ["시도했지만 실패한 것1", "것2", "것3"],
        "why_failed": "실패 이유 2문장"
    }},
    "dream_outcome": {{
        "ideal_result": "이상적인 결과",
        "timeline": "원하는 기간",
        "what_changes": "달라지는 것 2문장"
    }},
    "buying_triggers": {{
        "triggers": ["구매 요인1", "요인2", "요인3"],
        "objections": ["망설임 이유1", "이유2"]
    }},
    "marketing_hook": "마케팅 훅 한 문장"
}}"""
    return ask_ai(prompt, 0.6)


def analyze_competitor_reviews(topic):
    prompt = f"""주제: {topic}

이 주제 관련 전자책/도서의 부정적 리뷰를 분석해주세요.

[매우 중요]
- 모든 답변은 반드시 한국어로만 작성하세요.
- 영어, 러시아어 등 외국어 절대 사용 금지
- 한글과 숫자만 사용하세요.

JSON:
{{
    "analysis_scope": {{
        "books_analyzed": "287권",
        "reviews_analyzed": "3,842개",
        "negative_reviews": "892개 (23%)",
        "platforms": ["크몽", "예스24", "알라딘", "교보문고"]
    }},
    "negative_patterns": [
        {{
            "pattern": "불만 패턴 한국어로",
            "frequency": "67%",
            "example_reviews": ["실제 리뷰 예시 한국어로", "리뷰2"],
            "reader_emotion": "독자 감정 한국어로",
            "hidden_need": "숨겨진 니즈 한국어로",
            "solution": "해결책 한국어로"
        }},
        {{
            "pattern": "두 번째 불만",
            "frequency": "54%",
            "example_reviews": ["리뷰1", "리뷰2"],
            "reader_emotion": "감정",
            "hidden_need": "니즈",
            "solution": "해결책"
        }},
        {{
            "pattern": "세 번째 불만",
            "frequency": "41%",
            "example_reviews": ["리뷰1", "리뷰2"],
            "reader_emotion": "감정",
            "hidden_need": "니즈",
            "solution": "해결책"
        }}
    ],
    "hidden_needs_summary": {{
        "needs": ["숨겨진 니즈1", "니즈2", "니즈3"],
        "insight": "핵심 인사이트 2문장"
    }},
    "concept_suggestions": [
        {{
            "concept": "차별화 컨셉1 한국어로",
            "why_works": "이유 한국어로",
            "unique_point": "차별점 한국어로"
        }},
        {{
            "concept": "컨셉2",
            "why_works": "이유",
            "unique_point": "차별점"
        }}
    ],
    "success_formula": {{
        "must_have": ["필수1", "필수2", "필수3"],
        "must_avoid": ["금지1", "금지2"],
        "differentiation": "차별화 전략 한국어로 2문장"
    }}
}}"""
    return ask_ai(prompt, 0.6)


def generate_titles_bestseller(topic, persona, pains):
    prompt = f"""당신은 교보문고 종이책 베스트셀러와 크몽·클래스101 전자책 베스트셀러를 동시에 분석하는 제목 카피라이터입니다.

주제: {topic}
독자: {persona}
독자의 고민: {pains}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
가장 중요한 규칙 (이걸 어기면 전부 실패)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
제목은 반드시 "실제로 말이 되는 자연스러운 한국어 한 구절"이어야 합니다.
주제에서 뽑은 단어들을 기계적으로 이어 붙이지 마세요.

❌ 단어 짜깁기 실패 (절대 이렇게 만들지 말 것):
  "수면 매출 설계도"  → 수면+매출+설계도, 서로 관계없는 명사를 그냥 붙인 말
  "다이어트 부자 공식"  → 의미가 안 통하는 조합
  "관계 성장 엔진"      → 추상명사만 나열
  이런 제목은 소리내어 읽으면 "이게 무슨 말이지?" 싶고, 어떤 서점에도 존재하지 않습니다.

✅ 자연스러운 제목 = 사람이 실제로 쓰는 어순과 의미가 살아있는 구절
  (수면 주제 예) "잠든 사이 일어나는 일" / "초저녁의 기술" / "다시 잠드는 법"

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
두 시장의 제목 문법
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
[종이책 베스트셀러 — 품격·여운]
  돈의 속성 / 불변의 법칙 / 마흔에 읽는 쇼펜하우어 / 역행자 / 세이노의 가르침
  → 명사+의+명사, 한 단어 임팩트, 약간의 문학성. 광고 냄새가 없다.

[전자책 베스트셀러 — 구체적 약속·호기심]
  크몽·클래스101 상위권은 "독자가 무엇을 얻는지"가 살짝 드러난다.
  단, 유치하지 않게. "~하는 법/방법/노하우/비법/공식" 같은 흔한 꼬리표는 피하고,
  의외의 단어나 시선 전환을 한 번 넣어 호기심을 만든다.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
피해야 할 것
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
- 의미 없는 명사 나열 (가장 큰 실패 원인)
- 비밀/비법/공식/바이블/마법/머니/시스템/파이프라인/연금술
- 완벽한·궁극의·최고의·기적의 같은 과장 형용사
- "월 1000만원" 류 숫자 과시, "직장인을 위한" 류 타겟 명시

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
출력
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
'{topic}' 주제로 서로 다른 결의 제목을 정확히 5개.
(3개는 종이책 품격 톤, 2개는 전자책 호기심 톤)
길이는 자유 — 보통 2~7어절. 짧게 만드는 것보다 자연스러움이 우선이다.

각 제목마다 self_check를 채워, 소리내어 읽었을 때 말이 되는지 스스로 검증할 것.
self_check가 "어색하다/말이 안 된다"면 그 제목은 버리고 다시 만들 것.

JSON만 출력:
{{
    "titles": [
        {{"title": "제목", "subtitle": "호기심을 더하는 부제 한 줄 (20자 이내)", "concept": "이 제목이 매력적인 이유 한 줄", "self_check": "소리내어 읽으면 자연스러운가에 대한 한 문장 자기검증"}}
    ]
}}"""
    return ask_ai(prompt, 0.6, ensure_quality=True)


def analyze_text_content(text, source=""):
    prompt = f"""출처: {source}
내용: {text[:5000]}

분석:

JSON:
{{
    "title": "주제",
    "key_points": ["핵심1", "핵심2", "핵심3", "핵심4", "핵심5"],
    "insights": ["인사이트1", "인사이트2", "인사이트3"],
    "action_items": ["실행1", "실행2", "실행3"],
    "ebook_ideas": ["아이디어1", "아이디어2"],
    "summary": "요약 3문장"
}}"""
    return ask_ai(prompt, 0.5)


def summarize_all_knowledge(items, topic):
    """전체 학습 내용 통합 요약"""
    all_points = []
    all_tips = []
    all_ideas = []

    for item in items:
        if isinstance(item, dict):
            all_points.extend(item.get('key_points', []))
            all_tips.extend(item.get('actionable_tips', item.get('action_items', [])))
            all_ideas.extend(item.get('ebook_applications', item.get('ebook_ideas', [])))

    prompt = f"""전자책 주제: {topic}

학습한 모든 정보를 통합 분석해주세요.

수집된 핵심 포인트들:
{chr(10).join([f"• {p}" for p in all_points[:25]])}

실행 팁들:
{chr(10).join([f"• {t}" for t in all_tips[:15]])}

전자책 활용 아이디어:
{chr(10).join([f"• {i}" for i in all_ideas[:10]])}

JSON:
{{
    "integrated_summary": "전체 학습 내용 통합 요약 5문장",
    "core_insights": [
        "핵심 인사이트 1",
        "인사이트 2",
        "인사이트 3",
        "인사이트 4",
        "인사이트 5"
    ],
    "action_plan": [
        "즉시 실행할 것 1",
        "실행 2",
        "실행 3"
    ],
    "ebook_structure": [
        "추천 목차 1장",
        "2장",
        "3장",
        "4장"
    ],
    "unique_angle": "이 전자책만의 차별화된 관점",
    "study_plan": {{
        "week1": "1주차: 무엇을 할지",
        "week2": "2주차: 무엇을 할지",
        "week3": "3주차: 무엇을 할지",
        "week4": "4주차: 무엇을 할지"
    }},
    "expert_tips": [
        "전문가 팁 1",
        "팁 2",
        "팁 3"
    ]
}}"""
    return ask_ai(prompt, 0.6)


def generate_outline(topic, persona, pains, gaps=None):
    """한국 자기계발 베스트셀러 톤: 결제 버튼을 누르게 하는 자극형 목차"""

    # 페르소나/고통/시장 빈틈 정리 (있으면 활용, 없어도 작동)
    persona_block = f"[타겟 독자]\n{persona}\n" if persona else ""
    if isinstance(pains, list):
        pains_text = "\n".join(f"- {p}" for p in pains if p)
    else:
        pains_text = str(pains) if pains else ""
    pains_block = f"[독자가 지금 느끼는 통증]\n{pains_text}\n" if pains_text else ""
    if gaps:
        gaps_text = "\n".join(f"- {g}" for g in gaps) if isinstance(gaps, list) else str(gaps)
        gaps_block = f"[시장의 빈틈 - 경쟁자가 안 다루는 것]\n{gaps_text}\n"
    else:
        gaps_block = ""

    # 같은 책이면 기존 영문 기법 이름 재사용, 없으면 새로(고유) 생성
    method_name = get_or_create_method_name(topic, None, force_new=False)
    method_expansion = st.session_state.get('method_expansion', '')

    prompt = f"""당신은 한국 자기계발 분야 톱 0.1% 기획자입니다. 서점에서 단 5초간 목차만 본 사람이 책을 손에서 못 놓게 만드는 5장짜리 목차를 씁니다.

목차의 단 하나의 목적: 독자가 "이 책을 안 읽으면 평생 손해"라고 느끼게 만드는 것.
정보 전달은 본문이 한다. 목차는 100% 구매심리만 다룬다.

{method_lock_rule(method_name)}

[주제]: {topic}
{persona_block}{pains_block}{gaps_block}
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🛒 구매 결정 5초 룰 (모든 규칙 중 1순위)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

독자는 목차를 5초만 본다. 그 5초 안에 다음 3가지가 동시에 작동해야 결제한다.

[1] 정체성 변화 발견 — "이걸 읽으면 나는 OO한 사람이 된다"
   → 5개 챕터 제목을 이어 읽으면 한 사람의 결정적 변화가 보여야 한다.
   ✅ "단단해진 멘탈은 인생을 통째로 바꾼다" (변화 서사 O)
   ❌ "멘탈 관리의 다양한 기법" (정체성 변화 X — 즉시 폐기)

[2] 손실회피 작동 — "이걸 모르면 평생 OO한다"
   → 통념 박살(인지부조화) 챕터/소제목이 최소 3개 들어가야 한다.
   ✅ "의지로 버틴 사람일수록 더 크게 무너진다"
   ❌ "멘탈 관리의 중요성" (잃을 게 안 보임)

[3] 구체성 — 추상 명사 1개당 구체적 숫자/장면 1개
   → "많은 사람" 금지, "월급 280만원짜리 7년차 회사원" 가능
   → 시간(90초, 47일), 금액(34만원, 1억), 비율(99%, 8할) 적극 사용

25줄(챕터 5 + 소제목 20) 중 한 줄이라도 "그냥 정보"가 섞이면 그 목차는 평이해진다.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🧠 마케팅 뇌과학 8대 트리거 (목차 전체에 골고루 박을 것)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

1. **손해회피 + 충격 통계** — "1년 안에 99%가 다시 무너진다", "1억치 강의 들어도 망한다"
2. **인지부조화/통념 박살** — "의지로 버틴 사람일수록 더 크게 무너진다", "노력할수록 가난해진다"
3. **권위 어휘 (과학/임상)** — 뇌, 신경회로, N주 후, 임상, 데이터, 알고리즘
4. **임박감 + 절대성** — "이 90초를 놓치면 며칠 걸린다", "다시는 ~하지 않는다"
5. **정체성 전환 약속** — "회복한 뇌는 다시 무너지지 않는다", "단단해진 사람은 ~한다"
6. **인그룹 사회증명** — "○○를 익힌 사람들의 5년 뒤", "상위 1%만 도달하는"
7. **이중/삼중 보상** — "통장과 인간관계가 함께 변한다", "돈도 사람도 따라온다"
8. **호기심 갭** — "두 달 안에 가장 먼저 끊은 한 가지", "정확히 어디부터 멈추는가"

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🧲 호기심 갭(Information Gap) 강화 — 결제를 부르는 가장 강력한 무기
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

조지 로웬스타인의 정보격차 이론: 사람은 자기가 모르는 것이 '뭔지'는 알지만 '내용'은 모를 때 가장 강하게 끌린다. 25줄 중 최소 10줄에 이 갭을 박아야 목차만 보고 결제한다.

[호기심 갭 6대 공식 - 결과만 보이고 방법/이유는 본문으로 미루기]

1. **숫자 + 미공개 결과** — '정확히/딱'으로 시작해 결과만 보여주고 내용은 숨기기
   ✅ "정확히 47일째에 통장이 처음 뒤집힌 그 순간"
   ✅ "27만원짜리 첫 정산서가 알려준 단 한 가지"
   ❌ "47일 만에 돈을 버는 방법" (방법을 다 보여줘버림 → 결제 안 함)

2. **이미 벌어진 사건 + 원인 숨김** — '왜?'를 유발하는 결과만
   ✅ "3년 차 베테랑이 신입에게 6개월 만에 따라잡힌 단 하나의 이유"
   ✅ "월 1,000을 찍은 사람들이 가장 먼저 끊은 습관 한 가지"

3. **묘하게 구체적인 행동/대상 + 이유 숨김** — 디테일이 호기심을 폭발시킨다
   ✅ "성공한 부업러가 매일 밤 11시에 반드시 끄는 것"
   ✅ "1년 만에 1억 모은 사람들이 절대 안 쓰는 5단어"
   ✅ "월 500 넘긴 사람들 카톡 프로필에서 사라진 한 단어"

4. **반대 결과 미스터리** — 통념과 정반대 결과만 던지고 메커니즘은 본문
   ✅ "더 열심히 할수록 더 가난해진 7년의 비밀"
   ✅ "잠을 늘렸더니 매출이 2배가 된 이상한 메커니즘"

5. **'딱 하나' 절대성** — 수많은 변수 중 단 하나만 보여주기
   ✅ "월 100 / 월 500을 가르는 단 한 줄의 차이"
   ✅ "결국 모든 게 무너지는 사람들의 공통점 단 하나"

6. **시간 압축 미스터리** — 짧은 시간에 큰 일이 일어났는데 그 사이를 숨기기
   ✅ "퇴근 후 90분이 1년 뒤 인생을 갈라놓는다"
   ✅ "주말 4시간이 5년치 월급을 바꾼 그 과정"

[호기심 갭 만들 때 절대 어기지 말 것]
• 답을 같은 줄에 다 보여주지 마라. "왜 ~한가" "어떻게 ~하는가"로 끝나면 본문을 사야 알 수 있게.
• "방법" "비법" "노하우" 같은 말로 끝내면 갭이 닫힘 → 결제 안 함.
• 결과/사건/디테일은 보이고, 원리/메커니즘/순서는 숨겨라.
• 한 줄 안에 "장면 + 의문"이 같이 있어야 호기심이 작동한다.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🎯 단 하나의 미션
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

5개 챕터 제목만 빠르게 이어 읽었을 때 한 사람의 변화 이야기가 보이고, 한 줄 한 줄이 다음 챕터를 못 참게 만들어야 한다.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📐 챕터 제목 형식 (가장 중요)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[형식]
- 13~18자의 자연스러운 한국어 한 문장
- 단어 한두 개짜리 라벨 금지, 라벨 분리(— ㅣ :) 금지
- 평서문 또는 단언형. "~이유"로 끝나는 설명체는 한 PART에만 사용
- 명사 엔딩과 동사 엔딩을 챕터별로 섞어라 (5개 모두 명사 엔딩 금지)
- 5개를 이으면 [좌절 → 통념 박살 → 첫 사건 → 안정화 → 도약]의 5막

[좋은 예 - 성공적인 멘탈 관리 비결]
PART 1. 1년 안에 99%가 다시 무너지는 결정적 이유
PART 2. 의지로 잡으려는 순간 뇌는 반대로 움직인다
PART 3. 회복의 8할이 결판나는 폭발 직후 90초
PART 4. 한 번 회복한 뇌는 다시는 무너지지 않는다
PART 5. 단단해진 멘탈은 인생을 통째로 바꾼다

[좋은 예 - 30대 직장인 N잡 월 500]
PART 1. 직장인 99%가 부업 30일을 못 버티는 이유
PART 2. 노력보다 자리가 먼저다
PART 3. 첫 30만원이 통장에 찍힌 그날
PART 4. 새벽 3시에도 매출이 들어온다
PART 5. 월 500 다음, 억대 수익으로 가는 길

[나쁜 예 - 즉시 폐기]
- "발굴", "폭로", "전환" 같은 한두 단어짜리
- "발굴 — 부업의 90%는 첫 단추에서 망한다" (라벨 + 대시)
- "DPS의 첫 관문, 노력 없이도 돈이 따라오는 자리를 찾는 법" (시스템명 라벨화 + 너무 김)
- "이제 무너지는 게 더 이상 사건이 아니다" (사건이 아니다 ← 말이 안 됨)
- "한 번 흔들려도 다음 날엔 흔적도 없다" (AI식 과장)
- "월수도 시스템의 첫 설계" (의미 불명 + 설계라는 설명체 어휘)
- "MDS 파이프라인" (파이프라인이라는 영어 외래어를 시스템 접미사로)
- "주가 -12% 떨어져도 매도 안 하는 뇌 회로가 박혔다" (뇌 회로가 박히다 = 어법 어색, 비유 남용)
- "신경회로가 새로 깔리고 있다는 신호" 류 (한 번까진 OK, 같은 비유 두 번 X)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
💎 책의 시그니처 기법 (위 고정 규칙을 그대로 따른다)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

기법 이름은 위에서 "{method_name}"(영문 약자+한글 접미사)으로 이미 확정됨. 새로 짓지 말 것.
- 한글로 바꾸지 말고, 다른 표기로 바꾸지 말고, 그대로 사용한다.
- 아래 등장 규칙의 ○○○ 자리에는 전부 "{method_name}" 을 넣는다.
- "{method_name}" 을 글자 그대로 쓰고, 약자를 풀어쓰거나 접미사를 덧붙이지 않는다.

[약자의 의미를 목차에 '간접적으로' 녹여라]
이 기법의 약자 풀이: {method_expansion if method_expansion else "(풀이 없음 — 책 흐름으로 자연스럽게 암시)"}
- 약자가 의미 있는 머리글자라면, 각 단계가 PART 1~5 흐름이나 소제목에 자연스럽게 드러나 독자가 목차만 봐도 "이런 단계의 방법론이구나"를 감 잡게 하라.
- "Y: ~" 식 글자별 사전 나열·풀이 직접 받아쓰기 절대 금지. 호기심 자극 제목 안에 의미가 스며들게만.
- 풀이가 없으면 억지로 짜맞추지 말고 평소대로 호기심 중심으로 쓰되, 책 전체가 하나의 방법론을 단계적으로 다룬다는 느낌만 유지.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🎯 컨셉명 등장 규칙 - 3번, 자연스럽게 분산
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

이 컨셉명을 5개 PART 안에 정확히 3곳에 등장시킨다:

⚠️ 컨셉명 글자 수/단어 수를 단정하는 표현 절대 금지 (가장 자주 어기는 실수!)
   ❌ "○○○ 세 글자에서 시작한다" — 컨셉명이 3글자 아니면 거짓말이 됨
   ❌ "○○○ 네 글자만 기억하라" — 글자 수 단정 금지
   ❌ "단 한 단어가 모든 걸 바꾼다" — 단어 수 단정 금지
   ✅ 컨셉명이 'DPS'(3자)든 '단단한 매출 구조'(8자)든 '복리 자산 공식'(7자)이든 모두 자연스럽게 작동하는 문장만 사용

1. **PART 1 마지막 소제목** (도입) — 아래 풀에서 1개 선택, 컨셉명 글자 수와 무관하게 자연스러운 것
   - 예: "결국 모든 답은 '○○○' 안에 있었다"
   - 예: "이 책의 모든 페이지는 '○○○' 하나를 향해 간다"
   - 예: "지금부터 '○○○' 단 하나만 기억하면 된다"
   - 예: "여기서부터 진짜 이야기, '○○○'가 시작된다"
   - 예: "마지막에 도달하는 곳은 결국 '○○○'다"
   - 예: "'○○○'를 만나기 전과 후는 완전히 다른 게임이다"

2. **PART 3 또는 PART 4 안 (1곳)** (작동·전환점)
   - 예: "○○○가 본격 작동하기 시작하는 4가지 신호"
   - 예: "○○○를 처음 적용한 사람들이 가장 먼저 느낀 변화"
   - 예: "○○○가 통장에 처음 흔적을 남기는 순간"
   - 예: "○○○ 한 달 차에 가장 먼저 무너지는 한 가지"

3. **PART 5 마지막 소제목** (확장/사회증명)
   - 예: "○○○를 익힌 사람들의 5년 뒤가 완전히 다른 이유"
   - 예: "○○○로 자리잡은 사람들이 다시는 돌아가지 않는 이유"
   - 예: "○○○ 이후, 1년 만에 가장 크게 달라지는 단 한 가지"

❌ 챕터 제목에는 컨셉명 절대 등장 X (5개 챕터 제목엔 안 들어감)
❌ 매 PART에 박지 말 것. 정확히 3곳.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
✍️ 소제목 톤: 한국 자기계발 베스트셀러 + 마케팅 뇌과학
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

각 PART당 소제목 4개. 모두 다른 패턴 + 명사·동사 엔딩 섞기 (4개 모두 같은 엔딩 금지).

[좋은 패턴 풀 - 매번 다르게]
1. 통계 충격형 — "1년 안에 99%가 다시 무너진다", "한 달 차에 90%가 다시 무너지는 정확한 이유"
2. 인지부조화형 — "의지로 버틴 사람일수록 더 크게 무너진다"
3. 뇌과학 권위형 — "감정이 폭발할 때 뇌는 정확히 어디부터 멈추는가", "신경회로가 새로 깔리고 있다는 4가지 신호"
4. 임박감/손해회피형 — "이 90초를 놓치면 다시 일어서는 데 며칠이 걸린다"
5. 정체성 전환형 — "한 번 회복한 뇌는 다시는 무너지지 않는다"
6. 호기심 갭형 — "단번에 회복한 사람들이 모두 무의식적으로 하는 행동"
7. 이중/삼중 보상형 — "회복 후 6개월 만에 통장과 인간관계가 함께 변한다"
8. 인그룹 사회증명형 — "○○를 익힌 사람들의 5년 뒤가 완전히 다른 이유"

[★ 가장 중요 — 목차만 보고 결제하게 만들어라 (설명문 금지)]
이 목차의 단 하나의 목적: 서점에서 5초 훑은 사람이 "이건 사야 해"라고 결제하게 만드는 것.
정보 전달은 본문이 한다. 목차는 100% 구매 욕구만 자극한다.

지금 가장 흔한 실패 = "설명문처럼 평이함". 아래 셋 중 하나라도 걸리면 그 줄은 죽은 줄이니 다시 써라:
  (1) 정보 전달형("~하는 법", "~의 중요성", "~란 무엇인가") — 교과서 목차
  (2) 답을 다 말해버림 — 본문을 살 이유가 사라짐
  (3) 자극이 없음 — 심장이 안 뛰면 결제 안 한다

[자청식 후킹 7가지 무기 — PART마다 최소 2개 이상 섞어라]
1. 통념 정면 박살: "열심히 할수록 가난해지는 이유", "절약이 당신을 평생 가난하게 만든다"
2. 자기기만 적발(뜨끔하게): "공부만 하다 1년을 통째로 날리는 사람들의 5가지 변명"
3. 단정 선언(반박 불가 톤): "결국 전부 이거 하나에서 갈렸다", "답은 처음부터 정해져 있었다"
4. 날 선 대비: "버는 사람과 버는 척하는 사람", "3년 버틴 사람과 3개월에 접은 사람의 통장"
5. 충격 숫자: "10명 중 8명이 1년 안에 무너지는 진짜 이유", "상위 3%만 아는 한 가지"
6. 공포·손실 회피: "지금 이걸 모르면 5년 뒤 똑같은 자리에 있다", "당신이 놓치는 사이 벌어지는 일"
7. 미스터리 갭(정체를 가림): "퇴사 3개월 만에 다시 돌아온 회사원이 깨달은 단 하나"

[좋은 변환 — 평이 → 자극, 어법은 자연스럽게]
- "배당주 고르는 법" → "10년 배당을 받고도 한 푼도 못 쓴 사람들의 공통점"
- "복리의 중요성" → "같은 돈을 넣었는데 7년 뒤 잔고가 두 배로 갈린 이유"
- "감정 관리가 필요하다" → "계좌를 들여다본 횟수가 수익률을 갉아먹은 증거"
- "분산 투자를 하자" → "한 종목에 몰빵한 사람이 그해 오히려 덜 잃은 까닭"
- "꾸준함이 답이다" → "3년을 버틴 사람과 3개월에 그만둔 사람의 통장이 비슷했던 이유"

[규칙]
- 소제목 절반 이상에 구체 숫자(금액·기간·비율·인원)를 박는다.
- 통념을 뒤집거나 자기기만을 찌르는 줄을 PART마다 최소 1개.
- 답을 같은 줄에서 다 말하지 마라. "왜/어떻게/무엇이"의 정체는 본문에 숨긴다.
- 사람을 모욕하지 마라(욕설·인신공격 금지). 통념과 '행동'을 때리되 독자를 적으로 만들지 않는다.
- 자청·역행자 등 특정 작가의 '단어'는 쓰지 말 것. 톤만 가져온다.
- 자극을 위해 억지 비유나 말 안 되는 조합을 만들지 마라. 아래 [자연스러움 원칙]이 항상 우선.

[자연스러움 원칙 - 절대 어기지 말 것]
⚠️ 어법이 1순위다. 호기심보다 "말이 되는가"가 먼저다. 어법이 어색하면 호기심을 줄여서라도 자연스러운 문장으로 바꾼다.
- 모든 소제목/챕터 제목은 한국어 원어민 편집자가 손대지 않고 통과시킬, 문법적으로 완결된 자연스러운 문장이어야 한다.
- 단어를 억지로 조합한 "그럴듯해 보이지만 뜻이 안 통하는" 표현은 즉시 폐기. 특히 은유·비유를 무리하게 끼워 맞추지 마라.
  ❌ "배당 사이 파이프 굵기 차이는 정확히 얼마인가" (파이프 굵기? 배당 사이? — 무슨 말인지 알 수 없음)
  ❌ "수익이 자라는 토양의 산도를 맞추는 법" (억지 비유, 어법 어색)
  ✅ "같은 배당주를 사도 누구는 월세처럼 받고 누구는 못 받는 이유" (자연스럽고 호기심도 유발)
- 어법 검사: 주어와 서술어가 자연스럽게 연결되는가? ("회복이 굴러간다" X — 회복은 굴러가지 않음)
- 추상 개념 + 기계·물리·생물 동사 금지: 계좌·통장·재투자·수익·복리·멘탈·습관·시스템 같은 추상/사물에 "켜다/끄다/돌리다/감다/조이다/풀다/꽂다/심다/굴린다/얹는다/멈춘다/자란다/살아난다/숨쉰다/깨어난다" 같은 동사를 붙이지 마라. (계좌는 자라지 않고, 시스템은 살아나지 않는다)
  ❌ "재투자를 켠 사람과 끈 사람" / "수익률 화면을 끄니 계좌가 살아났다" / "내버려 두면 계좌가 자란다"
  ✅ "수익을 다시 넣은 사람과 빼서 쓴 사람의 3년 뒤" / "계좌를 덜 들여다본 해에 수익률이 더 높았던 이유" (대비·호기심은 살리되 주어-서술어가 말이 되게)
- 과장 형용사 금지: "흔적도 없다", "통째로", "완전히" (꼭 필요할 때만)
- 추상 X, 구체 O: "많은 사람" → "월급 280만원짜리 7년차 회사원"
- 도구/플랫폼명 적극: 네이버, 카카오, 노션, 카톡, 캘린더, 구글 시트
- 출력 직전, 모든 줄을 소리내어 읽어라. 원어민이 한 번에 이해 못 하거나 "이게 무슨 말이지?" 싶은 줄은 전부 다시 쓴다.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🚫 즉시 폐기 표현
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

특정 작가 고유어 (절대 금지): 자청 / 역행자 / 자의식 해체 / 유전자 역행 / 원시인 / 추월차선 / 아토믹 해빗 / 언카피어블
AI 클리셰: 졸업 / 정체 / 마지막 한 수 / 다른 차원 / 결정적 시그널 / 진짜 게임 / 흔적도 없다 / 사건이 아니다 / 회로가 박혔다
시스템 의인화: "○○가 멈춘 날", "○○ 위에 얹다", "○○를 졸업한", "회복이 굴러가다"
뇌과학 비유 남용: "뇌 회로가 박혔다", "뇌 회로가 새로 깔린다" (전체 목차에 뇌·신경회로는 사실 진술로 1~2회만, 비유 남용 X)
밍밍: 효과적인 / 성공적인 / ~의 모든 것 / ~하는 방법 / 알아야 할 / 의 중요성
유치 비유: 나침반 / 열쇠 / 보물 / 황금 / 마법 / 파이프라인 / 엔진 / 톱니바퀴 / 사이클 / 눈덩이
참고서: 첫걸음 / 완벽가이드 / 핵심정리 / 기초/중급/고급 / 첫 설계 / 첫 셋업
챕터 제목 라벨: "발굴 —", "1단계:", "STEP 1." 같은 분리 형식
의문문 문어체: "왜 ~는 ~하지 못하는가" 식의 한 PART에 1개까지만
콜론(:) — 단 한 번도 쓰지 마라
숫자 중복 금지: 전체 목차에서 같은 숫자(예: 3시간 + 3개월) 두 번 등장 금지
숫자 표기 - 부호 금지: "-12%" "+30%" 같은 부호 사용 X. "12% 폭락에도", "30% 상승하면" 식으로

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📝 출력 형식 (이 형식 외 어떤 텍스트도 출력 금지)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[책 시그니처 컨셉]
○○○ + (시스템/구조/공식/알고리즘) | 한 줄 설명

PART 1. [13~18자, 좌절 + 통계 충격]
- [통계 충격 또는 결정적 함정형]
- [인지부조화/통념 박살형]
- [패턴 - 위 8개 중 다른 것]
- [컨셉명 첫 등장: 자연스러운 도입]

PART 2. [13~18자, 통념 박살 + 뇌과학 권위]
- [패턴]
- [패턴 - 다른 것]
- [패턴 - 또 다른 것]
- [패턴 - 또 다른 것]

PART 3. [13~18자, 첫 사건/결정적 순간]
- [패턴]
- [임박감/손해회피형]
- [패턴]
- [컨셉명 등장 가능: "○○○를 처음 적용한 사람들이 가장 먼저 느낀 변화" — PART 4에 넣을 거면 여기는 일반 패턴]

PART 4. [13~18자, 정체성 전환 선언]
- [컨셉명 등장 가능: "○○○가 본격 작동하는 4가지 신호" — PART 3에 안 넣었다면 여기에]
- [뇌과학 권위형 또는 통계형]
- [패턴]
- [패턴]

PART 5. [13~18자, 도약/이중 보상]
- [호기심 갭형]
- [이중/삼중 보상형]
- [패턴 - 또 다른 것]
- [컨셉명 세 번째 등장: 인그룹 사회증명]

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔍 출력 전 자가 점검 (반드시 통과)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

출력 직전에 5가지 모두 통과해야 한다. 하나라도 No면 다시 써라.

체크 1. 5개 챕터 제목만 이어 읽었을 때 "X였던 사람이 Y로 바뀐다"는 변화 서사가 또렷한가?
체크 2. 통념 박살(인지부조화) 패턴이 5개 챕터+20개 소제목 안에 3개 이상 박혀있는가?
체크 3. 구체적 숫자(시간/금액/비율)가 8개 이상 등장하는가?
체크 4. 평이한 표현("~의 방법", "~의 모든 것", "~의 중요성", "효과적인", "성공적인")이 0개인가?
체크 5. 5초간 훑어본 가상 독자가 "이건 안 사면 손해"라고 느낄 만한 손실회피 트리거가 챕터 제목 5개 중 2개 이상에 있는가?
체크 6. 호기심 갭(결과만 보이고 방법/이유는 숨김)이 20개 소제목 중 10개 이상에 박혀 있는가? — "결과만 보이는데 본문을 사야 알 수 있는 한 줄"이 절반 이상이어야 결제 전환됨.
체크 7. 컨셉명 글자 수를 단정하는 표현("세 글자", "네 글자", "단 한 단어")이 단 하나도 없는가? — 하나라도 있으면 즉시 전체 다시 쓰기.
체크 8. (어법 - 가장 중요) 모든 챕터 제목·소제목을 소리내어 읽었을 때, 원어민이 한 번에 이해되고 어법이 자연스러운가? "배당 사이 파이프 굵기 차이는…" 같은 억지 조합·뜻 모를 비유가 단 하나라도 있으면 그 줄을 자연스럽게 다시 써라.
체크 9. (어법) 추상 개념에 기계·물리 동사를 붙인 줄("재투자를 켠/끈", "수익을 돌린다" 류)이 하나도 없는가? 있으면 즉시 자연스러운 표현으로 교체.
체크 10. (자극) 통념을 정면으로 뒤집거나 자기기만을 찌르는 줄이 PART마다 최소 1개 있는가? 전부 정보 전달형이면 평이한 것 — 다시 써라.

목차만 출력. 콜론 금지. 매 소제목 다른 패턴. 명사·동사 엔딩 섞기. 어법 어색한 표현 즉시 폐기. 각 PART는 정확히 소제목 4개. 컨셉명 글자수 단정 금지. 자가점검 결과는 출력하지 말 것."""
    return ask_ai(prompt, 0.85, ensure_quality=True)


def generate_content_premium(subtopic, chapter, questions, answers, topic, persona):
    """자기계발 후킹 + 권석천 칼럼 깊이의 몰입형 본문"""
    import random

    qa_pairs = ""
    for i, (q, a) in enumerate(zip(questions, answers), 1):
        if a.strip():
            qa_pairs += f"\n질문{i}: {q}\n답변{i}: {a}\n"

    # 다양한 시작 스타일
    hook_styles = [
        "도발적 질문으로 시작 (예: '왜 99%는 여기서 실패할까요?')",
        "충격적 고백으로 시작 (예: '저도 2년간 완전히 틀리고 있었습니다.')",
        "반전 사실로 시작 (예: '사실 정반대였습니다.')",
        "구체적 숫자로 시작 (예: '정확히 23일 만에 달라졌습니다.')",
        "생생한 에피소드로 시작 (예: '그날 밤 컴퓨터 앞에서 깨달았습니다.')",
        "단호한 선언으로 시작 (예: '핵심부터 말씀드리겠습니다.')",
        "대화체로 시작 (예: '\"진짜요?\" 처음 들었을 때 저도 의심했습니다.')",
        "before/after로 시작 (예: '1년 전의 저는 완전히 다른 상황이었습니다.')",
        "상식 뒤집기로 시작 (예: '열심히 하면 된다? 틀렸습니다.')",
        "사건 장면으로 시작 (예: '2023년 3월의 일이었습니다.')",
        "인용으로 시작 (예: '한 후배가 이런 말을 했습니다.')",
    ]
    current_hook = random.choice(hook_styles)

    # 표 포함 여부 (랜덤하게 30% 확률)
    include_table = random.random() < 0.3

    prompt = f"""당신은 한국 자기계발 베스트셀러 작가입니다. '{subtopic}'에 대해 첫 문장으로 독자를 멈춰 세우고, 권석천 기자 칼럼처럼 정확한 디테일과 통찰로 끝까지 끌고 가는 본문을 씁니다.

[주제]: {topic}
[챕터]: {chapter}
[참고 내용]
{qa_pairs}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔥 첫 문장이 전부다
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

반드시 이 스타일로 시작:
👉 {current_hook}

(일반론·교훈으로 시작하면 즉시 폐기. 사건/숫자/대사/의문 중 하나로만 시작)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
✍️ 본문 톤: 자기계발 후킹 + 권석천 칼럼 깊이
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[문체]
- 합쇼체 기본 ("~입니다", "~합니다") + 가끔 구어체 ("~거든요", "~더라고요")
- 현재 시제로 장면을 그리듯
- 짧은 문장과 긴 문장을 교차해 리듬
- 추상보다 구체. "많은 사람" → "월급 280만원짜리 7년차 회사원"

[권석천식 깊이]
- 사실 → 분석 → 통찰 순서로 전개
- 가설 검증식 흐름: "왜 그럴까. 이유를 되짚어봤습니다"
- 사회적 맥락이나 통계, 책/논문 인용을 자연스럽게 섞기
- 결론 강요 X, 독자가 스스로 깨닫게 단서를 깔기

[자기계발식 후킹]
- 첫 문장: 사건/숫자/대사/의문
- 본문 중간: 작은 반전 1회 ("그런데 진짜 흥미로운 건 그 다음이었습니다")
- 마지막 문장: 발견의 결과로서의 통찰 한 줄

[표 포함: {'예' if include_table else '아니오'}]
{'''
📊 본문 중간에 비교표 1개 필수:
<table style="width:100%; border-collapse:collapse; margin:20px 0;">
<tr style="background:#1a1a1a;"><th style="border:1px solid #333;padding:12px;color:#C9A24B;">구분</th><th style="border:1px solid #333;padding:12px;color:#C9A24B;">기존 방식</th><th style="border:1px solid #333;padding:12px;color:#C9A24B;">새로운 방식</th></tr>
<tr><td style="border:1px solid #333;padding:10px;">항목</td><td style="border:1px solid #333;padding:10px;">내용</td><td style="border:1px solid #333;padding:10px;">내용</td></tr>
</table>
''' if include_table else ''}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🎯 구체성을 끝까지
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

- 숫자: "많이" X → "월 340만원, 정확히 47일" O
- 사례: 수강생/지인/현장 (이름은 가끔만)
- 실제 도구명: 네이버, 카카오, 노션, 구글 시트 등

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🚫 절대 금지
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

❌ 자청, 역행자, 자의식 해체, 유전자 역행, 원시인, 추월차선, 아토믹 해빗 (특정 작가/책 고유 표현)
❌ 유치한 비유: 마법, 황금열쇠, 나침반, 로켓, 눈덩이, 톱니바퀴, 파이프라인
❌ AI 어휘: 중요합니다, 따라서, 결론적으로, ~를 통해, 다양한, 효과적인, 진정한
❌ 직접 호칭: 여러분, 당신, 독자님
❌ 형식: 1. 2. 첫째, 글머리 기호, 이모지
❌ 같은 이름 반복 (민준, 지수가 계속 X)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔥 몰입·충격 어조 (자청식 몰입감 — 특정 작가 고유어는 제외)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

- 도입부터 통념을 정면으로 깬다. 독자가 "어, 내가 알던 거랑 반대네" 하고 멈칫하게.
- 글 중간마다 '작은 충격'을 심어라: 예상과 반대되는 사실, 뒤집히는 데이터, 의외의 결론.
- 한 호흡에 읽히는 리듬 — 짧은 단정문으로 치고, 긴 문장으로 풀고, 다시 끊는다.
- "왜?"를 계속 만들어 다음 문단을 안 읽고는 못 배기게 한다(궁금증 갭).
- 충격 뒤엔 반드시 '그래서 무엇을 어떻게'의 실질이 온다. 겁만 주고 끝내지 않는다.
- 자청·역행자 등 특정 작가의 '단어'는 절대 쓰지 않는다. 몰입감과 태도만 가져온다.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🎓 전문성 (아마추어 글과 가르는 지점)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

- 추상 주장 1개당 근거 1개(데이터·연구·구체 사례·숫자)를 반드시 붙인다.
- 메커니즘을 설명하라. "왜 그렇게 되는가"의 원리와 과정을 단계로 풀어라.
- 바로 따라 할 수 있게: 무엇을·어디서·어떤 순서로·얼마나, 구체적으로.
- 흔한 조언의 한계를 짚고, 더 정확한 기준·예외를 제시한다(깊이의 증거).

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📏 분량: 2400~3000자 (최소 2400자 이상, 전문성 있는 밀도로)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

⚠️ 길이는 사례·데이터·단계별 설명·메커니즘으로 채운다. 같은 말 반복·미사여구로 늘리면 즉시 실패.

'{subtopic}' 본문 작성.
- 시작: {current_hook}
- 권석천 칼럼처럼 사실 → 분석 → 통찰 순서로 인과 추적
- 마지막 한 줄에 발견된 통찰 하나
- {'비교표 1개 포함' if include_table else '순수 텍스트만'}

⛔ 절대 금지: 본문 첫 줄에 소제목('{subtopic}')을 다시 쓰지 마라.
   소제목은 위에 이미 표시되므로, 본문은 곧장 첫 후킹 문장으로 시작한다.
   ❌ 잘못: "{subtopic}\\n\\n그날 새벽 두 시였습니다..." (소제목 반복)
   ✅ 올바름: "그날 새벽 두 시였습니다..." (바로 본문 시작)"""
    return ask_ai(prompt, 0.75, ensure_quality=True)


def format_content_html(content):
    """본문을 HTML 형식으로 변환 (강조 표시 적용)"""
    if not content:
        return ""
    # 「」 → 주황색 볼드
    formatted = re.sub(r'「([^」]+)」', r'<b style="color:#e67e22;">\1</b>', content)
    # ★ → 주황색 볼드 문장
    formatted = re.sub(r'★\s*(.+?)(?=\n|$)', r'<p style="color:#e67e22;font-weight:700;margin:20px 0;font-size:17px;">★ \1</p>', formatted)
    # 문단 구분 (빈 줄) → 문단 간격
    formatted = formatted.replace('\n\n', '</p><p style="color:#000000;margin:25px 0;line-height:2.0;font-size:17px;">')
    # 단일 줄바꿈 제거 (문단 내 연결)
    formatted = formatted.replace('\n', ' ')
    formatted = f'<p style="color:#000000;margin:25px 0;line-height:2.0;font-size:17px;">{formatted}</p>'
    # 빈 <p> 태그 정리
    formatted = re.sub(r'<p[^>]*>\s*</p>', '', formatted)
    return formatted


def generate_questions(subtopic, chapter, topic):
    prompt = f"""'{topic}' 전자책 '{chapter}' 챕터의 '{subtopic}' 작성용 질문 3개:

Q1: [질문]
Q2: [질문]
Q3: [질문]"""
    return ask_ai(prompt, 0.7)


# ==========================================
# 메인 UI
# ==========================================
# 비디오 배경 헤더 — GitHub/Streamlit 배포 시에도 사용자 제공 영상을 자동 탐색
HEADER_VIDEO_FILENAME = "13649905_1920_1080_30fps.mp4"
HEADER_VIDEO_NAMES = {
    "webm": ["title_bg.webm", "writey_title_bg.webm", "hero_bg.webm"],
    "mp4": [HEADER_VIDEO_FILENAME, "title_bg.mp4", "writey_title_bg.mp4", "hero_bg.mp4"],
}
HEADER_VIDEO_DIRS = ("", "assets", "static", "media", "videos")


def _candidate_header_video_dirs():
    bases = []
    try:
        app_dir = Path(__file__).resolve().parent
        bases.extend(app_dir / name if name else app_dir for name in HEADER_VIDEO_DIRS)
    except Exception:
        pass

    try:
        cwd = Path.cwd()
        bases.extend(cwd / name if name else cwd for name in HEADER_VIDEO_DIRS)
    except Exception:
        pass

    # 로컬 미리보기용 fallback. 배포 시에는 앱 폴더 또는 assets/ 폴더에 올리면 됩니다.
    bases.extend([
        Path.home() / "Downloads",
        Path.home() / "Desktop",
        Path("/Users/hyunwoo/Downloads"),
        Path("/Users/hyunwoo/Desktop"),
    ])

    unique = []
    seen = set()
    for base in bases:
        key = str(base)
        if key not in seen:
            seen.add(key)
            unique.append(base)
    return unique


def _find_header_video(exts=("mp4", "webm")):
    found = {}
    for ext in exts:
        for base in _candidate_header_video_dirs():
            for name in HEADER_VIDEO_NAMES.get(ext, [f"title_bg.{ext}"]):
                p = base / name
                if p.exists():
                    found[ext] = str(p)
                    break
            if ext in found:
                break
    return found

_header_videos = _find_header_video()
header_video_webm_b64 = get_video_base64(_header_videos["webm"]) if "webm" in _header_videos else None
header_video_mp4_b64 = get_video_base64(_header_videos["mp4"]) if "mp4" in _header_videos else None
header_video_b64 = header_video_mp4_b64 or header_video_webm_b64  # 이번에 지정한 mp4 영상을 최우선 사용

if header_video_b64:
    st.markdown("""
    <div class="writey-brandbar">
        <div class="writey-brand-left">
            <span class="writey-wordmark">WRITEY</span>
            <span class="writey-cashtag">CASHMAKER STUDIO</span>
        </div>
        <span class="writey-author">남현우 작가</span>
    </div>
    """, unsafe_allow_html=True)

    # 영상 히어로는 components.html(iframe)로 렌더 → 큰 base64 data URI도 정상 재생
    _sources = ""
    if header_video_mp4_b64:
        _sources += f'<source src="data:video/mp4;base64,{header_video_mp4_b64}" type="video/mp4">'
    if header_video_webm_b64:
        _sources += f'<source src="data:video/webm;base64,{header_video_webm_b64}" type="video/webm">'
    components.html(f"""
    <!DOCTYPE html><html><head><meta charset="utf-8">
    <style>
        @import url('https://cdn.jsdelivr.net/gh/orioncactus/pretendard/dist/web/static/pretendard.css');
        html,body {{ margin:0; padding:0; background:transparent; overflow:hidden; }}
        .hero {{
            position:relative; width:100%; height:430px;
            border-radius:10px; overflow:hidden;
            border:1px solid rgba(240,213,139,0.28);
            box-shadow:0 34px 90px rgba(0,0,0,0.58), inset 0 1px 0 rgba(255,255,255,0.14);
            font-family:'Pretendard','Apple SD Gothic Neo',sans-serif;
            background:#07080b;
            isolation:isolate;
        }}
        .hero video {{
            position:absolute; top:0; left:0; width:100%; height:100%;
            object-fit:cover; filter:brightness(0.68) saturate(1.22) contrast(1.12);
            transform:scale(1.035);
        }}
        .hero .veil {{
            position:absolute; inset:0;
            background:
                radial-gradient(circle at 18% 22%, rgba(240,213,139,0.18), transparent 32%),
                radial-gradient(circle at 76% 10%, rgba(123,211,200,0.12), transparent 30%),
                linear-gradient(90deg, rgba(5,6,8,0.86) 0%, rgba(5,6,8,0.52) 47%, rgba(5,6,8,0.88) 100%),
                linear-gradient(180deg, rgba(5,6,8,0.16) 0%, rgba(5,6,8,0.72) 100%);
            z-index:1;
        }}
        .hero .grain {{
            position:absolute; inset:0; z-index:2; opacity:.28; mix-blend-mode:soft-light;
            background-image:
                linear-gradient(rgba(255,255,255,.05) 1px, transparent 1px),
                linear-gradient(90deg, rgba(255,255,255,.035) 1px, transparent 1px);
            background-size:42px 42px;
        }}
        .hero .shine {{
            position:absolute; inset:-40% -10% auto auto; width:58%; height:140%;
            background:linear-gradient(120deg, transparent 0%, rgba(255,255,255,0.18) 46%, transparent 68%);
            transform:rotate(13deg); z-index:3; opacity:.5;
        }}
        .hero .center {{
            position:absolute; inset:0; z-index:4; display:flex; flex-direction:column;
            justify-content:center; align-items:flex-start; text-align:left;
            padding:58px 64px;
        }}
        .hero .panel {{
            max-width:720px;
        }}
        .hero .eyebrow-row {{
            display:flex; align-items:center; gap:12px; margin-bottom:18px;
        }}
        .hero .status-dot {{
            width:9px; height:9px; border-radius:99px; background:#7bd3c8;
            box-shadow:0 0 24px rgba(123,211,200,.92);
        }}
        .eyebrow {{
            color:#8fe3d7; font-size:12px; letter-spacing:.2em;
            text-transform:uppercase; font-weight:850;
        }}
        .title {{
            font-size:94px; font-weight:920; letter-spacing:-0.045em;
            color:#fff6e6; margin:0; line-height:.9;
            text-shadow:0 18px 55px rgba(0,0,0,.52);
        }}
        .title span {{
            color:#d7b86a;
        }}
        .subhead {{
            color:#efe1c3; font-size:22px; margin:24px 0 0;
            font-weight:650; letter-spacing:-.01em;
        }}
        .tagline {{
            color:rgba(245,241,232,.74); font-size:15px; margin:10px 0 0;
            font-weight:500; letter-spacing:0; max-width:560px; line-height:1.7;
        }}
        .hero .hud {{
            position:absolute; right:34px; bottom:32px; z-index:5;
            display:grid; grid-template-columns:repeat(3, minmax(0, 1fr)); gap:10px;
            width:min(520px, calc(100% - 68px));
        }}
        .hero .hud-card {{
            border:1px solid rgba(255,255,255,.16);
            background:rgba(8,9,12,.46);
            backdrop-filter:blur(14px) saturate(1.2);
            border-radius:8px; padding:13px 14px;
            box-shadow:inset 0 1px 0 rgba(255,255,255,.08);
        }}
        .hero .hud-label {{
            color:rgba(245,241,232,.54); font-size:10px; letter-spacing:.13em;
            text-transform:uppercase; font-weight:800; margin-bottom:3px;
        }}
        .hero .hud-value {{
            color:#fff6e6; font-size:18px; font-weight:850;
        }}
        .hero .bottom-line {{
            position:absolute; left:64px; bottom:43px; width:130px; height:2px; z-index:5;
            background:linear-gradient(90deg,#d7b86a,rgba(215,184,106,0));
        }}
        @media (max-width: 760px) {{
            .hero {{ height:500px; }}
            .hero .center {{ padding:42px 28px; justify-content:flex-start; }}
            .title {{ font-size:62px; }}
            .subhead {{ font-size:18px; }}
            .hero .hud {{ left:28px; right:28px; bottom:28px; width:auto; grid-template-columns:1fr; }}
            .hero .bottom-line {{ left:28px; bottom:196px; }}
        }}
    </style></head>
    <body>
        <div class="hero">
            <video autoplay muted loop playsinline preload="auto">
                {_sources}
            </video>
            <div class="veil"></div>
            <div class="grain"></div>
            <div class="shine"></div>
            <div class="center">
                <div class="panel">
                    <div class="eyebrow-row"><span class="status-dot"></span><div class="eyebrow">Premium E-Book Operating System</div></div>
                    <h1 class="title">WRITE<span>Y</span></h1>
                    <p class="subhead">기획부터 종이책급 출력까지, 한 화면에서 완성합니다</p>
                    <p class="tagline">인터뷰, 시장 분석, 목차 설계, 본문 작성, 표지와 WORD 내보내기를 하나의 프리미엄 제작 흐름으로 묶었습니다.</p>
                </div>
            </div>
            <div class="bottom-line"></div>
            <div class="hud">
                <div class="hud-card"><div class="hud-label">Mode</div><div class="hud-value">Studio</div></div>
                <div class="hud-card"><div class="hud-label">Output</div><div class="hud-value">DOCX</div></div>
                <div class="hud-card"><div class="hud-label">Cover</div><div class="hud-value">Premium</div></div>
            </div>
        </div>
    </body></html>
    """, height=452)
else:
    st.markdown("""
    <div class="writey-brandbar">
        <div class="writey-brand-left">
            <span class="writey-wordmark">WRITEY</span>
            <span class="writey-cashtag">CASHMAKER</span>
        </div>
        <span class="writey-author">남현우 작가</span>
    </div>
    <div class="interview-hero">
        <div class="eyebrow">AI EBOOK Writer</div>
        <h1 class="title">WRITEY</h1>
        <div class="divider"></div>
        <p class="tagline">6개의 질문에 답하면 AI가 목차부터 본문까지 완성합니다</p>
    </div>
    """, unsafe_allow_html=True)

# ==========================================
# 인터뷰 모드 (interview_completed가 False일 때)
# ==========================================
if not st.session_state.get('interview_completed', False):
    st.markdown("""
    <div class="section-title-box">
        <span class="section-step">INTERVIEW</span>
        <h2>나만의 전자책 만들기</h2>
        <p>몇 가지 질문에 답하면 AI가 전자책을 완성해드립니다</p>
    </div>
    """, unsafe_allow_html=True)

    # 인터뷰 진행 상태
    if 'interview_step' not in st.session_state:
        st.session_state['interview_step'] = 1

    step = st.session_state['interview_step']
    total_steps = 6

    # 진행률 표시
    st.progress(step / total_steps)
    st.caption(f"질문 {step} / {total_steps}")

    st.markdown("---")

    # 인터뷰 데이터 임시 저장
    if 'temp_interview' not in st.session_state:
        st.session_state['temp_interview'] = {}

    # ========== STEP 1: 기본 정보 ==========
    if step == 1:
        st.markdown("""
        <div style="background:linear-gradient(135deg, rgba(201,162,75,0.1) 0%, rgba(201,162,75,0.05) 100%);padding:30px;border-radius:15px;border-left:4px solid var(--gold);margin-bottom:30px;">
            <h3 style="color:var(--gold);margin:0 0 10px 0;">👋 먼저 당신에 대해 알려주세요</h3>
            <p style="color:var(--text);margin:0;opacity:0.9;">전자책의 저자로서 기본 정보를 입력해주세요</p>
        </div>
        """, unsafe_allow_html=True)

        with st.form(key="step1_form"):
            author_name = st.text_input("저자명 (필명 가능)", value=st.session_state['temp_interview'].get('author_name', ''), placeholder="예: 김성장, 머니메이커 등")
            field = st.text_input("당신의 전문 분야는?", value=st.session_state['temp_interview'].get('field', ''), placeholder="예: 주식투자, 블로그 수익화, 다이어트, 영어회화 등")

            exp_options = ["선택하세요", "1년 미만", "1~2년", "3~5년", "5~10년", "10년 이상"]
            saved_exp = st.session_state['temp_interview'].get('experience_years', '선택하세요')
            exp_index = exp_options.index(saved_exp) if saved_exp in exp_options else 0
            experience = st.selectbox("이 분야 경험은?", exp_options, index=exp_index)

            col1, col2 = st.columns([1, 1])
            with col2:
                submitted = st.form_submit_button("다음 →", use_container_width=True, type="primary")

            if submitted:
                if not author_name.strip() or not field.strip() or experience == "선택하세요":
                    st.error("모든 항목을 입력해주세요")
                else:
                    st.session_state['temp_interview']['author_name'] = author_name.strip()
                    st.session_state['temp_interview']['field'] = field.strip()
                    st.session_state['temp_interview']['experience_years'] = experience
                    st.session_state['interview_step'] = 2
                    st.rerun()

    # ========== STEP 2: 주제와 노하우 ==========
    elif step == 2:
        st.markdown("""
        <div style="background:linear-gradient(135deg, rgba(201,162,75,0.1) 0%, rgba(201,162,75,0.05) 100%);padding:30px;border-radius:15px;border-left:4px solid var(--gold);margin-bottom:30px;">
            <h3 style="color:var(--gold);margin:0 0 10px 0;">📚 어떤 내용을 담을까요?</h3>
            <p style="color:var(--text);margin:0;opacity:0.9;">당신만의 핵심 노하우를 알려주세요</p>
        </div>
        """, unsafe_allow_html=True)

        col_prev, col_next = st.columns([1, 1])
        with col_prev:
            if st.button("← 이전", key="interview_prev_2", use_container_width=True):
                st.session_state['interview_step'] = 1
                st.rerun()

        with st.form(key="step2_form"):
            topic = st.text_input("전자책 주제", value=st.session_state['temp_interview'].get('topic', ''), placeholder="예: 월 100만원 배당 투자, 하루 1시간 블로그로 월 300 벌기")
            core_method = st.text_area("당신만의 핵심 방법/노하우는?", value=st.session_state['temp_interview'].get('core_method', ''), height=120, placeholder="예: 저는 고배당 ETF를 활용해서 안정적으로 수익을 내는 방법을 알려드립니다. 핵심은 분산투자와 복리의 마법입니다...")

            col1, col2 = st.columns([1, 1])
            with col2:
                submitted = st.form_submit_button("다음 →", use_container_width=True, type="primary")

            if submitted:
                if not topic.strip() or not core_method.strip():
                    st.error("모든 항목을 입력해주세요")
                else:
                    st.session_state['temp_interview']['topic'] = topic.strip()
                    st.session_state['temp_interview']['core_method'] = core_method.strip()
                    st.session_state['interview_step'] = 3
                    st.rerun()

    # ========== STEP 3: 타겟 독자 (AI 추천) ==========
    elif step == 3:
        st.markdown("""
        <div style="background:linear-gradient(135deg, rgba(201,162,75,0.1) 0%, rgba(201,162,75,0.05) 100%);padding:30px;border-radius:15px;border-left:4px solid var(--gold);margin-bottom:30px;">
            <h3 style="color:var(--gold);margin:0 0 10px 0;">🎯 누구를 위한 책인가요?</h3>
            <p style="color:var(--text);margin:0;opacity:0.9;">AI가 시장 데이터를 분석해 최적의 타겟을 추천해드립니다</p>
        </div>
        """, unsafe_allow_html=True)

        topic = st.session_state['temp_interview'].get('topic', '')

        # AI 타겟 분석 (캐시)
        if 'ai_target_suggestions' not in st.session_state or st.session_state.get('ai_target_topic') != topic:
            if st.button("🔍 AI 타겟 분석 시작", key="analyze_target", use_container_width=True, type="primary"):
                with st.spinner("시장 데이터 분석 중..."):
                    result = suggest_targets(topic)
                    parsed = parse_json(result)
                    if parsed and parsed.get('personas'):
                        st.session_state['ai_target_suggestions'] = parsed['personas']
                        st.session_state['ai_target_topic'] = topic
                        st.rerun()
                    else:
                        st.error("분석 실패. 다시 시도해주세요.")

        # AI 추천 결과 표시
        if st.session_state.get('ai_target_suggestions'):
            st.markdown("### 📊 AI 추천 타겟")
            personas = st.session_state['ai_target_suggestions']

            selected_idx = st.session_state.get('selected_target_idx', 0)

            for idx, persona in enumerate(personas[:3]):
                is_selected = (idx == selected_idx)
                border_color = "var(--gold)" if is_selected else "var(--line)"
                bg_color = "rgba(201,162,75,0.1)" if is_selected else "rgba(20,20,20,0.5)"

                pain_list = persona.get('pain_points', [])[:3]
                pains_text = " / ".join(pain_list) if pain_list else "고민 분석 중..."

                st.markdown(f"""
                <div style="background:{bg_color};border:1px solid {border_color};border-radius:10px;padding:15px;margin-bottom:10px;">
                    <div style="font-weight:bold;color:var(--gold);margin-bottom:5px;">{persona.get('name', '타겟')}</div>
                    <div style="font-size:13px;color:var(--text2);margin-bottom:8px;">{persona.get('demographics', '')}</div>
                    <div style="font-size:12px;color:var(--text);opacity:0.8;">💭 {pains_text}</div>
                </div>
                """, unsafe_allow_html=True)

                if st.button(f"✓ 이 타겟 선택" if not is_selected else "✓ 선택됨", key=f"select_target_{idx}", use_container_width=True, disabled=is_selected):
                    st.session_state['selected_target_idx'] = idx
                    st.rerun()

            st.markdown("---")

            # 선택된 타겟 정보 자동 입력
            selected_persona = personas[selected_idx] if selected_idx < len(personas) else personas[0]
            default_reader = f"{selected_persona.get('name', '')} ({selected_persona.get('demographics', '')})"
            default_problem = " ".join(selected_persona.get('pain_points', [])[:3])

            st.markdown("##### 선택된 타겟 정보 (수정 가능)")
            target_reader = st.text_input("타겟 독자", value=st.session_state['temp_interview'].get('target_reader', '') or default_reader, key="target_reader_input")
            target_problem = st.text_area("이 독자들의 가장 큰 고민/문제", value=st.session_state['temp_interview'].get('target_problem', '') or default_problem, height=80, key="target_problem_input")

            col1, col2 = st.columns([1, 1])
            with col1:
                if st.button("← 이전", key="interview_prev_3", use_container_width=True):
                    st.session_state['interview_step'] = 2
                    st.rerun()
            with col2:
                if st.button("다음 →", key="interview_next_3", use_container_width=True, type="primary"):
                    if not target_reader or not target_problem:
                        st.error("타겟 독자와 고민을 입력해주세요")
                    else:
                        st.session_state['temp_interview']['target_reader'] = target_reader
                        st.session_state['temp_interview']['target_problem'] = target_problem
                        st.session_state['interview_step'] = 4
                        st.rerun()
        else:
            # AI 분석 전 직접 입력 옵션
            st.markdown("---")
            st.markdown("##### 또는 직접 입력")
            target_reader = st.text_input("타겟 독자", value=st.session_state['temp_interview'].get('target_reader', ''), placeholder="예: 30대 직장인, 투자 초보자")
            target_problem = st.text_area("이 독자들의 가장 큰 고민/문제는?", value=st.session_state['temp_interview'].get('target_problem', ''), height=80, placeholder="예: 월급만으로는 부족하고, 어디서부터 시작해야 할지 모르겠다...")

            col1, col2 = st.columns([1, 1])
            with col1:
                if st.button("← 이전", key="interview_prev_3_manual", use_container_width=True):
                    st.session_state['interview_step'] = 2
                    st.rerun()
            with col2:
                if st.button("다음 →", key="interview_next_3_manual", use_container_width=True, type="primary"):
                    if not target_reader or not target_problem:
                        st.error("타겟 독자와 고민을 입력해주세요")
                    else:
                        st.session_state['temp_interview']['target_reader'] = target_reader
                        st.session_state['temp_interview']['target_problem'] = target_problem
                        st.session_state['interview_step'] = 4
                        st.rerun()

    # ========== STEP 4: 스토리 & 경력 ==========
    elif step == 4:
        st.markdown("""
        <div style="background:linear-gradient(135deg, rgba(201,162,75,0.1) 0%, rgba(201,162,75,0.05) 100%);padding:30px;border-radius:15px;border-left:4px solid var(--gold);margin-bottom:30px;">
            <h3 style="color:var(--gold);margin:0 0 10px 0;">💪 당신의 이야기를 들려주세요</h3>
            <p style="color:var(--text);margin:0;opacity:0.9;">독자들이 공감할 수 있는 진솔한 경험담과 경력</p>
        </div>
        """, unsafe_allow_html=True)

        struggle_story = st.text_area("처음 시작할 때 겪었던 어려움/실패는?", value=st.session_state['temp_interview'].get('struggle_story', ''), height=100, placeholder="예: 처음에는 주식으로 500만원을 잃었습니다. 유튜브 정보만 믿고 투자했다가 큰 손실을 봤죠...")
        breakthrough = st.text_area("어떻게 극복하고 성과를 냈나요?", value=st.session_state['temp_interview'].get('breakthrough', ''), height=100, placeholder="예: 그 후 기본서 10권을 정독하고, 나만의 원칙을 세웠습니다. 1년 후 손실을 모두 만회하고 수익을 내기 시작했습니다...")

        st.markdown("---")
        st.markdown("##### 📌 작가 경력/경험 (선택)")
        author_career = st.text_area("관련 경력이나 자격, 성과가 있다면?", value=st.session_state['temp_interview'].get('author_career', ''), height=100, placeholder="예: 금융회사 7년 근무, 투자 관련 유튜브 구독자 5만명, 월 수익 3천만원 달성, CFA 자격증 보유, 강의 경력 3년...")

        col1, col2 = st.columns([1, 1])
        with col1:
            if st.button("← 이전", key="interview_prev_4", use_container_width=True):
                st.session_state['interview_step'] = 3
                st.rerun()
        with col2:
            if st.button("다음 →", key="interview_next_4", use_container_width=True, type="primary"):
                if not struggle_story or not breakthrough:
                    st.error("어려움/실패와 극복 스토리는 필수입니다")
                else:
                    st.session_state['temp_interview']['struggle_story'] = struggle_story
                    st.session_state['temp_interview']['breakthrough'] = breakthrough
                    st.session_state['temp_interview']['author_career'] = author_career
                    st.session_state['interview_step'] = 5
                    st.rerun()

    # ========== STEP 5: 마무리 ==========
    elif step == 5:
        st.markdown("""
        <div style="background:linear-gradient(135deg, rgba(201,162,75,0.1) 0%, rgba(201,162,75,0.05) 100%);padding:30px;border-radius:15px;border-left:4px solid var(--gold);margin-bottom:30px;">
            <h3 style="color:var(--gold);margin:0 0 10px 0;">✨ 마지막으로!</h3>
            <p style="color:var(--text);margin:0;opacity:0.9;">독자에게 전하고 싶은 메시지</p>
        </div>
        """, unsafe_allow_html=True)

        why_write = st.text_area("왜 이 책을 쓰려고 하나요?", value=st.session_state['temp_interview'].get('why_write', ''), height=80, placeholder="예: 저처럼 헤매는 사람들이 시행착오 없이 바로 성과를 낼 수 있도록 도와주고 싶습니다...")
        final_message = st.text_area("독자에게 마지막으로 전하고 싶은 말", value=st.session_state['temp_interview'].get('final_message', ''), height=80, placeholder="예: 누구나 할 수 있습니다. 포기하지 않으면 반드시 성공합니다...")

        # 입력 내용 미리보기
        st.markdown("---")
        st.markdown("### 📋 입력 내용 확인")

        preview_data = st.session_state['temp_interview']
        st.markdown(f"""
        <div style="background:rgba(20,20,20,0.8);padding:20px;border-radius:10px;border:1px solid var(--line);">
            <p><b>저자:</b> {preview_data.get('author_name', '')}</p>
            <p><b>분야:</b> {preview_data.get('field', '')} ({preview_data.get('experience_years', '')})</p>
            <p><b>주제:</b> {preview_data.get('topic', '')}</p>
            <p><b>타겟:</b> {preview_data.get('target_reader', '')}</p>
        </div>
        """, unsafe_allow_html=True)

        col1, col2 = st.columns([1, 1])
        with col1:
            if st.button("← 이전", key="interview_prev_5", use_container_width=True):
                st.session_state['interview_step'] = 4
                st.rerun()
        with col2:
            if st.button("📋 목차 생성하기", key="interview_generate_outline", use_container_width=True, type="primary"):
                if not get_api_key():
                    st.error("사이드바에서 API 키를 먼저 입력해주세요")
                elif not why_write or not final_message:
                    st.error("모든 항목을 입력해주세요")
                else:
                    st.session_state['temp_interview']['why_write'] = why_write
                    st.session_state['temp_interview']['final_message'] = final_message

                    # 목차만 먼저 생성
                    progress_box = st.empty()
                    interview_data = st.session_state['temp_interview']
                    success = generate_outline_only(interview_data, progress_box)

                    if success:
                        import time
                        time.sleep(1)
                        st.session_state['interview_step'] = 6  # 목차 확인 단계로 이동
                        st.rerun()

    # ========== STEP 6: 목차 확인 및 본문 생성 ==========
    elif step == 6:
        st.markdown("""
        <div style="background:linear-gradient(135deg, rgba(201,162,75,0.1) 0%, rgba(201,162,75,0.05) 100%);padding:30px;border-radius:15px;border-left:4px solid var(--gold);margin-bottom:30px;">
            <h3 style="color:var(--gold);margin:0 0 10px 0;">📋 목차 확인 및 수정</h3>
            <p style="color:var(--text);margin:0;opacity:0.9;">생성된 목차를 확인하고, 직접 수정하거나 AI로 재생성할 수 있습니다</p>
        </div>
        """, unsafe_allow_html=True)

        # 제목 표시
        book_title = st.session_state.get('book_title', '')
        subtitle = st.session_state.get('subtitle', '')
        book_concept = st.session_state.get('book_concept', '')

        if book_title:
            st.markdown(f"""
            <div style="background:rgba(30,30,30,0.9);padding:25px;border-radius:15px;border:0.5px solid var(--gold);margin-bottom:20px;text-align:center;">
                <h2 style="color:var(--gold);margin:0 0 10px 0;font-size:32px;">{book_title}</h2>
                <p style="color:var(--text2);margin:0;font-size:18px;">{subtitle}</p>
            </div>
            """, unsafe_allow_html=True)

        # 컨셉 표시
        if book_concept:
            with st.expander("💡 이 책의 고유 컨셉 보기", expanded=False):
                st.markdown(f"""
                <div style="background:rgba(201,162,75,0.1);padding:20px;border-radius:10px;border-left:3px solid var(--gold);">
                    {book_concept.replace(chr(10), '<br>')}
                </div>
                """, unsafe_allow_html=True)

        st.markdown("---")

        # 목차 표시 및 편집
        outline = st.session_state.get('outline', [])
        chapters = st.session_state.get('chapters', {})

        if outline:
            st.markdown("### 📖 목차 구성")
            st.caption("각 챕터와 소제목을 직접 수정하거나, 🔄 버튼으로 AI가 새로 생성합니다")

            for i, ch in enumerate(outline):
                ch_data = chapters.get(ch, {})
                subtopics = ch_data.get('subtopics', [])

                # 챕터 헤더
                st.markdown(f"""
                <div style="background:linear-gradient(90deg, rgba(201,162,75,0.2) 0%, rgba(30,30,30,0.9) 100%);
                            padding:15px 20px;border-radius:10px;margin:20px 0 10px 0;
                            border-left:4px solid var(--gold);">
                    <span style="color:var(--gold);font-weight:bold;font-size:18px;">PART {i+1}</span>
                </div>
                """, unsafe_allow_html=True)

                # 챕터 제목 편집
                new_ch_name = st.text_input(
                    f"챕터 {i+1} 제목",
                    value=ch,
                    key=f"ch_edit_{i}",
                    label_visibility="collapsed"
                )

                # 챕터 이름 변경 적용
                if new_ch_name != ch and new_ch_name.strip():
                    # 목차에서 이름 변경
                    st.session_state['outline'][i] = new_ch_name.strip()
                    # chapters 딕셔너리에서도 키 변경
                    st.session_state['chapters'][new_ch_name.strip()] = st.session_state['chapters'].pop(ch)
                    st.rerun()

                # 소제목들
                for j, sub in enumerate(subtopics):
                    col1, col2 = st.columns([0.5, 5.5])
                    with col1:
                        st.markdown(f"<div style='color:var(--text2);padding-top:8px;'>•</div>", unsafe_allow_html=True)
                    with col2:
                        new_sub = st.text_input(
                            f"소제목 {j+1}",
                            value=sub,
                            key=f"sub_edit_{i}_{j}",
                            label_visibility="collapsed"
                        )
                        # 소제목 변경 적용
                        if new_sub != sub and new_sub.strip():
                            st.session_state['chapters'][ch]['subtopics'][j] = new_sub.strip()
                            # subtopic_data도 업데이트
                            old_data = st.session_state['chapters'][ch]['subtopic_data'].pop(sub, {'questions': [], 'answers': [], 'content': ''})
                            st.session_state['chapters'][ch]['subtopic_data'][new_sub.strip()] = old_data
                            st.rerun()

            st.markdown("---")

        # 하단 버튼
        col1, col2 = st.columns([1, 1])
        with col1:
            if st.button("← 이전", key="interview_prev_6", use_container_width=True):
                st.session_state['interview_step'] = 5
                st.rerun()
        with col2:
            if st.button("✍️ 본문 생성하기", key="generate_body", use_container_width=True, type="primary"):
                progress_box = st.empty()
                interview_data = st.session_state.get('interview_data', st.session_state['temp_interview'])
                success = generate_body_from_outline(interview_data, progress_box)

                if success:
                    import time
                    time.sleep(1)
                    st.session_state['current_page'] = 7  # 최종 출력 페이지로 이동
                    st.rerun()

    st.stop()

# ==========================================
# 여기서부터 기존 페이지 로직 (인터뷰 완료 후)
# ==========================================

# 페이지 네비게이션 (간소화: 4단계)
simple_pages = ["주제", "목차", "본문", "완성"]
page_mapping = [0, 4, 5, 7]  # 실제 페이지 인덱스
current = st.session_state['current_page']

# 현재 페이지가 간소화된 네비게이션의 어디에 해당하는지
def get_simple_index(current_page):
    if current_page <= 0:
        return 0
    elif current_page <= 4:
        return 1
    elif current_page <= 5:
        return 2
    else:
        return 3

simple_current = get_simple_index(current)

project_title = st.session_state.get('book_title') or st.session_state.get('topic') or "새 전자책 프로젝트"
topic_label = st.session_state.get('topic') or "주제를 입력하면 프로젝트가 시작됩니다"
outline_count = len(st.session_state.get('outline', []))
subtopic_total = sum(len(ch.get('subtopics', [])) for ch in st.session_state.get('chapters', {}).values())
body_done = sum(
    1
    for ch in st.session_state.get('chapters', {}).values()
    for s in ch.get('subtopic_data', {}).values()
    if s.get('content')
)
workspace_progress = int(round(sum([
    bool(st.session_state.get('topic')),
    bool(st.session_state.get('target_persona')),
    bool(st.session_state.get('outline')),
    body_done > 0,
]) / 4 * 100))

st.markdown(f"""
<section class="workspace-hero">
    <div class="workspace-hero-top">
        <div>
            <div class="workspace-kicker">WRITEY E-BOOK STUDIO</div>
            <h1 class="workspace-title">{html.escape(str(project_title))}</h1>
            <p class="workspace-desc">{html.escape(str(topic_label))}</p>
        </div>
        <div class="workspace-badge">{simple_current + 1}/4 · {html.escape(simple_pages[simple_current])}</div>
    </div>
    <div class="workspace-metrics">
        <div class="workspace-metric">
            <div class="workspace-metric-label">진행률</div>
            <div class="workspace-metric-value">{workspace_progress}%</div>
        </div>
        <div class="workspace-metric">
            <div class="workspace-metric-label">목차</div>
            <div class="workspace-metric-value">{outline_count}개</div>
        </div>
        <div class="workspace-metric">
            <div class="workspace-metric-label">본문</div>
            <div class="workspace-metric-value">{body_done}/{subtopic_total}</div>
        </div>
        <div class="workspace-metric">
            <div class="workspace-metric-label">현재 단계</div>
            <div class="workspace-metric-value">{html.escape(simple_pages[simple_current])}</div>
        </div>
    </div>
</section>
""", unsafe_allow_html=True)

# 프리미엄 네비게이션 바 (4단계)
st.markdown('<div class="premium-nav-container">', unsafe_allow_html=True)
cols = st.columns(4)
for i, (col, page) in enumerate(zip(cols, simple_pages)):
    with col:
        if i == simple_current:
            st.markdown(f'<div class="nav-item active">{i+1}. {page}</div>', unsafe_allow_html=True)
        else:
            if st.button(f"{i+1}. {page}", key=f"nav_{i}", use_container_width=True):
                st.session_state['current_page'] = page_mapping[i]
                st.rerun()
st.markdown('</div>', unsafe_allow_html=True)

st.markdown('<div class="section-divider"></div>', unsafe_allow_html=True)

# API 키가 없으면 환영 화면 표시
if not get_api_key():
    st.markdown("""
    <section class="onboarding-card">
        <div class="onboarding-eyebrow">First Setup</div>
        <div class="onboarding-title">Claude API 키를 연결하면 바로 제작을 시작할 수 있습니다.</div>
        <p class="onboarding-copy">
            Writey는 주제 분석, 타겟 선정, 목차 구성, 본문 생성, 최종 다운로드까지 한 흐름으로 이어지는 전자책 제작 스튜디오입니다.
            처음 한 번만 왼쪽 사이드바에 API 키를 저장하면 다음 접속부터 바로 사용할 수 있어요.
        </p>
        <div class="api-callout">
            <b>먼저 왼쪽 사이드바의 Claude API 키 입력란에 sk-ant-api03-... 형식의 키를 붙여넣어 주세요.</b><br>
            입력한 키는 브라우저에 저장되어 다음 접속 때 자동으로 불러옵니다.
        </div>
        <div class="setup-grid">
            <div class="setup-step">
                <div class="setup-number">1</div>
                <div class="setup-title">Anthropic 가입</div>
                <div class="setup-copy">Anthropic Console에서 계정을 만들고 로그인합니다. Google 계정 가입이 가장 빠릅니다.</div>
            </div>
            <div class="setup-step">
                <div class="setup-number">2</div>
                <div class="setup-title">결제 수단 등록</div>
                <div class="setup-copy">Settings의 Billing 메뉴에서 카드를 등록하고 $5~10 정도의 크레딧을 충전합니다.</div>
            </div>
            <div class="setup-step">
                <div class="setup-number">3</div>
                <div class="setup-title">API 키 발급</div>
                <div class="setup-copy">API Keys 메뉴에서 Create Key를 누른 뒤 생성된 키를 복사해 사이드바에 붙여넣습니다.</div>
            </div>
        </div>
    </section>
    """, unsafe_allow_html=True)

    link_col1, link_col2, link_col3 = st.columns(3)
    with link_col1:
        st.link_button("Anthropic 가입", "https://console.anthropic.com/", use_container_width=True, type="primary")
    with link_col2:
        st.link_button("Billing 열기", "https://console.anthropic.com/settings/billing", use_container_width=True)
    with link_col3:
        st.link_button("API Keys 열기", "https://console.anthropic.com/settings/keys", use_container_width=True)

    st.warning("⚠️ API 키는 생성 시 한 번만 보여줍니다. 꼭 복사해두세요!")
    st.link_button("📺 Claude API 키 발급 방법 (유튜브)", "https://www.youtube.com/results?search_query=anthropic+claude+api+key+발급", use_container_width=True)

    st.stop()  # API 키 없으면 여기서 멈춤

# ==========================================
# PAGE 0: 주제 & 시장분석
# ==========================================
if current == 0:
    st.markdown("""
    <div class="section-title-box">
        <span class="section-step">STEP 01</span>
        <h2>주제 선정 & 시장 분석</h2>
        <p>AI가 전자책의 성공 가능성을 분석합니다</p>
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns([1, 1], gap="large")

    with col1:
        st.markdown("### 주제 입력")

        topic = st.text_input("어떤 주제로 전자책을 쓸까요?", value=st.session_state['topic'], placeholder="예: 주식 배당으로 월 100만원", key="p0_topic", label_visibility="collapsed")
        if topic != st.session_state['topic']:
            st.session_state['topic'] = topic
            st.session_state['score_details'] = None

        # 빠른 제작 버튼 (자동 모드)
        st.markdown("""
        <div class="quick-start-card">
            <div class="eyebrow">빠른 제작 모드</div>
            <div class="headline">목차와 본문을 한 번에 자동 완성</div>
        </div>
        """, unsafe_allow_html=True)

        if st.button("🚀 빠른 제작 시작", use_container_width=True, key="p0_auto", type="primary"):
            if not topic:
                st.error("주제를 먼저 입력해주세요")
            elif not get_api_key():
                st.error("사이드바에서 API 키를 입력해주세요")
            else:
                progress_box = st.empty()
                success = auto_generate_all(topic, progress_box)
                if success:
                    import time
                    time.sleep(1)
                    st.session_state['current_page'] = 5  # 본문 페이지로 이동
                    st.rerun()

        st.markdown("---")
        st.caption("또는 시장 분석부터 단계별로 진행:")

        if st.button("📊 시장 분석 먼저 하기", use_container_width=True, key="p0_analyze"):
            if not topic:
                st.error("주제를 입력해주세요")
            elif not get_api_key():
                st.error("사이드바에서 API 키를 입력해주세요")
            else:
                with st.spinner("AI가 시장을 분석하고 있습니다..."):
                    result = analyze_market_deep(topic)
                    parsed = parse_json(result)
                    if parsed:
                        st.session_state['score_details'] = parsed
                        st.rerun()

    with col2:
        if st.session_state.get('score_details'):
            d = st.session_state['score_details']
            score = d.get('total_score', 0)
            verdict = d.get('verdict', '')
            v_class = "verdict-go" if "추천" in verdict else ("verdict-wait" if "보류" in verdict else "verdict-no")

            st.markdown(f"""
            <div class="score-card">
                <div class="score-number">{score}</div>
                <div style="font-size:14px;color:var(--text-dim);margin-top:8px;">종합 점수</div>
                <div style="margin-top:24px;"><span class="{v_class}">{verdict}</span></div>
            </div>
            """, unsafe_allow_html=True)

            st.markdown(f"""
            <div class="result-card" style="margin-top:20px;">
                <div style="font-size:13px;color:var(--text-dim);margin-bottom:8px;">AI 분석 요약</div>
                <div style="font-size:15px;color:var(--text-bright);line-height:1.7;">{d.get('verdict_reason', '')}</div>
            </div>
            """, unsafe_allow_html=True)

            sd = d.get('search_data', {})
            if sd:
                st.markdown(f"""
                <div class="data-card" style="margin-top:16px;">
                    <b>검색 데이터</b><br><br>
                    <div style="display:grid;grid-template-columns:1fr 1fr;gap:12px;">
                        <div>• 네이버: <b>{sd.get('naver_monthly', 'N/A')}</b></div>
                        <div>• 구글: <b>{sd.get('google_monthly', 'N/A')}</b></div>
                        <div>• 블로그: <b>{sd.get('naver_blog_posts', 'N/A')}</b></div>
                        <div>• 유튜브: <b>{sd.get('youtube_videos', 'N/A')}</b></div>
                    </div>
                </div>
                """, unsafe_allow_html=True)

            ms = d.get('market_size', {})
            comp = d.get('competition', {})

            c1, c2 = st.columns(2)
            with c1:
                st.markdown(f'<div class="stat-box"><div class="stat-value">{ms.get("level", "")}</div><div class="stat-label">시장 규모 ({ms.get("score", 0)}점)</div></div>', unsafe_allow_html=True)
            with c2:
                st.markdown(f'<div class="stat-box"><div class="stat-value">{comp.get("level", "")}</div><div class="stat-label">경쟁 강도 ({comp.get("score", 0)}점)</div></div>', unsafe_allow_html=True)

            if comp.get('your_opportunity'):
                st.success(f"**차별화 기회:** {comp.get('your_opportunity', '')}")

            # 경쟁 도서 검색 - 주제 키워드로 직접 검색
            current_topic = st.session_state.get('topic', '')
            if current_topic:
                st.markdown("""
                <div style="margin-top:35px;">
                    <div style="display:flex;align-items:center;gap:12px;margin-bottom:25px;">
                        <div style="width:50px;height:50px;background:linear-gradient(135deg, var(--gold) 0%, var(--gold-dark) 100%);border-radius:12px;display:flex;align-items:center;justify-content:center;">
                            <span style="font-size:26px;">🔍</span>
                        </div>
                        <div>
                            <h4 style="color:var(--gold);margin:0;font-size:22px;font-weight:600;">경쟁 도서 직접 확인하기</h4>
                            <p style="color:var(--text2);margin:4px 0 0 0;font-size:14px;">각 플랫폼에서 이 주제의 책들을 살펴보세요</p>
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)

                # 플랫폼별 검색 URL 매핑
                platforms = [
                    {
                        'name': '크몽 전자책',
                        'icon': '📘',
                        'url': 'https://kmong.com/search?c=ebook&q=',
                        'desc': '전자책/PDF 마켓',
                        'gradient': 'linear-gradient(135deg, #667eea 0%, #764ba2 100%)'
                    },
                    {
                        'name': '리디북스',
                        'icon': '📗',
                        'url': 'https://ridibooks.com/search?q=',
                        'desc': '국내 최대 전자책',
                        'gradient': 'linear-gradient(135deg, #11998e 0%, #38ef7d 100%)'
                    },
                    {
                        'name': 'YES24',
                        'icon': '📙',
                        'url': 'https://www.yes24.com/Product/Search?domain=BOOK&query=',
                        'desc': '종합 서점',
                        'gradient': 'linear-gradient(135deg, #f093fb 0%, #f5576c 100%)'
                    },
                    {
                        'name': '교보문고',
                        'icon': '📕',
                        'url': 'https://search.kyobobook.co.kr/search?keyword=',
                        'desc': '국내 대표 서점',
                        'gradient': 'linear-gradient(135deg, #4facfe 0%, #00f2fe 100%)'
                    },
                    {
                        'name': '클래스101',
                        'icon': '🎓',
                        'url': 'https://class101.net/search?query=',
                        'desc': '온라인 클래스',
                        'gradient': 'linear-gradient(135deg, #fa709a 0%, #fee140 100%)'
                    },
                    {
                        'name': '탈잉',
                        'icon': '👨‍🏫',
                        'url': 'https://taling.me/search?query=',
                        'desc': '재능 마켓',
                        'gradient': 'linear-gradient(135deg, #a8edea 0%, #fed6e3 100%)'
                    }
                ]

                search_query = urllib.parse.quote(current_topic)

                cols = st.columns(3)
                for idx, platform in enumerate(platforms):
                    with cols[idx % 3]:
                        search_url = platform['url'] + search_query
                        st.markdown(f"""
                        <a href="{search_url}" target="_blank" style="text-decoration:none;display:block;margin-bottom:15px;">
                            <div style="background:rgba(25,25,25,0.9);border:1px solid rgba(201,162,75,0.3);border-radius:16px;overflow:hidden;transition:all 0.3s ease;">
                                <div style="height:80px;background:{platform['gradient']};display:flex;align-items:center;justify-content:center;">
                                    <span style="font-size:40px;">{platform['icon']}</span>
                                </div>
                                <div style="padding:18px;text-align:center;">
                                    <div style="font-size:17px;color:var(--text);font-weight:700;margin-bottom:6px;">
                                        {platform['name']}
                                    </div>
                                    <div style="font-size:13px;color:var(--text2);margin-bottom:12px;">
                                        {platform['desc']}
                                    </div>
                                    <div style="background:linear-gradient(135deg, var(--gold) 0%, var(--gold-dark) 100%);color:var(--dark);padding:10px 16px;border-radius:8px;font-size:13px;font-weight:700;">
                                        🔍 "{current_topic[:15]}{'...' if len(current_topic) > 15 else ''}" 검색
                                    </div>
                                </div>
                            </div>
                        </a>
                        """, unsafe_allow_html=True)
        else:
            st.markdown("""
            <div class="empty-state">
                <div class="empty-state-text">
                    주제를 입력하고 <b>AI 시장 분석</b>을 시작하세요<br>
                    검색량, 경쟁 강도, 수익 가능성을 분석합니다
                </div>
            </div>
            """, unsafe_allow_html=True)

    st.markdown('<div class="next-section"></div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 2, 1])
    with c2:
        if st.button("다음 단계로 타겟 설정", key="p0_next", use_container_width=True):
            go_next()
            st.rerun()


# ==========================================
# PAGE 1: 타겟 & 컨셉
# ==========================================
elif current == 1:
    st.markdown("""
    <div class="section-title-box">
        <span class="section-step">STEP 02</span>
        <h2>타겟 설정 & 제목 생성</h2>
        <p>구매할 사람을 정하고 끌리는 제목을 만듭니다</p>
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns([1, 1])

    with col1:
        st.markdown("### 타겟 선정")

        if st.button("AI 타겟 추천", key="p1_target"):
            if st.session_state['topic'] and get_api_key():
                with st.spinner("분석 중..."):
                    result = suggest_targets(st.session_state['topic'])
                    parsed = parse_json(result)
                    if parsed:
                        st.session_state['suggested_targets'] = parsed
                        st.rerun()

        if st.session_state.get('suggested_targets'):
            personas = st.session_state['suggested_targets'].get('personas', [])[:3]

            for i, p in enumerate(personas):
                target_name = p.get('name', '')
                target_demo = p.get('demographics', '')
                target_needs = p.get('needs', '')
                target_pains = p.get('pain_points', [])

                st.markdown(f"""<div class="data-card">
                    <b>{html.escape(str(target_name))}</b><br>
                    <small>{html.escape(str(target_demo))}</small><br>
                    <small style="color:var(--gold);">{html.escape(str(target_needs))}</small>
                </div>""", unsafe_allow_html=True)

                if st.button(f"이 타겟 선택", key=f"sel_target_{i}", use_container_width=True):
                    selected_target = f"{target_name} - {target_demo}"
                    st.session_state['target_persona'] = selected_target
                    st.session_state['p1_persona'] = selected_target
                    st.session_state['pain_points'] = ", ".join(target_pains[:5])
                    st.session_state['suggested_targets'] = None
                    st.rerun()

        st.markdown("---")
        st.markdown("### 선택된 타겟")
        persona = st.text_area("타겟:", value=st.session_state.get('target_persona', ''), height=60, key="p1_persona", placeholder="AI 추천에서 선택하거나 직접 입력")
        st.session_state['target_persona'] = persona

        if st.button("고민 심층 분석", key="p1_analyze", use_container_width=True):
            if not persona:
                st.error("타겟을 먼저 입력해주세요")
            elif not get_api_key():
                st.error("API 키를 입력해주세요")
            else:
                with st.spinner("심층 분석 중..."):
                    r = analyze_pains_deep(st.session_state['topic'], persona)
                    parsed = parse_json(r)
                    if parsed:
                        st.session_state['analyzed_pains'] = parsed
                        surface = parsed.get('surface_pains', {}).get('pains', [])
                        hidden = parsed.get('hidden_pains', {}).get('pains', [])
                        st.session_state['pain_points'] = ", ".join((surface + hidden)[:6])
                        st.rerun()
                    else:
                        st.error("분석 실패. 다시 시도해주세요")

        if st.session_state.get('analyzed_pains'):
            p = st.session_state['analyzed_pains']
            st.markdown("**표면적 고민**")
            for pain in p.get('surface_pains', {}).get('pains', []):
                st.write(f"• {pain}")
            st.markdown("**숨겨진 진짜 고민**")
            for pain in p.get('hidden_pains', {}).get('pains', []):
                st.write(f"• {pain}")
            if p.get('marketing_hook'):
                st.info(f"**마케팅 훅:** {p.get('marketing_hook', '')}")

    with col2:
        st.markdown("### 베스트셀러급 제목 생성")

        # 선택된 제목이 있으면 상단에 확정 표시
        if st.session_state.get('book_title'):
            st.markdown(f"""
            <div style="background:linear-gradient(135deg,#10b981,#059669);padding:16px 20px;border-radius:12px;margin-bottom:20px;">
                <div style="color:white;font-size:12px;margin-bottom:6px;">✓ 확정된 제목</div>
                <div style="color:white;font-size:20px;font-weight:700;">{html.escape(st.session_state.get('book_title', ''))}</div>
                <div style="color:rgba(255,255,255,0.85);font-size:14px;margin-top:4px;">{html.escape(st.session_state.get('subtitle', ''))}</div>
            </div>
            """, unsafe_allow_html=True)

        pain_points = st.text_area("독자의 고민:", value=st.session_state['pain_points'], height=60, key="p1_pains")
        st.session_state['pain_points'] = pain_points

        if st.button("베스트셀러 제목 생성", key="p1_title"):
            if st.session_state['topic']:
                with st.spinner("베스트셀러 패턴 분석 중..."):
                    r = generate_titles_bestseller(st.session_state['topic'], st.session_state['target_persona'], st.session_state['pain_points'])
                    parsed = parse_json(r)
                    if parsed:
                        st.session_state['generated_titles'] = parsed
                        st.rerun()

        if st.session_state.get('generated_titles'):
            titles_list = st.session_state['generated_titles'].get('titles', [])[:5]
            for i, t in enumerate(titles_list):
                title_val = t.get('title', '')
                subtitle_val = t.get('subtitle', '')
                concept_val = t.get('concept', '')

                st.markdown(f"""
                <div class="title-card">
                    <div class="title-main">{html.escape(title_val)}</div>
                    <div class="title-sub">{html.escape(subtitle_val)}</div>
                    <div style="font-size:11px;color:var(--gold);margin-top:12px;letter-spacing:2px;">{html.escape(concept_val)}</div>
                </div>
                """, unsafe_allow_html=True)

                if st.button(f"✓ 이 제목으로 확정", key=f"sel_title_{i}", use_container_width=True):
                    st.session_state['book_title'] = title_val
                    st.session_state['subtitle'] = subtitle_val
                    st.toast(f"'{title_val}' 제목이 확정되었습니다!")
                    st.rerun()

        # 직접 입력 옵션
        st.markdown("---")
        st.markdown("#### 또는 직접 입력")
        manual_title = st.text_input("제목 입력", key="manual_title_v3")
        manual_subtitle = st.text_input("부제 입력", key="manual_subtitle_v3")
        if st.button("✓ 직접 입력한 제목으로 확정", key="manual_confirm_v3", use_container_width=True):
            if manual_title:
                st.session_state['book_title'] = manual_title
                st.session_state['subtitle'] = manual_subtitle if manual_subtitle else ''
                st.toast(f"'{manual_title}' 제목이 확정되었습니다!")
                st.rerun()

    st.markdown('<div class="next-section"></div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 1, 1])
    with c1:
        if st.button("이전", key="p1_prev", use_container_width=True):
            go_prev()
            st.rerun()
    with c3:
        if st.button("다음 경쟁분석", key="p1_next", use_container_width=True):
            go_next()
            st.rerun()


# ==========================================
# PAGE 2: 경쟁도서 분석
# ==========================================
elif current == 2:
    st.markdown("""
    <div class="section-title-box">
        <span class="section-step">STEP 03</span>
        <h2>경쟁 도서 분석</h2>
        <p>기존 도서의 부정 리뷰를 분석해서 숨은 니즈를 찾습니다</p>
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns([1, 1])

    with col1:
        st.markdown("### 부정 리뷰 분석")

        if st.button("경쟁 도서 분석하기", use_container_width=True, key="p2_analyze"):
            if not st.session_state['topic']:
                st.error("주제를 먼저 입력해주세요")
            elif not get_api_key():
                st.error("API 키를 입력해주세요")
            else:
                with st.spinner("경쟁 도서 분석 중..."):
                    result = analyze_competitor_reviews(st.session_state['topic'])
                    parsed = parse_json(result)
                    if parsed:
                        st.session_state['review_analysis'] = parsed
                        concepts = parsed.get('concept_suggestions', [])
                        st.session_state['market_gaps'] = [c.get('concept', '') for c in concepts]
                        st.rerun()

        if st.session_state.get('review_analysis'):
            a = st.session_state['review_analysis']
            scope = a.get('analysis_scope', {})
            col_s1, col_s2 = st.columns(2)
            with col_s1:
                st.markdown(f'<div class="stat-box"><div class="stat-value">{scope.get("books_analyzed", "N/A")}</div><div class="stat-label">분석 도서</div></div>', unsafe_allow_html=True)
            with col_s2:
                st.markdown(f'<div class="stat-box"><div class="stat-value">{scope.get("negative_reviews", "N/A")}</div><div class="stat-label">부정 리뷰</div></div>', unsafe_allow_html=True)

    with col2:
        st.markdown("### 분석 결과")

        if st.session_state.get('review_analysis'):
            a = st.session_state['review_analysis']

            patterns = a.get('negative_patterns', [])
            if patterns:
                st.markdown("#### 독자 불만 패턴")
                for i, p in enumerate(patterns[:3], 1):
                    st.markdown(f"""<div class="data-card">
                        <b>{i}. {p.get('pattern', '')} ({p.get('frequency', '')})</b>
                    </div>""", unsafe_allow_html=True)
                    for rev in p.get('example_reviews', []):
                        st.caption(f'"{rev}"')
                    st.info(f"**숨겨진 니즈:** {p.get('hidden_need', '')}")
                    st.success(f"**해결책:** {p.get('solution', '')}")

            concepts = a.get('concept_suggestions', [])
            if concepts:
                st.markdown("#### 차별화 컨셉")
                for c in concepts[:2]:
                    st.markdown(f"""
                    <div class="info-card">
                        <b>「{html.escape(c.get('concept', ''))}」</b><br>
                        <span style="color:rgba(255,255,255,0.7);">{html.escape(c.get('why_works', ''))}</span>
                    </div>
                    """, unsafe_allow_html=True)
        else:
            st.markdown('<div style="text-align:center;padding:60px;background:rgba(255,255,255,0.03);border-radius:16px;border:1px solid rgba(201,162,75,0.15);"><p style="color:rgba(255,255,255,0.5);">분석 버튼을 눌러주세요</p></div>', unsafe_allow_html=True)

    st.markdown('<div class="next-section"></div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 1, 1])
    with c1:
        if st.button("이전", key="p2_prev", use_container_width=True):
            go_prev()
            st.rerun()
    with c3:
        if st.button("다음 학습", key="p2_next", use_container_width=True):
            go_next()
            st.rerun()


# ==========================================
# PAGE 3: 학습 & 리서치
# ==========================================
elif current == 3:
    st.markdown("""
    <div class="section-title-box">
        <span class="section-step">STEP 04</span>
        <h2>학습 & 리서치</h2>
        <p>베스트셀러 분석, 트렌드 파악, 핵심 인사이트를 수집합니다</p>
    </div>
    """, unsafe_allow_html=True)

    # 탭으로 구분
    tab1, tab2, tab3 = st.tabs(["레퍼런스 추천", "트렌드 분석", "경쟁서 분석"])

    # ========== 탭1: 레퍼런스 추천 & 아이디어 ==========
    with tab1:
        topic = st.session_state.get('topic', '')

        col1, col2 = st.columns([1, 1])

        with col1:
            st.markdown("### 레퍼런스 자동 추천")
            st.markdown('<p style="color:var(--text2);font-size:13px;">주제에 맞는 참고 자료를 AI가 자동으로 추천합니다</p>', unsafe_allow_html=True)

            if not topic:
                st.warning("먼저 시장분석 페이지에서 주제를 입력해주세요")
            else:
                st.markdown(f'<p style="color:var(--accent);font-size:14px;margin:10px 0;">현재 주제: <b>{html.escape(topic)}</b></p>', unsafe_allow_html=True)

                ref_category = st.selectbox("추천 카테고리", ["베스트셀러 도서", "핵심 개념/이론", "성공 사례", "전문가 인사이트"], key="ref_cat")

                if st.button("레퍼런스 추천받기", use_container_width=True, key="auto_ref_btn"):
                    if not get_api_key():
                        st.error("API 키를 입력해주세요")
                    else:
                        with st.spinner("관련 레퍼런스 심층 분석 중..."):
                            prompt = f"""'{topic}' 주제로 전자책을 쓰려고 합니다.
'{ref_category}' 카테고리에서 참고할 만한 자료 3개를 추천해주세요.

중요: 마치 이 책/자료를 직접 읽은 것처럼 아주 상세하게 설명해주세요.

각 추천 자료에 대해 다음을 포함해주세요:
1. 제목과 저자
2. 책/자료의 핵심 메시지 (10문장 이상으로 상세히)
3. 주요 챕터/섹션별 핵심 내용
4. 저자의 핵심 주장과 근거
5. 실제 사례나 스토리
6. 전자책에 활용할 수 있는 구체적 인사이트

중요: 책의 모든 주요 챕터를 빠짐없이 요약해주세요. 일부만 하지 말고 전체 목차를 다 포함해주세요.

JSON 형식으로 응답:
{{
    "recommendations": [
        {{
            "title": "자료 제목",
            "author": "저자/출처",
            "core_message": "이 책의 핵심 메시지와 주장을 10문장 이상으로 상세하게 설명",
            "chapters": [
                {{"name": "1장 제목", "summary": "이 챕터의 핵심 내용 2-3문장"}},
                {{"name": "2장 제목", "summary": "이 챕터의 핵심 내용 2-3문장"}},
                {{"name": "3장 제목", "summary": "이 챕터의 핵심 내용 2-3문장"}},
                {{"name": "4장 제목", "summary": "이 챕터의 핵심 내용 2-3문장"}},
                {{"name": "5장 제목", "summary": "이 챕터의 핵심 내용 2-3문장"}},
                {{"name": "6장 제목", "summary": "이 챕터의 핵심 내용 2-3문장"}},
                {{"name": "7장 제목", "summary": "이 챕터의 핵심 내용 2-3문장"}},
                {{"name": "8장 제목", "summary": "이 챕터의 핵심 내용 2-3문장"}}
            ],
            "key_arguments": ["저자의 핵심 주장1과 근거", "핵심 주장2와 근거", "핵심 주장3과 근거"],
            "real_examples": ["책에 나온 실제 사례/스토리 1", "사례 2", "사례 3"],
            "key_insights": ["전자책에 활용할 인사이트 1", "인사이트 2", "인사이트 3", "인사이트 4", "인사이트 5"],
            "application": "내 전자책에 구체적으로 활용하는 방법 (3문장 이상)"
        }}
    ]
}}"""
                            result = ask_ai(prompt, 0.8)
                            parsed = parse_json(result)
                            if parsed and parsed.get('recommendations'):
                                st.session_state['recommended_refs'] = parsed['recommendations']
                                st.rerun()
                            else:
                                st.error("추천 생성 실패")

                # 추천된 레퍼런스 표시
                if st.session_state.get('recommended_refs'):
                    st.markdown("---")
                    st.markdown("#### 추천 레퍼런스")

                    for i, ref in enumerate(st.session_state['recommended_refs']):
                        st.markdown(f"""<div class="data-card">
                            <b>{html.escape(str(ref.get('title', '')))}</b>
                            <br><small style="color:var(--text2);">{html.escape(str(ref.get('author', '')))}</small>
                        </div>""", unsafe_allow_html=True)

                        # 핵심 메시지
                        if ref.get('core_message'):
                            st.markdown("**핵심 메시지**")
                            st.markdown(f'<p style="color:#e0e0e0;line-height:1.8;">{html.escape(str(ref.get("core_message", "")))}</p>', unsafe_allow_html=True)

                        # 챕터 요약
                        if ref.get('chapters'):
                            st.markdown("**챕터별 요약**")
                            for ch in ref.get('chapters', []):
                                st.markdown(f"""<div style="background:rgba(255,255,255,0.05);padding:12px 16px;margin:8px 0;border-left:3px solid var(--gold);">
                                    <b style="color:var(--gold);">{html.escape(str(ch.get('name', '')))}</b><br>
                                    <span style="color:#e0e0e0;">{html.escape(str(ch.get('summary', '')))}</span>
                                </div>""", unsafe_allow_html=True)

                        # 핵심 주장
                        if ref.get('key_arguments'):
                            st.markdown("**저자의 핵심 주장**")
                            for arg in ref.get('key_arguments', []):
                                st.info(arg)

                        # 실제 사례
                        if ref.get('real_examples'):
                            st.markdown("**실제 사례**")
                            for ex in ref.get('real_examples', []):
                                st.markdown(f'<p style="color:#e0e0e0;">- {html.escape(str(ex))}</p>', unsafe_allow_html=True)

                        # 활용 인사이트
                        if ref.get('key_insights'):
                            st.markdown("**활용 인사이트**")
                            for insight in ref.get('key_insights', []):
                                st.success(insight)

                        # 적용 방법
                        if ref.get('application'):
                            st.markdown("**내 책에 적용하는 방법**")
                            st.markdown(f'<p style="color:#e0e0e0;line-height:1.8;">{html.escape(str(ref.get("application", "")))}</p>', unsafe_allow_html=True)

                        if st.button("이 자료 저장하기", key=f"save_ref_{i}", use_container_width=True):
                            ref_item = {
                                'title': ref.get('title', ''),
                                'source': ref.get('author', ''),
                                'core_message': ref.get('core_message', ''),
                                'chapters': ref.get('chapters', []),
                                'key_arguments': ref.get('key_arguments', []),
                                'real_examples': ref.get('real_examples', []),
                                'key_insights': ref.get('key_insights', []),
                                'application': ref.get('application', ''),
                                'type': 'recommended',
                                'added_at': datetime.now().strftime('%Y-%m-%d %H:%M')
                            }
                            st.session_state['knowledge_hub'].append(ref_item)
                            st.success("저장 완료")
                            st.rerun()

                        st.markdown("---")

        with col2:
            st.markdown("### 저장된 자료 & 아이디어 도출")
            hub = st.session_state.get('knowledge_hub', [])

            if hub:
                st.caption(f"총 {len(hub)}개 자료 저장됨")

                for i, item in enumerate(hub):
                    title = item.get('title', item.get('main_topic', item.get('source', f'자료 {i+1}')))

                    st.markdown(f"""<div class="data-card">
                        <b>{html.escape(str(title))}</b><br>
                        <small>{html.escape(str(item.get('source', '')))} | {item.get('added_at', '')}</small>
                    </div>""", unsafe_allow_html=True)

                    # 핵심 메시지 전체 표시
                    if item.get('core_message'):
                        st.write(item['core_message'])

                    # 핵심 인사이트 표시
                    if item.get('key_insights'):
                        st.markdown("**핵심 인사이트:**")
                        for insight in item.get('key_insights', [])[:3]:
                            st.success(insight)

                    # 적용 방법 표시
                    if item.get('application'):
                        st.info(f"적용법: {item['application']}")

                    col_a, col_b = st.columns([1, 1])
                    with col_a:
                        if st.button("상세보기", key=f"view_ref_{i}"):
                            st.session_state[f'show_detail_{i}'] = not st.session_state.get(f'show_detail_{i}', False)
                            st.rerun()
                    with col_b:
                        if st.button("삭제", key=f"del_ref_{i}"):
                            st.session_state['knowledge_hub'].pop(i)
                            st.rerun()

                    # 상세 보기 토글
                    if st.session_state.get(f'show_detail_{i}', False):
                        if item.get('chapters'):
                            st.markdown("**챕터 요약:**")
                            for ch in item.get('chapters', []):
                                st.markdown(f"""<div style="background:rgba(255,255,255,0.05);padding:10px 14px;margin:6px 0;border-left:2px solid var(--gold);">
                                    <b style="color:var(--gold);">{html.escape(str(ch.get('name', '')))}</b><br>
                                    <span style="color:#e0e0e0;font-size:14px;">{html.escape(str(ch.get('summary', '')))}</span>
                                </div>""", unsafe_allow_html=True)
                        if item.get('key_arguments'):
                            st.markdown("**핵심 주장:**")
                            for arg in item.get('key_arguments', []):
                                st.markdown(f'<p style="color:#e0e0e0;">- {html.escape(str(arg))}</p>', unsafe_allow_html=True)
                        if item.get('real_examples'):
                            st.markdown("**실제 사례:**")
                            for ex in item.get('real_examples', []):
                                st.markdown(f'<p style="color:#e0e0e0;">- {html.escape(str(ex))}</p>', unsafe_allow_html=True)

                    st.markdown("---")

                st.markdown("---")
                st.markdown("#### 아이디어 도출")
                st.markdown('<p style="color:var(--text2);font-size:13px;">수집된 자료를 바탕으로 전자책 아이디어를 생성합니다</p>', unsafe_allow_html=True)

                if st.button("아이디어 생성하기", use_container_width=True, key="ideate_btn"):
                    if not get_api_key():
                        st.error("API 키를 입력해주세요")
                    else:
                        with st.spinner("아이디어 생성 중..."):
                            hub_summary = ""
                            for item in hub[:5]:
                                hub_summary += f"\n[{item.get('title', '')}]\n"
                                if item.get('core_message'):
                                    hub_summary += f"핵심 메시지: {item.get('core_message', '')}\n"
                                if item.get('key_insights'):
                                    hub_summary += f"인사이트: {', '.join(item.get('key_insights', []))}\n"
                                if item.get('application'):
                                    hub_summary += f"적용법: {item.get('application', '')}\n"

                            prompt = f"""다음 수집된 자료들을 철저히 분석하여 '{topic}' 주제의 전자책 아이디어를 도출해주세요:

수집된 자료:
{hub_summary}

위 자료들의 공통점, 차이점, 빈틈을 분석하고 다음을 포함해서 아이디어를 생성해주세요:
1. 기존 책들과 확실히 다른 차별화된 콘셉트
2. 독자의 문제를 해결하는 독특한 관점
3. 구체적인 목차/콘텐츠 구성 아이디어
4. 타겟 독자에게 강하게 어필할 포인트

JSON 형식으로 응답:
{{
    "main_concept": "핵심 콘셉트 한 문장 (경쟁작과 어떻게 다른지 명확히)",
    "unique_angles": ["독특한 관점 1 (왜 이 관점이 효과적인지 설명)", "관점 2", "관점 3"],
    "content_ideas": ["챕터 아이디어 1", "챕터 아이디어 2", "챕터 아이디어 3", "챕터 아이디어 4", "챕터 아이디어 5"],
    "appeal_points": ["어필 포인트 1", "포인트 2", "포인트 3"],
    "title_suggestions": ["제목 제안 1 (부제 포함)", "제목 제안 2 (부제 포함)", "제목 제안 3 (부제 포함)"],
    "differentiation": "경쟁작 대비 구체적인 차별화 전략 (3문장 이상)"
}}"""
                            result = ask_ai(prompt, 0.9)
                            parsed = parse_json(result)
                            if parsed:
                                st.session_state['generated_ideas'] = parsed
                                st.rerun()
                            else:
                                st.error("아이디어 생성 실패")

                # 생성된 아이디어 표시
                if st.session_state.get('generated_ideas'):
                    ideas = st.session_state['generated_ideas']

                    st.markdown(f"""<div class="summary-hub">
                        <b>핵심 콘셉트</b><br>
                        {html.escape(str(ideas.get('main_concept', '')))}
                    </div>""", unsafe_allow_html=True)

                    if ideas.get('unique_angles'):
                        st.markdown("**독특한 관점**")
                        for angle in ideas.get('unique_angles', []):
                            st.info(angle)

                    if ideas.get('title_suggestions'):
                        st.markdown("**제목 제안**")
                        for title in ideas.get('title_suggestions', []):
                            st.success(title)

                    if ideas.get('content_ideas'):
                        st.markdown("**콘텐츠 아이디어**")
                        for idea in ideas.get('content_ideas', []):
                            st.write(f"- {idea}")

                    if ideas.get('differentiation'):
                        st.markdown(f"""<div class="data-card">
                            <b>차별화 전략</b><br>
                            <small>{html.escape(str(ideas.get('differentiation', '')))}</small>
                        </div>""", unsafe_allow_html=True)
            else:
                st.markdown('<div style="text-align:center;padding:60px 20px;background:var(--card);border:1px solid var(--line);"><p style="color:var(--text2);">추천받은 레퍼런스를 저장하면<br>아이디어를 도출할 수 있습니다</p></div>', unsafe_allow_html=True)

    # ========== 탭2: 트렌드 분석 ==========
    with tab2:
        st.markdown("### 시장 트렌드 분석")
        st.markdown('<p style="color:var(--text2);">현재 인기 있는 전자책 트렌드와 키워드를 파악합니다</p>', unsafe_allow_html=True)

        col1, col2 = st.columns([1, 1])

        with col1:
            st.markdown("#### 트렌드 키워드 분석")
            trend_topic = st.text_input("분석할 분야", key="trend_topic", placeholder="예: 재테크, 자기계발, 다이어트...")

            if st.button("트렌드 분석", use_container_width=True, key="trend_btn"):
                if not trend_topic:
                    st.error("분야를 입력하세요")
                elif not get_api_key():
                    st.error("API 키를 입력해주세요")
                else:
                    with st.spinner("트렌드 분석 중..."):
                        prompt = f"""'{trend_topic}' 분야의 전자책 시장 트렌드를 분석해주세요.

JSON 형식으로 응답:
{{
    "hot_keywords": ["인기 키워드 1", "키워드 2", "키워드 3", "키워드 4", "키워드 5"],
    "rising_topics": ["떠오르는 주제 1", "주제 2", "주제 3"],
    "reader_needs": ["독자가 원하는 것 1", "원하는 것 2", "원하는 것 3"],
    "content_gaps": ["시장에서 부족한 콘텐츠 1", "부족한 콘텐츠 2"],
    "recommended_angles": ["추천 접근 방식 1", "접근 방식 2", "접근 방식 3"],
    "avoid": ["피해야 할 것 1", "피해야 할 것 2"]
}}"""
                        result = ask_ai(prompt, 0.8)
                        parsed = parse_json(result)
                        if parsed:
                            st.session_state['trend_analysis'] = parsed
                            st.rerun()

            if st.session_state.get('trend_analysis'):
                ta = st.session_state['trend_analysis']
                if ta.get('hot_keywords'):
                    st.write("**인기 키워드**")
                    st.write(" | ".join(ta.get('hot_keywords', [])))
                if ta.get('rising_topics'):
                    st.write("**떠오르는 주제**")
                    for t in ta.get('rising_topics', []):
                        st.write(f"- {t}")

        with col2:
            st.markdown("#### 독자 니즈")
            if st.session_state.get('trend_analysis'):
                ta = st.session_state['trend_analysis']
                if ta.get('reader_needs'):
                    st.write("**독자가 원하는 것**")
                    for n in ta.get('reader_needs', []):
                        st.info(n)
                if ta.get('content_gaps'):
                    st.write("**시장 빈틈**")
                    for g in ta.get('content_gaps', []):
                        st.success(g)
                if ta.get('recommended_angles'):
                    st.write("**추천 접근법**")
                    for r in ta.get('recommended_angles', []):
                        st.write(f"- {r}")
            else:
                st.markdown('<div style="text-align:center;padding:60px;background:var(--card);border:1px solid var(--line);"><p style="color:var(--text2);">분야를 입력하고<br>트렌드 분석을 시작하세요</p></div>', unsafe_allow_html=True)

    # ========== 탭3: 경쟁서 분석 ==========
    with tab3:
        st.markdown("### 경쟁 도서 분석")
        st.markdown('<p style="color:var(--text2);">경쟁 전자책의 목차, 리뷰, 강점을 분석합니다</p>', unsafe_allow_html=True)

        col1, col2 = st.columns([1, 1])

        with col1:
            st.markdown("#### 경쟁서 정보 입력")
            comp_title = st.text_input("책 제목", key="comp_title", placeholder="예: 돈의 심리학")
            comp_toc = st.text_area("목차 (복사/붙여넣기)", height=150, key="comp_toc", placeholder="1장. 제목\n2장. 제목\n...")
            comp_reviews = st.text_area("대표 리뷰 (선택)", height=100, key="comp_reviews", placeholder="인상적인 리뷰를 붙여넣으세요...")

            if st.button("경쟁서 분석", use_container_width=True, key="comp_btn"):
                if not comp_title or not comp_toc:
                    st.error("제목과 목차를 입력하세요")
                elif not get_api_key():
                    st.error("API 키를 입력해주세요")
                else:
                    with st.spinner("분석 중..."):
                        prompt = f"""다음 경쟁 도서를 분석해주세요:

제목: {comp_title}
목차:
{comp_toc}

리뷰: {comp_reviews if comp_reviews else '없음'}

JSON 형식으로 응답:
{{
    "book_summary": "이 책의 핵심 콘셉트",
    "target_audience": "예상 타겟 독자",
    "strengths": ["강점 1", "강점 2", "강점 3"],
    "weaknesses": ["약점/빈틈 1", "약점 2"],
    "unique_selling_point": "이 책만의 차별점",
    "improvement_opportunities": ["내 책에서 더 잘할 수 있는 것 1", "기회 2", "기회 3"],
    "key_chapters": ["핵심 챕터 1", "챕터 2"],
    "content_structure": "콘텐츠 구성 방식"
}}"""
                        result = ask_ai(prompt, 0.7)
                        parsed = parse_json(result)
                        if parsed:
                            parsed['title'] = comp_title
                            parsed['added_at'] = datetime.now().strftime('%Y-%m-%d %H:%M')
                            if 'competitor_analysis' not in st.session_state:
                                st.session_state['competitor_analysis'] = []
                            st.session_state['competitor_analysis'].append(parsed)
                            st.success("분석 완료")
                            st.rerun()

        with col2:
            st.markdown("#### 분석 결과")
            comps = st.session_state.get('competitor_analysis', [])

            if comps:
                for i, comp in enumerate(comps):
                    st.markdown(f"""<div class="data-card">
                        <b>{html.escape(str(comp.get('title', f'경쟁서 {i+1}')))}</b>
                        <br><small>{comp.get('added_at', '')}</small>
                    </div>""", unsafe_allow_html=True)
                    st.caption(comp.get('book_summary', ''))

                    if comp.get('improvement_opportunities'):
                        for o in comp.get('improvement_opportunities', [])[:2]:
                            st.success(f"차별화: {o}")

                    if st.button("삭제", key=f"del_comp_{i}"):
                        st.session_state['competitor_analysis'].pop(i)
                        st.rerun()
            else:
                st.markdown('<div style="text-align:center;padding:60px;background:var(--card);border:1px solid var(--line);"><p style="color:var(--text2);">경쟁 도서 정보를 입력하고<br>분석해보세요</p></div>', unsafe_allow_html=True)

    st.markdown('<div class="next-section"></div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 1, 1])
    with c1:
        if st.button("이전", key="p3_prev", use_container_width=True):
            go_prev()
            st.rerun()
    with c3:
        if st.button("다음", key="p3_next", use_container_width=True):
            go_next()
            st.rerun()


# ==========================================
# PAGE 4: 목차 설계
# ==========================================
elif current == 4:
    st.markdown("""
    <div class="section-title-box">
        <span class="section-step">STEP 05</span>
        <h2>목차 설계</h2>
        <p>독자의 호기심을 자극하는 목차를 만듭니다</p>
    </div>
    """, unsafe_allow_html=True)

    if st.session_state.get('market_gaps'):
        st.success(f"{len(st.session_state['market_gaps'])}개 차별화 포인트 반영")

    col1, col2 = st.columns([1, 1])

    with col1:
        st.markdown("### 목차 생성")

        st.markdown("""
        <div class="info-card">
            <b>🔥 목차 작성 팁</b><br><br>
            • 설명하지 말고 <b>궁금하게</b><br>
            • 구체적 <b>숫자 + 결과</b> 보여주기<br>
            • <b>실패담/고백</b>으로 공감 얻기<br>
            • "99%가 모르는" <b>비밀</b> 암시<br>
            • <b>반전</b>이 있을 것 같은 느낌<br><br>
            <span style="color:var(--gold);">❌ "시간관리의 중요성"</span><br>
            <span style="color:#50c878;">✓ "20대에 이걸 몰라서 5년 날렸다"</span>
        </div>
        """, unsafe_allow_html=True)

        if st.button("목차 생성하기", use_container_width=True, key="p4_outline_btn"):
            if not st.session_state.get('topic'):
                st.error("주제를 입력하세요")
            elif not get_api_key():
                st.error("API 키를 입력해주세요")
            else:
                with st.spinner("목차 생성 중..."):
                    result = generate_outline(
                        st.session_state['topic'],
                        st.session_state.get('target_persona', ''),
                        st.session_state.get('pain_points', ''),
                        st.session_state.get('market_gaps', [])
                    )

                    if result:
                        lines = result.split('\n')
                        chapters = []
                        current_ch = None
                        subtopics = {}

                        for line in lines:
                            orig_line = line
                            line = line.strip()
                            if not line:
                                continue

                            # 마크다운 정리 (먼저 정리한 후 검사)
                            clean_line = re.sub(r'^[#\*\s]+', '', line).strip()
                            clean_line = clean_line.replace('**', '').replace('*', '').strip()

                            # PART 또는 챕터 형식 인식 (더 유연하게)
                            is_chapter = False

                            # PART 형식 (다양한 변형)
                            if re.search(r'PART\s*\d+', clean_line, re.IGNORECASE):
                                is_chapter = True
                            # 파트 형식 (한글)
                            elif re.search(r'파트\s*\d+', clean_line):
                                is_chapter = True
                            # Chapter 형식
                            elif re.search(r'(Chapter|챕터)\s*\d+', clean_line, re.IGNORECASE):
                                is_chapter = True
                            # "1. 제목" 형식
                            elif re.match(r'^\d+[\.\)]\s*.+', clean_line) and not clean_line.startswith('-'):
                                is_chapter = True
                            # 숫자로 시작하는 제목 (예: "1 첫번째 파트")
                            elif re.match(r'^\d+\s+[가-힣A-Za-z]', clean_line):
                                is_chapter = True

                            if is_chapter:
                                name = clean_line
                                if name and len(name) > 3:
                                    current_ch = name
                                    chapters.append(current_ch)
                                    subtopics[current_ch] = []

                            # 소제목 - 다양한 형식 지원
                            elif current_ch:
                                is_subtopic = False
                                st_name = ""

                                # "-" 또는 "•" 또는 "·" 로 시작
                                if re.match(r'^\s*[\-\•\·]\s*', line):
                                    is_subtopic = True
                                    st_name = re.sub(r'^[\s\-\•\·]+', '', line).strip()
                                # 들여쓰기 된 내용
                                elif orig_line.startswith('  ') or orig_line.startswith('\t'):
                                    if not any(x in line.upper() for x in ['PART', 'CHAPTER', '파트']):
                                        is_subtopic = True
                                        st_name = line.strip().lstrip('-•· ')
                                # "  1)" 또는 "  a)" 형식
                                elif re.match(r'^\s+[\da-z][\)\.]', orig_line):
                                    is_subtopic = True
                                    st_name = re.sub(r'^[\s\da-z\)\.\-]+', '', line).strip()

                                if is_subtopic:
                                    st_name = st_name.replace('**', '').replace('*', '').replace('#', '').strip()
                                    # 소제목이 충분히 길고 유효한 경우만 추가
                                    if st_name and len(st_name) > 3 and not re.match(r'^(PART|파트|Chapter|챕터)', st_name, re.IGNORECASE):
                                        subtopics[current_ch].append(st_name)

                        if chapters:
                            st.session_state['outline'] = chapters
                            st.session_state['chapters'] = {}
                            for ch in chapters:
                                st.session_state['chapters'][ch] = {
                                    'subtopics': subtopics.get(ch, []),
                                    'subtopic_data': {s: {'questions': [], 'answers': [], 'content': ''} for s in subtopics.get(ch, [])}
                                }
                            st.success(f"{len(chapters)}개 챕터 생성!")
                            st.rerun()
                        else:
                            st.error("목차 생성 실패. 다시 시도해주세요.")
                    else:
                        st.error("AI 응답 없음. 다시 시도해주세요.")

    with col2:
        st.markdown("### 현재 목차")

        if st.session_state.get('outline'):
            # 수정 모드 토글
            if 'edit_outline_mode' not in st.session_state:
                st.session_state['edit_outline_mode'] = False

            col_view, col_edit = st.columns([1, 1])
            with col_view:
                if st.button("👁 보기 모드", use_container_width=True, disabled=not st.session_state['edit_outline_mode']):
                    st.session_state['edit_outline_mode'] = False
                    st.rerun()
            with col_edit:
                if st.button("✏️ 수정 모드", use_container_width=True, disabled=st.session_state['edit_outline_mode']):
                    st.session_state['edit_outline_mode'] = True
                    st.rerun()

            st.markdown("---")

            if st.session_state['edit_outline_mode']:
                # 수정 모드
                st.markdown('<p style="color:var(--gold);font-size:14px;">📝 제목을 직접 수정할 수 있습니다</p>', unsafe_allow_html=True)

                updated_outline = []
                updated_chapters = {}

                for ch_idx, ch in enumerate(st.session_state['outline']):
                    # 챕터 제목 수정
                    new_ch_title = st.text_input(
                        f"PART {ch_idx + 1}",
                        value=ch,
                        key=f"edit_ch_{ch_idx}"
                    )
                    updated_outline.append(new_ch_title)
                    updated_chapters[new_ch_title] = {'subtopics': [], 'subtopic_data': {}}

                    # 소제목 수정
                    subtopics = st.session_state['chapters'].get(ch, {}).get('subtopics', [])
                    new_subtopics = []
                    for st_idx, st_name in enumerate(subtopics):
                        new_st = st.text_input(
                            f"  └ 소제목 {st_idx + 1}",
                            value=st_name,
                            key=f"edit_st_{ch_idx}_{st_idx}",
                            label_visibility="collapsed"
                        )
                        if new_st.strip():
                            new_subtopics.append(new_st)
                            # 기존 데이터 유지
                            old_data = st.session_state['chapters'].get(ch, {}).get('subtopic_data', {}).get(st_name, {'questions': [], 'answers': [], 'content': ''})
                            updated_chapters[new_ch_title]['subtopic_data'][new_st] = old_data

                    updated_chapters[new_ch_title]['subtopics'] = new_subtopics
                    st.markdown("---")

                # 저장 버튼
                if st.button("💾 수정 내용 저장", use_container_width=True, type="primary"):
                    st.session_state['outline'] = updated_outline
                    st.session_state['chapters'] = updated_chapters
                    st.session_state['edit_outline_mode'] = False
                    st.success("목차가 수정되었습니다!")
                    st.rerun()

            else:
                # 보기 모드 - 예쁘게 표시
                for ch_idx, ch in enumerate(st.session_state['outline']):
                    st.markdown(f"""
                    <div style="background:linear-gradient(135deg, rgba(201,162,75,0.15) 0%, rgba(201,162,75,0.05) 100%);
                                padding:16px 20px;border-radius:12px;margin-bottom:8px;border-left:4px solid var(--gold);">
                        <span style="color:var(--gold);font-size:13px;font-weight:600;">PART {ch_idx + 1}</span>
                        <p style="color:var(--text);font-size:17px;font-weight:600;margin:8px 0 0 0;">{ch}</p>
                    </div>
                    """, unsafe_allow_html=True)

                    subtopics = st.session_state['chapters'].get(ch, {}).get('subtopics', [])
                    for st_idx, st_name in enumerate(subtopics):
                        st.markdown(f"""
                        <div style="padding:10px 20px 10px 35px;color:var(--text);font-size:15px;">
                            <span style="color:var(--gold);margin-right:8px;">•</span>{st_name}
                        </div>
                        """, unsafe_allow_html=True)

                    st.markdown("<div style='margin-bottom:20px;'></div>", unsafe_allow_html=True)

        else:
            st.markdown('<div style="text-align:center;padding:60px;background:rgba(255,255,255,0.03);border-radius:16px;border:1px solid rgba(201,162,75,0.15);"><p style="color:rgba(255,255,255,0.5);">목차를 생성해주세요</p></div>', unsafe_allow_html=True)

    st.markdown('<div class="next-section"></div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 1, 1])
    with c1:
        if st.button("이전", key="p4_prev", use_container_width=True):
            go_prev()
            st.rerun()
    with c3:
        if st.button("다음 본문", key="p4_next", use_container_width=True):
            go_next()
            st.rerun()


# ==========================================
# PAGE 5: 본문 작성
# ==========================================
elif current == 5:
    st.markdown("""
    <div class="section-title-box">
        <span class="section-step">STEP 06</span>
        <h2>본문 작성</h2>
        <p>AI가 각 챕터의 콘텐츠를 작성합니다</p>
    </div>
    """, unsafe_allow_html=True)

    if not st.session_state.get('outline'):
        st.warning("먼저 목차를 설계하세요")
    else:
        col_sel1, col_sel2 = st.columns([1, 1])
        with col_sel1:
            selected_ch = st.selectbox("챕터", st.session_state['outline'], key="p5_chapter")

        # 선택된 챕터가 있고 chapters에 존재하는지 확인
        if selected_ch and selected_ch in st.session_state.get('chapters', {}):
            ch_data = st.session_state['chapters'][selected_ch]
            subtopics_list = ch_data.get('subtopics', [])

            # 소제목이 있는 경우에만 선택박스 표시
            selected_st = None
            if subtopics_list:
                with col_sel2:
                    selected_st = st.selectbox("소제목", subtopics_list, key="p5_subtopic")

            # 진행률 표시
            completed = sum(1 for s in subtopics_list if ch_data.get('subtopic_data', {}).get(s, {}).get('content'))
            total = len(subtopics_list)
            if total > 0:
                st.progress(completed / total)
                st.caption(f"{completed}/{total} 완료")

            # 소제목이 선택된 경우에만 편집 UI 표시
            if selected_st:
                # subtopic_data 초기화 확인
                if 'subtopic_data' not in ch_data:
                    ch_data['subtopic_data'] = {}
                if selected_st not in ch_data['subtopic_data']:
                    ch_data['subtopic_data'][selected_st] = {'questions': [], 'answers': [], 'content': ''}

                st_data = ch_data['subtopic_data'][selected_st]

                col1, col2 = st.columns([1, 1])

                # 버튼 키를 위한 고유 식별자
                st_key = f"{selected_ch}_{selected_st}".replace(" ", "_")

                with col1:
                    st.markdown("### 인터뷰")
                    if st.button("질문 생성", key=f"gen_q_{st_key}"):
                        if not get_api_key():
                            st.error("API 키를 입력해주세요")
                        else:
                            with st.spinner("생성 중..."):
                                q_text = generate_questions(selected_st, selected_ch, st.session_state['topic'])
                                if q_text:
                                    questions = re.findall(r'Q\d+:\s*(.+)', q_text)
                                    if not questions:
                                        questions = [q.strip() for q in q_text.split('\n') if '?' in q][:3]
                                    if questions:
                                        st_data['questions'] = questions
                                        st_data['answers'] = [''] * len(questions)
                                        st.rerun()
                                    else:
                                        st.error("질문 생성에 실패했습니다")

                    if st_data.get('questions'):
                        for i, q in enumerate(st_data['questions']):
                            st.markdown(f"**Q{i+1}.** {q}")
                            # answers 리스트 크기 확인
                            while len(st_data.get('answers', [])) <= i:
                                st_data['answers'].append('')
                            st_data['answers'][i] = st.text_area(f"A{i+1}", value=st_data['answers'][i], height=80, key=f"ans_{st_key}_{i}", label_visibility="collapsed")

                with col2:
                    st.markdown("### 본문")
                    has_ans = st_data.get('questions') and any(a.strip() for a in st_data.get('answers', []))

                    if has_ans:
                        if st.button("본문 생성", key=f"gen_content_{st_key}", use_container_width=True, type="primary"):
                            if not get_api_key():
                                st.error("API 키를 입력해주세요")
                            else:
                                with st.spinner("본문 작성 중... (1~2분 소요)"):
                                    content = generate_content_premium(selected_st, selected_ch, st_data['questions'], st_data['answers'], st.session_state['topic'], st.session_state['target_persona'])
                                    if content:
                                        st_data['content'] = content
                                        st.success("본문 생성 완료!")
                                        st.rerun()
                                    else:
                                        st.error("본문 생성에 실패했습니다. 다시 시도해주세요.")
                    else:
                        st.info("왼쪽에서 질문에 답변을 입력하면 본문을 생성할 수 있습니다")

                    # 본문 표시
                    current_content = st_data.get('content', '')
                    if current_content:
                        # HTML 형식으로 변환하여 표시
                        formatted_html = format_content_html(current_content)
                        st.markdown(f"""
                        <style>
                        .content-preview-box {{
                            background:#ffffff !important;
                            padding:25px 30px;
                            border-radius:12px;
                            border:1px solid rgba(201,162,75,0.3);
                            margin:15px 0;
                            font-family:'S-CoreDream', sans-serif !important;
                            font-size:17px;
                            max-height:500px;
                            overflow-y:auto;
                        }}
                        .content-preview-box,
                        .content-preview-box p,
                        .content-preview-box span,
                        .content-preview-box div {{
                            color:#000000 !important;
                            -webkit-text-fill-color:#000000 !important;
                        }}
                        .content-preview-box b[style*="color:#e67e22"],
                        .content-preview-box p[style*="color:#e67e22"] {{
                            color:#e67e22 !important;
                            -webkit-text-fill-color:#e67e22 !important;
                        }}
                        </style>
                        <div class="content-preview-box">
                            {formatted_html}
                        </div>
                        """, unsafe_allow_html=True)
                        st.caption(f"📝 {len(current_content.replace(' ', '').replace(chr(10), '')):,}자")

                        # 이미지 추가 기능
                        st.markdown("---")
                        st.markdown("**📷 이미지 추가**")
                        uploaded_img = st.file_uploader("이미지 업로드", type=['png', 'jpg', 'jpeg'], key=f"img_{st_key}", label_visibility="collapsed")
                        if uploaded_img:
                            # 이미지 저장
                            if 'images' not in st_data:
                                st_data['images'] = []
                            img_b64 = base64.b64encode(uploaded_img.read()).decode()
                            st_data['images'].append({'name': uploaded_img.name, 'data': img_b64})
                            st.success(f"이미지 '{uploaded_img.name}' 추가됨!")
                            st.rerun()

                        # 추가된 이미지 표시
                        if st_data.get('images'):
                            st.caption(f"추가된 이미지: {len(st_data['images'])}개")
                            for idx, img in enumerate(st_data['images']):
                                col_img, col_del = st.columns([4, 1])
                                with col_img:
                                    st.image(f"data:image/png;base64,{img['data']}", caption=img['name'], width=200)
                                with col_del:
                                    if st.button("삭제", key=f"del_img_{st_key}_{idx}"):
                                        st_data['images'].pop(idx)
                                        st.rerun()

                        # 수정 기능
                        st.markdown("---")
                        with st.expander("✏️ 본문 직접 수정"):
                            st.caption("「중요단어」 → 주황색 강조 | ★ 문장 → 핵심 강조")
                            edited = st.text_area("본문 편집", value=current_content, height=400, key=f"content_{st_key}", label_visibility="collapsed")
                            if edited != current_content:
                                st_data['content'] = edited
                                st.rerun()
                    else:
                        st.markdown('<div style="text-align:center;padding:80px 20px;background:rgba(255,255,255,0.03);border-radius:12px;border:1px dashed rgba(201,162,75,0.3);"><p style="color:var(--text2);font-size:16px;">본문이 아직 없습니다<br>질문에 답변 후 "본문 생성" 버튼을 누르세요</p></div>', unsafe_allow_html=True)
            else:
                st.info("이 챕터에는 소제목이 없습니다. 목차를 다시 생성해주세요.")

        st.markdown("---")
        st.markdown("### 전체 본문")
        full_content = get_full_content()
        if full_content:
            char_count = len(full_content.replace(' ', '').replace('\n', ''))
            est_pages = estimate_docx_pages(
                st.session_state.get('chapters', {}),
                st.session_state.get('outline', [])
            )
            st.success(f"총 {char_count:,}자 | 약 {est_pages}페이지 (WORD A5 기준)")

    st.markdown('<div class="next-section"></div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 1, 1])
    with c1:
        if st.button("이전", key="p5_prev", use_container_width=True):
            go_prev()
            st.rerun()
    with c3:
        if st.button("다음 출력", key="p5_next", use_container_width=True):
            go_next()
            st.rerun()


# ==========================================
# PAGE 6: 표지 디자인
# ==========================================
elif current == 6:
    st.markdown("""
    <div class="section-title-box">
        <span class="section-step">STEP 07</span>
        <h2>표지 디자인</h2>
        <p>전문 디자인 툴로 고품질 표지를 만드세요</p>
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns([1, 1])

    with col1:
        st.markdown("### 표지 정보 정리")

        # 이전 페이지에서 설정한 제목/부제 자동 연동
        saved_title = st.session_state.get('book_title', '')
        saved_subtitle = st.session_state.get('subtitle', '')

        cover_title = st.text_input("표지 제목", value=saved_title, key="cover_title", placeholder="예: 돈의 속성")
        cover_subtitle = st.text_input("부제목", value=saved_subtitle, key="cover_subtitle", placeholder="예: 당신이 모르는 부의 법칙")
        cover_author = st.text_input("저자명", key="cover_author", placeholder="예: 홍길동")

        st.markdown("---")
        st.markdown("### AI 표지 스타일 추천")

        if st.button("내 주제에 맞는 표지 스타일 추천받기", use_container_width=True, key="ai_cover_suggest"):
            topic = st.session_state.get('topic', '')
            if not topic:
                st.error("시장분석 페이지에서 주제를 먼저 입력해주세요")
            elif not get_api_key():
                st.error("API 키를 입력해주세요")
            else:
                with st.spinner("베스트셀러 표지 분석 중..."):
                    prompt = f"""'{topic}' 주제의 전자책 표지 디자인을 추천해주세요.

이 분야의 실제 베스트셀러 책 표지를 분석해서 추천해주세요.

JSON 형식으로 응답:
{{
    "recommended_style": "추천 스타일명",
    "color_scheme": "추천 색상 조합 (예: 검정 배경 + 금색 텍스트)",
    "design_concept": "디자인 콘셉트 설명 (2문장)",
    "typography_tip": "타이포그래피 팁 (폰트 스타일, 크기 등)",
    "reference_books": ["참고할 베스트셀러 표지 1", "표지 2", "표지 3"],
    "canva_search_keyword": "Canva에서 검색할 키워드 (영문)"
}}"""
                    result = ask_ai(prompt, 0.7)
                    parsed = parse_json(result)
                    if parsed:
                        st.session_state['cover_suggestion'] = parsed
                        st.rerun()

        if st.session_state.get('cover_suggestion'):
            sug = st.session_state['cover_suggestion']
            st.markdown(f"""<div class="data-card">
                <b>추천 스타일: {html.escape(str(sug.get('recommended_style', '')))}</b><br>
                <small>색상: {html.escape(str(sug.get('color_scheme', '')))}</small>
            </div>""", unsafe_allow_html=True)
            st.write(sug.get('design_concept', ''))
            if sug.get('typography_tip'):
                st.info(f"💡 타이포그래피 팁: {sug.get('typography_tip', '')}")
            if sug.get('reference_books'):
                st.markdown("**참고 베스트셀러:**")
                for book in sug.get('reference_books', []):
                    st.caption(f"- {book}")
            if sug.get('canva_search_keyword'):
                st.session_state['canva_keyword'] = sug.get('canva_search_keyword', '')

    with col2:
        st.markdown("### 표지 미리보기")

        _label_to_id = {v: k for k, v in COVER_TEMPLATES.items()}
        _cover_labels = list(_label_to_id.keys())
        if st.session_state.get("cover_template_choice") not in _cover_labels:
            st.session_state["cover_template_choice"] = _cover_labels[0]
        _choice = st.radio(
            "표지 스타일",
            _cover_labels,
            key="cover_template_choice",
        )
        _tmpl = _label_to_id[_choice]

        _svg = build_cover_svg(
            _tmpl,
            cover_title or "제목을 입력하세요",
            cover_subtitle,
            cover_author,
        )
        st.markdown(
            '<div style="max-width:320px;margin:12px auto 18px;border-radius:8px;'
            'overflow:hidden;box-shadow:0 24px 70px rgba(0,0,0,0.55);">'
            f'{_svg}</div>',
            unsafe_allow_html=True,
        )
        st.download_button(
            "표지 다운로드 (SVG · 고해상도 벡터)",
            _svg,
            file_name=f"{(cover_title or 'cover')}_cover.svg",
            mime="image/svg+xml",
            use_container_width=True,
            key="cover_svg_dl",
        )
        st.caption(
            "SVG는 무손실 벡터 파일입니다. 브라우저에서 열어 캡처하거나, "
            "Canva·Figma·미리캔버스에 올려 PNG/JPG로 내보낼 수 있습니다."
        )
        st.caption(
            "베스트셀러 브리핑 스타일은 cover_photo.jpg/png 또는 author_photo.jpg/png를 "
            "앱 파일 옆이나 assets 폴더에 두면 하단 인물 사진까지 자동 반영됩니다."
        )

        st.markdown("---")
        st.markdown("### Canva로 표지 만들기")

        st.markdown("""
        <div class="data-card">
            <p style="font-size:16px;margin-bottom:15px;">
                <b>Canva</b>는 전문 디자이너 수준의 표지를 무료로 만들 수 있는 온라인 툴입니다.
            </p>
            <p style="color:var(--text2);font-size:14px;">
                ✓ 수천 개의 프로 템플릿<br>
                ✓ 드래그 앤 드롭 편집<br>
                ✓ 무료 이미지/아이콘<br>
                ✓ 한글 폰트 지원
            </p>
        </div>
        """, unsafe_allow_html=True)

        st.markdown("---")

        # Canva 검색 키워드 설정
        canva_keyword = st.session_state.get('canva_keyword', 'book cover')

        # Canva 책 표지 템플릿 링크
        canva_url = f"https://www.canva.com/templates/?query={canva_keyword}%20book%20cover"

        st.markdown(f"""
        <a href="{canva_url}" target="_blank" style="
            display:block;
            background:linear-gradient(135deg,#7c3aed,#6366f1);
            color:white;
            padding:18px 24px;
            border-radius:12px;
            text-decoration:none;
            text-align:center;
            font-size:18px;
            font-weight:600;
            margin-bottom:15px;
            transition:transform 0.2s;
        ">
            🎨 Canva에서 표지 만들기
        </a>
        """, unsafe_allow_html=True)

        st.markdown("---")
        st.markdown("### 사용 방법")
        st.markdown("""
        1. **Canva 열기** - 위 버튼 클릭 (무료 가입)
        2. **템플릿 선택** - 마음에 드는 디자인 클릭
        3. **텍스트 수정** - 제목, 부제목, 저자명 입력
        4. **다운로드** - PNG 또는 PDF로 저장
        """)

        st.markdown("---")
        st.markdown("### 추천 검색어")

        search_keywords = [
            "ebook cover", "book cover minimalist",
            "book cover gold", "book cover business",
            "korean book cover", "self help book cover"
        ]

        cols = st.columns(2)
        for i, kw in enumerate(search_keywords):
            with cols[i % 2]:
                if st.button(kw, key=f"canva_kw_{i}", use_container_width=True):
                    st.session_state['canva_keyword'] = kw
                    st.rerun()

        st.markdown("---")

        # 복사할 텍스트
        if cover_title or cover_subtitle or cover_author:
            st.markdown("### 복사할 텍스트")
            copy_text = f"제목: {cover_title}\n부제목: {cover_subtitle}\n저자: {cover_author}"
            st.code(copy_text, language=None)

    st.markdown('<div class="next-section"></div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 1, 1])
    with c1:
        if st.button("이전", key="p6_prev", use_container_width=True):
            go_prev()
            st.rerun()
    with c3:
        if st.button("다음", key="p6_next", use_container_width=True):
            go_next()
            st.rerun()


# ==========================================
# PAGE 7: 최종 출력
# ==========================================
elif current == 7:
    st.markdown("""
    <div class="section-title-box">
        <span class="section-step">STEP 08</span>
        <h2>최종 출력</h2>
        <p>완성된 전자책을 다운로드하세요</p>
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns([1.5, 1])

    with col1:
        st.markdown("### 다운로드")

        final_title = st.text_input("제목", value=st.session_state.get('book_title', ''), key="p6_title")
        final_subtitle = st.text_input("부제", value=st.session_state.get('subtitle', ''), key="p6_subtitle")

        full = f"{final_title}\n{final_subtitle}\n\n{'='*50}\n\n"
        for ch in st.session_state.get('outline', []):
            if ch in st.session_state.get('chapters', {}):
                ch_data = st.session_state['chapters'][ch]
                ch_content = ""
                for s in ch_data.get('subtopics', []):
                    c = ch_data.get('subtopic_data', {}).get(s, {}).get('content', '')
                    if c:
                        ch_content += f"\n\n【{s}】\n\n{clean_content(c)}"
                if ch_content:
                    full += f"\n\n{ch}\n{'-'*40}{ch_content}\n"

        st.markdown("**미리보기**")
        st.text_area("전체 내용", value=full, height=300, disabled=True, key="p7_preview")

        # 저자명 가져오기
        author_name = st.session_state.get('author_name', '') or st.session_state.get('interview_data', {}).get('author_name', '')

        # 다운로드 버튼 3개
        st.markdown("### 📥 다운로드")

        c1, c2, c3 = st.columns(3)
        with c1:
            st.download_button("📄 TXT", full, file_name=f"{final_title or 'ebook'}.txt", use_container_width=True, key="p7_txt")
        with c2:
            # HTML 내보내기 - 특수문자 이스케이프 처리
            escaped_title = html.escape(final_title)
            escaped_content = html.escape(full).replace('\n', '<br>')
            html_content = f"""<!DOCTYPE html>
<html lang="ko">
<head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{escaped_title}</title>
    <style>
        @import url('https://cdn.jsdelivr.net/gh/orioncactus/pretendard/dist/web/static/pretendard.css');
        @import url('https://fonts.googleapis.com/css2?family=Noto+Serif+KR:wght@500;600;700;900&display=swap');
        * {{ box-sizing: border-box; }}
        body {{
            max-width: 860px;
            margin: 0 auto;
            padding: 80px 44px;
            font-family: 'Noto Serif KR', 'Nanum Myeongjo', serif;
            line-height: 2.05;
            background: #f6f0e2;
            color: #211c17;
            word-break: keep-all;
        }}
        body::before {{
            content: "";
            display: block;
            height: 360px;
            margin: -20px 0 54px;
            border: 1px solid #b28734;
            background: linear-gradient(135deg, #211c17, #100f0d);
            box-shadow: 0 26px 70px rgba(33, 28, 23, 0.22);
        }}
        body::after {{
            content: "CASHMAKER";
            display: block;
            margin-top: 72px;
            padding-top: 26px;
            border-top: 1px solid #d9cfbc;
            font-family: 'Pretendard', sans-serif;
            font-size: 12px;
            font-weight: 800;
            letter-spacing: 0.16em;
            color: #9b4830;
            text-align: center;
        }}
        h1, h2, h3 {{
            font-family: 'Noto Serif KR', serif;
            color: #211c17;
            line-height: 1.24;
        }}
        h1 {{
            font-size: 42px;
            margin: 0 0 12px;
            letter-spacing: -0.01em;
        }}
        h2 {{
            font-size: 22px;
            margin: 48px 0 18px;
            padding-bottom: 10px;
            border-bottom: 2px solid #b28734;
        }}
        p {{
            margin: 0 0 1.05em;
        }}
        @media (max-width: 640px) {{
            body {{ padding: 54px 22px; }}
            body::before {{ height: 250px; }}
            h1 {{ font-size: 32px; }}
        }}
    </style>
</head>
<body>
{escaped_content}
</body>
</html>"""
            st.download_button("🌐 HTML", html_content, file_name=f"{final_title or 'ebook'}.html", use_container_width=True, key="p7_html")

        with c3:
            # DOCX 다운로드
            if DOCX_AVAILABLE:
                docx_data, docx_error = create_ebook_docx(
                    final_title,
                    final_subtitle,
                    author_name,
                    st.session_state.get('chapters', {}),
                    st.session_state.get('outline', []),
                    st.session_state.get('interview_data', {})
                )
                if docx_data:
                    st.download_button(
                        "📘 WORD",
                        docx_data,
                        file_name=f"{final_title or 'ebook'}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        key="p7_docx"
                    )
                else:
                    st.button("📘 WORD", disabled=True, use_container_width=True, key="p7_docx_disabled")
                    if docx_error:
                        st.caption(f"⚠️ {docx_error[:30]}")
            else:
                st.button("📘 WORD", disabled=True, use_container_width=True, key="p7_docx_na")
                st.caption("pip install python-docx")

        total = len(full.replace(' ', '').replace('\n', ''))
        if total > 0:
            est_pages = estimate_docx_pages(
                st.session_state.get('chapters', {}),
                st.session_state.get('outline', [])
            )
            st.success(f"총 {total:,}자 | 약 {est_pages}페이지 (WORD A5 기준)")
            st.caption("표지·판권·프롤로그·에필로그·저자소개와 소제목별 페이지 나눔까지 반영한 추정치입니다. 워드의 글꼴·자동 줄나눔에 따라 1~2페이지 오차가 있을 수 있습니다.")

    with col2:
        st.markdown("### 현황")
        total_st = sum(len(ch.get('subtopics', [])) for ch in st.session_state.get('chapters', {}).values())
        done = sum(1 for ch in st.session_state.get('chapters', {}).values() for s in ch.get('subtopic_data', {}).values() if s.get('content'))

        if total_st > 0:
            st.progress(done / total_st)
            st.write(f"**완료:** {done}/{total_st}")

    st.markdown('<div class="next-section"></div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 2, 1])
    with c1:
        if st.button("이전", key="p7_prev", use_container_width=True):
            go_prev()
            st.rerun()


st.markdown("""
<div class="app-footer">
    <b>CASHMAKER</b> | 제작: 남현우 작가
</div>
""", unsafe_allow_html=True)
